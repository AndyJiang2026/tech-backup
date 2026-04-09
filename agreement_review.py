#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
业务约定书审查 - 在乙方原文基础上用红色标注修改建议
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def add_run(paragraph, text, color='black', bold=False, font_size=11):
    """添加带颜色的文本片段"""
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color == 'red':
        run.font.color.rgb = RGBColor(255, 0, 0)
    else:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return run

def create_agreement_review():
    """创建审查后的协议文档"""
    doc = Document()
    
    # 设置文档样式
    doc.styles['Normal'].font.name = 'SimSun'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    doc.styles['Normal'].font.size = Pt(11)
    
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('业务约定书（审查修改版）')
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    
    # 说明
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note_run = note.add_run('说明：黑色为乙方原文，红色为修改建议')
    note_run.font.size = Pt(10)
    note_run.font.color.rgb = RGBColor(255, 0, 0)
    note_run.font.italic = True
    
    doc.add_paragraph()
    
    # 标题
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '业务约定书', 'black', True, 14)
    
    doc.add_paragraph()
    
    # 甲方信息
    p = doc.add_paragraph()
    add_run(p, '委托方（甲方）：', 'black', True)
    add_run(p, '未来新视界科技（北京）有限公司', 'black')
    
    p = doc.add_paragraph()
    add_run(p, '联系人：', 'black', True)
    add_run(p, '请填写', 'black')
    
    p = doc.add_paragraph()
    add_run(p, '联系电话：', 'black', True)
    add_run(p, '请填写企业固定电话', 'black')
    
    p = doc.add_paragraph()
    add_run(p, '手机：', 'black', True)
    add_run(p, '请填写', 'black')
    
    p = doc.add_paragraph()
    add_run(p, '联系地址：', 'black', True)
    add_run(p, '请填写企业实际办公地址', 'black')
    
    doc.add_paragraph()
    
    # 乙方信息
    p = doc.add_paragraph()
    add_run(p, '受托方（乙方）：', 'black', True)
    add_run(p, '北京一嘉二科技服务有限公司', 'black')
    
    p = doc.add_paragraph()
    add_run(p, '联系人：', 'black', True)
    add_run(p, '黄明兰', 'black')
    
    p = doc.add_paragraph()
    add_run(p, '联系电话：', 'black', True)
    add_run(p, '010-58030270', 'black')
    
    p = doc.add_paragraph()
    add_run(p, '手机：', 'black', True)
    add_run(p, '18310808734', 'black')
    
    p = doc.add_paragraph()
    add_run(p, '联系地址：', 'black', True)
    add_run(p, '北京市丰台区马家堡东路 106 号自然新天地 A 座 1010-1011', 'black')
    
    doc.add_paragraph()
    
    # 前言
    p = doc.add_paragraph()
    add_run(p, '依据《中华人民共和国民法典》，在甲乙双方保证其主体合法的基础上，甲方委托乙方就申报政府资助资金、资质类项目提供咨询服务，为促使项目申报成功，甲、乙双方本着平等自愿、互惠互利的原则，达成如下协议，双方共同恪守。', 'black')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '乙方协助甲方申请项目为：', 'black')
    
    p = doc.add_paragraph()
    add_run(p, '1、朝阳区促进文化产业高质量发展的若干措施支持项目', 'black')
    
    p = doc.add_paragraph()
    add_run(p, '2、国家电影局电影精品专项资金资助项目', 'black')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # ========== 第一条 ==========
    p = doc.add_paragraph()
    add_run(p, '第一条、甲方的权利和义务', 'black', True, 12)
    
    p = doc.add_paragraph()
    add_run(p, '1、甲方应及时完整地向乙方提供必要的相关基础材料，保证所提交的材料及文件内容真实、合法且有效，并依法承担由此产生的全部法律责任。', 'black')
    
    p = doc.add_paragraph()
    add_run(p, '2、配合乙方做好与相关第三方的沟通事宜，协助乙方办理双方认为其他必要的事宜。', 'black')
    
    p = doc.add_paragraph()
    add_run(p, '3、甲方应支付乙方咨询顾问费，在材料申报过程中，如有涉及财务审计、专利或著作权申请、用户报告、查新等第三方的所有费用，应由甲方自行委托并独立承担，与应支付乙方的咨询顾问费用无关。', 'black')
    
    p = doc.add_paragraph()
    add_run(p, '4、甲方指定的项目联系人，配合乙方开展相关工作、提供资料等，如甲方更换项目联系人，应当及时通知乙方。', 'black')
    
    p = doc.add_paragraph()
    add_run(p, '5、乙方为甲方上述资金申请的唯一合作人，甲方须严格为乙方的服务内容及双方合作事宜保密。', 'black')
    add_run(p, '\n【修改建议：删除"唯一合作人"限制，改为"乙方为甲方上述资金申请的主要合作方之一"】', 'red')
    
    doc.add_paragraph()
    
    # ========== 第二条 ==========
    p = doc.add_paragraph()
    add_run(p, '第二条、乙方的权利和义务', 'black', True, 12)
    
    p = doc.add_paragraph()
    add_run(p, '1、在甲方提供真实、完整且与公司实际经营情况一致的基础数据及材料的前提下，乙方为甲方申报上述项目提供咨询服务。', 'black')
    
    p = doc.add_paragraph()
    add_run(p, '2、乙方应尽职完成本合同约定的咨询服务内容，维护甲方利益。', 'black')
    add_run(p, '\n【修改建议：增加"乙方应保证申报材料的准确性和专业性，因乙方过错导致申报失败的，应承担相应责任"】', 'red')
    
    p = doc.add_paragraph()
    add_run(p, '3、甲方有权就乙方服务范围内的事项，向乙方提出口头或书面询问，乙方应及时作出答复。', 'black')
    add_run(p, '\n【修改建议：明确"及时"为"2 个工作日内"】', 'red')
    
    p = doc.add_paragraph()
    add_run(p, '4、由于项目申报在流程上有一定的阶段性，在每个阶段的工作中，乙方有责任配合甲方完成申报工作，向申报部门提交全部申报材料后即视为乙方已完成全部咨询服务内容，如在合同履行过程中甲方对乙方的工作表示不满，应及时向乙方提出，乙方应及时作出解释或整改。但甲方不得在成功申请到政府资助金后，又以乙方履行瑕疵为由提出拒付或少付咨询顾问费。', 'black')
    add_run(p, '\n【修改建议：删除"但甲方不得..."条款，改为"如乙方服务存在重大瑕疵导致甲方损失的，甲方有权要求相应赔偿"】', 'red')
    
    p = doc.add_paragraph()
    add_run(p, '5、乙方有义务严格遵守顾问服务的行业标准，保护甲方的合法权益，严守甲方的机密，严格为双方合作事宜保密。', 'black')
    
    doc.add_paragraph()
    
    # ========== 第三条 ==========
    p = doc.add_paragraph()
    add_run(p, '第三条、佣金及付款方式', 'black', True, 12)
    
    p = doc.add_paragraph()
    add_run(p, '1、咨询顾问费按项目立项资金总额的 10% 支付。', 'black')
    add_run(p, '\n【修改建议：改为"8%"，或设置阶梯费率（500 万以下 10%，500 万以上 8%）】', 'red')
    
    p = doc.add_paragraph()
    add_run(p, '2、合同一经签订生效，项目即开始运作。', 'black')
    
    p = doc.add_paragraph()
    add_run(p, '3、付款时间：成功申请到政府资助资金后，甲方承诺，在收到资助项目拨款后的五个工作日内，按政府拨款数额的 10%，支付乙方咨询顾问费。若政府资助金分批支付的，则甲方每收到一笔政府资助金，应在收到后五个工作日内，按该金额的 10% 向乙方支付，直至全部付清。', 'black')
    add_run(p, '\n【修改建议：将"五个工作日"改为"十五个工作日"，给予甲方更充足的财务处理时间】', 'red')
    
    doc.add_paragraph()
    
    # ========== 第四条 ==========
    p = doc.add_paragraph()
    add_run(p, '第四条、违约责任', 'black', True, 12)
    
    p = doc.add_paragraph()
    add_run(p, '本合同签订后，双方均应履行合同规定的全部条款，否则，违约方应赔偿给守约方因此而造成的一切损失。', 'black')
    
    p = doc.add_paragraph()
    add_run(p, '1、如果乙方工作人员在整理甲方的申报材料时发现甲方项目和企业条件有严重不足而无法申报本合同规定的项目时，双方可协商申报另外的科技计划或专项资金项目。若因自然灾害、政府政策调整等不可抗力导致合同无法履行，双方可延期或终止合同，不承担违约责任。', 'black')
    
    p = doc.add_paragraph()
    add_run(p, '2、合同签订生效后，如乙方无正当理由，擅自中止对甲方的服务，则应赔偿由此造成的损失；如甲方单方面停止履行合同，并自行进行申报，视同本合同仍然生效，仍须按上述支付费用条款支付乙方的咨询顾问费。', 'black')
    add_run(p, '\n【修改建议：增加"如乙方未按约定提供服务或服务质量严重不达标，甲方有权解除合同并要求退还已支付费用"】', 'red')
    
    p = doc.add_paragraph()
    add_run(p, '3、成功申请到政府资助金后，如甲方拒绝接受政府资助，仍视同乙方已完成合同服务内容，甲方仍须支付上述咨询顾问费，按照同批其他企业的最高立项金额在第一批拨款后的十个工作日内，一次性付清。', 'black')
    add_run(p, '\n【修改建议：删除此条款，或改为"如甲方无正当理由拒绝接受资助，应按实际获批金额的 50% 支付咨询费"】', 'red')
    
    p = doc.add_paragraph()
    add_run(p, '4、如甲方逾期付款，乙方将向甲方另行加收利息，利息按应付款额每日万分之八计算，直至付清。', 'black')
    add_run(p, '\n【修改建议：改为"万分之三"，万分之八年化约 29%，过高】', 'red')
    
    p = doc.add_paragraph()
    add_run(p, '5、甲方每收到一笔本项目的政府资助金后，应及时告知乙方（不超过 3 个工作日）收到款项的日期和数额，如果乙方向甲方询问关于本项目政府资助金回款情况，甲方拒绝告知的，则乙方有权随时要求甲方按原始申请资金额度支付本项目全部的咨询顾问费，甲方不予支付的，视为甲方逾期付款，应承担相应违约责任。', 'black')
    add_run(p, '\n【修改建议：删除"按原始申请资金额度支付全部咨询费"条款，改为"甲方应如实告知，如故意隐瞒，乙方有权通过合法途径查询并追究违约责任"】', 'red')
    
    doc.add_paragraph()
    
    # ========== 第五条 ==========
    p = doc.add_paragraph()
    add_run(p, '第五条、争议的处理', 'black', True, 12)
    
    p = doc.add_paragraph()
    add_run(p, '1、双方因履行本协议发生的争议，应首先通过友好协商解决，协商不成时，任何一方可向乙方所在地人民法院提起诉讼解决。', 'black')
    add_run(p, '\n【修改建议：改为"甲方所在地"或"合同签订地"人民法院，或约定仲裁】', 'red')
    
    p = doc.add_paragraph()
    add_run(p, '2、本合同执行期为自合同签订之日起，三年内有效（付款时间不受此条款限制）。', 'black')
    add_run(p, '\n【修改建议：改为"至项目申报服务完成之日止"，或明确"有效期两年"】', 'red')
    
    p = doc.add_paragraph()
    add_run(p, '3、本合同一式两份，双方各持一份，本合同经双方盖章后立即生效。', 'black')
    
    doc.add_paragraph()
    
    # 签字页
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '（以下无正文）', 'black', False, 10)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 甲方签字
    p = doc.add_paragraph()
    add_run(p, '甲方：未来新视界科技（北京）有限公司', 'black', True)
    p = doc.add_paragraph()
    add_run(p, '（盖章）', 'black')
    p = doc.add_paragraph()
    add_run(p, '法定代表人：________________', 'black')
    p = doc.add_paragraph()
    add_run(p, '（授权代表人）', 'black')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # 乙方签字
    p = doc.add_paragraph()
    add_run(p, '乙方：北京一嘉二科技服务有限公司', 'black', True)
    p = doc.add_paragraph()
    add_run(p, '（盖章）', 'black')
    p = doc.add_paragraph()
    add_run(p, '法定代表人：________________', 'black')
    p = doc.add_paragraph()
    add_run(p, '（授权代表人）', 'black')
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '签约时间：2026 年 3 月 27 日', 'black')
    p = doc.add_paragraph()
    add_run(p, '签约地点：北京市', 'black')
    
    return doc

def create_summary():
    """创建审查意见汇总"""
    doc = Document()
    
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('合同审查意见汇总')
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    
    doc.add_paragraph()
    
    # 高风险条款
    p = doc.add_paragraph()
    add_run(p, '一、高风险条款（建议必须修改）', 'black', True, 12)
    
    p = doc.add_paragraph()
    add_run(p, '1. 第四条第 3 款：即使拒绝资助也要全额付款', 'black')
    add_run(p, '\n【风险：极不公平，建议删除或大幅修改】', 'red', True)
    
    p = doc.add_paragraph()
    add_run(p, '2. 第四条第 4 款：逾期利息万分之八/日', 'black')
    add_run(p, '\n【风险：年化约 29%，过高，建议改为万分之三】', 'red', True)
    
    p = doc.add_paragraph()
    add_run(p, '3. 第四条第 5 款：拒绝告知就要付全款', 'black')
    add_run(p, '\n【风险：过于苛刻，建议删除惩罚性条款】', 'red', True)
    
    p = doc.add_paragraph()
    add_run(p, '4. 第五条第 1 款：争议管辖在乙方所在地', 'black')
    add_run(p, '\n【风险：增加甲方维权成本，建议改为甲方所在地】', 'red', True)
    
    doc.add_paragraph()
    
    # 中风险条款
    p = doc.add_paragraph()
    add_run(p, '二、中风险条款（建议争取修改）', 'black', True, 12)
    
    p = doc.add_paragraph()
    add_run(p, '1. 第一条第 5 款："唯一合作人"限制', 'black')
    add_run(p, '\n【建议：删除排他性条款】', 'red')
    
    p = doc.add_paragraph()
    add_run(p, '2. 第二条第 4 款：不得拒付少付费用', 'black')
    add_run(p, '\n【建议：增加甲方索赔权利】', 'red')
    
    p = doc.add_paragraph()
    add_run(p, '3. 第三条第 1 款：10% 费率', 'black')
    add_run(p, '\n【建议：协商降至 8% 或设置阶梯费率】', 'red')
    
    p = doc.add_paragraph()
    add_run(p, '4. 第三条第 3 款：5 个工作日付款', 'black')
    add_run(p, '\n【建议：延长至 15 个工作日】', 'red')
    
    doc.add_paragraph()
    
    # 低风险提示
    p = doc.add_paragraph()
    add_run(p, '三、低风险提示（可协商）', 'black', True, 12)
    
    p = doc.add_paragraph()
    add_run(p, '1. 第二条第 3 款："及时答复"不明确', 'black')
    add_run(p, '\n【建议：明确为 2 个工作日】', 'red')
    
    p = doc.add_paragraph()
    add_run(p, '2. 第五条第 2 款：三年有效期', 'black')
    add_run(p, '\n【建议：改为服务完成之日或两年】', 'red')
    
    doc.add_paragraph()
    
    # 谈判策略
    p = doc.add_paragraph()
    add_run(p, '四、谈判策略建议', 'black', True, 12)
    
    p = doc.add_paragraph()
    add_run(p, '1. 优先争取修改高风险条款（4 条）', 'black')
    p = doc.add_paragraph()
    add_run(p, '2. 其次协商中风险条款（4 条）', 'black')
    p = doc.add_paragraph()
    add_run(p, '3. 低风险条款可作为让步筹码', 'black')
    
    doc.add_paragraph()
    
    # 总体评价
    p = doc.add_paragraph()
    add_run(p, '五、总体评价', 'black', True, 12)
    
    p = doc.add_paragraph()
    add_run(p, '该合同整体偏向乙方利益，存在多处不公平条款。建议重点修改第四条违约责任相关条款，争取更平等的权利义务关系。', 'black')
    
    return doc

if __name__ == '__main__':
    # 生成审查后的协议
    agreement_doc = create_agreement_review()
    agreement_path = '/home/admin/.openclaw/workspace-chief-agent/业务约定书（审查修改版）.docx'
    agreement_doc.save(agreement_path)
    print('✓ 审查修改版已生成：' + agreement_path)
    
    # 生成审查意见汇总
    summary_doc = create_summary()
    summary_path = '/home/admin/.openclaw/workspace-chief-agent/合同审查意见汇总.docx'
    summary_doc.save(summary_path)
    print('✓ 审查意见汇总已生成：' + summary_path)
    
    print('\n✅ 全部文档生成完成！')
