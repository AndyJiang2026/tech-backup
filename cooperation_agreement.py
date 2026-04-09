#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成合作协议 Word 文档 - 红色字体标注修改处
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def add_colored_run(paragraph, text, color='red', bold=False, font_size=12):
    """添加带颜色的文本片段"""
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    if color == 'red':
        run.font.color.rgb = RGBColor(255, 0, 0)
    elif color == 'black':
        run.font.color.rgb = RGBColor(0, 0, 0)
    return run

def create_cooperation_agreement():
    """创建合作协议文档"""
    doc = Document()
    
    # 设置文档样式
    doc.styles['Normal'].font.name = 'SimSun'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    doc.styles['Normal'].font.size = Pt(12)
    
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('合作协议')
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()  # 空行
    
    # 甲方乙方信息
    info_para = doc.add_paragraph()
    add_colored_run(info_para, '甲方（委托方）：', 'black', bold=True)
    add_colored_run(info_para, '[甲方公司名称]\n', 'black')
    add_colored_run(info_para, '乙方（服务方）：', 'black', bold=True)
    add_colored_run(info_para, '[乙方公司名称]\n', 'black')
    add_colored_run(info_para, '签订日期：', 'black', bold=True)
    add_colored_run(info_para, '2026 年  月  日', 'black')
    
    doc.add_paragraph()  # 空行
    
    # 前言
    preamble = doc.add_paragraph()
    preamble_para = preamble.add_run('鉴于甲乙双方本着平等互利、诚实信用的原则，经友好协商，就合作事宜达成如下协议：')
    preamble_para.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()  # 空行
    
    # ========== 第一条 合作内容 ==========
    h1 = doc.add_paragraph()
    h1_run = h1.add_run('第一条 合作内容')
    h1_run.font.size = Pt(14)
    h1_run.font.bold = True
    h1_run.font.color.rgb = RGBColor(0, 0, 0)
    
    content_para = doc.add_paragraph()
    add_colored_run(content_para, '1.1 乙方为甲方提供以下服务：', 'black')
    content_para.add_run('\n')
    add_colored_run(content_para, '（1）[具体服务内容 1]', 'black')
    content_para.add_run('\n')
    add_colored_run(content_para, '（2）[具体服务内容 2]', 'black')
    content_para.add_run('\n')
    add_colored_run(content_para, '（3）[具体服务内容 3]', 'black')
    
    doc.add_paragraph()  # 空行
    
    # ========== 第二条 双方权利义务 ==========
    h2 = doc.add_paragraph()
    h2_run = h2.add_run('第二条 双方权利义务')
    h2_run.font.size = Pt(14)
    h2_run.font.bold = True
    h2_run.font.color.rgb = RGBColor(0, 0, 0)
    
    # 甲方权利义务
    para_a = doc.add_paragraph()
    add_colored_run(para_a, '2.1 甲方权利义务：', 'black', bold=True)
    para_a.add_run('\n')
    add_colored_run(para_a, '（1）甲方有权要求乙方按照约定提供服务', 'black')
    para_a.add_run('\n')
    add_colored_run(para_a, '（2）甲方应按时支付服务费用', 'black')
    para_a.add_run('\n')
    add_colored_run(para_a, '（3）甲方应为乙方提供必要的工作条件和配合', 'black')
    
    # 乙方权利义务
    para_b = doc.add_paragraph()
    add_colored_run(para_b, '2.2 乙方权利义务：', 'black', bold=True)
    para_b.add_run('\n')
    add_colored_run(para_b, '（1）乙方有权按约定收取服务费用', 'black')
    para_b.add_run('\n')
    para_b2 = doc.add_paragraph()
    para_b2_run = para_b2.add_run('（2）乙方应保证服务质量，按约定完成工作')
    para_b2_run.font.color.rgb = RGBColor(0, 0, 0)
    para_b2.add_run('\n')
    add_colored_run(para_b2, '（3）乙方应在收到甲方反馈后 24 小时内响应（原为 48 小时）', 'red')
    
    doc.add_paragraph()  # 空行
    
    # ========== 第三条 费用及支付方式 ==========
    h3 = doc.add_paragraph()
    h3_run = h3.add_run('第三条 费用及支付方式')
    h3_run.font.size = Pt(14)
    h3_run.font.bold = True
    h3_run.font.color.rgb = RGBColor(0, 0, 0)
    
    fee_para1 = doc.add_paragraph()
    add_colored_run(fee_para1, '3.1 服务费用总额：人民币', 'black')
    add_colored_run(fee_para1, '[金额] 元（大写：', 'black')
    add_colored_run(fee_para1, '[大写金额]', 'black')
    add_colored_run(fee_para1, '）', 'black')
    
    fee_para2 = doc.add_paragraph()
    add_colored_run(fee_para2, '3.2 支付方式：', 'black', bold=True)
    fee_para2.add_run('\n')
    add_colored_run(fee_para2, '（1）合同签订后 5 个工作日内，甲方向乙方支付合同总额的 50% 作为预付款', 'black')
    fee_para2.add_run('\n')
    add_colored_run(fee_para2, '（2）乙方完成全部服务并经甲方验收合格后', 'black')
    add_colored_run(fee_para2, '15 个工作日', 'red', bold=True)
    add_colored_run(fee_para2, '内（', 'black')
    add_colored_run(fee_para2, '原为 30 天', 'red')
    add_colored_run(fee_para2, '），甲方向乙方支付剩余 50% 尾款', 'black')
    
    fee_para3 = doc.add_paragraph()
    add_colored_run(fee_para3, '3.3 乙方应在付款前向甲方开具等额有效的增值税专用发票', 'black')
    
    doc.add_paragraph()  # 空行
    
    # ========== 第四条 知识产权归属 ==========
    h4 = doc.add_paragraph()
    h4_run = h4.add_run('第四条 知识产权归属')
    h4_run.font.size = Pt(14)
    h4_run.font.bold = True
    h4_run.font.color.rgb = RGBColor(0, 0, 0)
    
    ip_para1 = doc.add_paragraph()
    add_colored_run(ip_para1, '4.1 乙方在履行本合同过程中产生的所有工作成果（包括但不限于文档、代码、设计稿、方案等）的知识产权', 'black')
    add_colored_run(ip_para1, '归甲方所有', 'red', bold=True)
    add_colored_run(ip_para1, '（', 'black')
    add_colored_run(ip_para1, '原为"双方共有"', 'red')
    add_colored_run(ip_para1, '）', 'black')
    
    ip_para2 = doc.add_paragraph()
    add_colored_run(ip_para2, '4.2 乙方保证其提供的服务不侵犯任何第三方的知识产权。如因乙方原因导致甲方被第三方主张侵权，乙方应承担全部赔偿责任', 'black')
    
    ip_para3 = doc.add_paragraph()
    add_colored_run(ip_para3, '4.3 甲方有权对乙方的工作成果进行任何形式的修改、使用、许可或转让，无需另行征得乙方同意', 'black')
    
    doc.add_paragraph()  # 空行
    
    # ========== 第五条 保密条款 ==========
    h5 = doc.add_paragraph()
    h5_run = h5.add_run('第五条 保密条款')
    h5_run.font.size = Pt(14)
    h5_run.font.bold = True
    h5_run.font.color.rgb = RGBColor(0, 0, 0)
    
    secret_para1 = doc.add_paragraph()
    add_colored_run(secret_para1, '5.1 双方应对在合作过程中获知的对方商业秘密、技术秘密、客户信息等保密信息严格保密', 'black')
    
    secret_para2 = doc.add_paragraph()
    add_colored_run(secret_para2, '5.2 保密期限：', 'black', bold=True)
    add_colored_run(secret_para2, '自合同签订之日起', 'black')
    add_colored_run(secret_para2, '5 年', 'red', bold=True)
    add_colored_run(secret_para2, '（', 'black')
    add_colored_run(secret_para2, '原为 2 年', 'red')
    add_colored_run(secret_para2, '）。保密期限不因合同终止而解除', 'black')
    
    secret_para3 = doc.add_paragraph()
    add_colored_run(secret_para3, '5.3 任何一方违反保密义务，应向守约方支付违约金人民币', 'black')
    add_colored_run(secret_para3, '[金额] 元', 'black')
    add_colored_run(secret_para3, '，并赔偿由此造成的全部损失', 'black')
    
    doc.add_paragraph()  # 空行
    
    # ========== 第六条 违约责任 ==========
    h6 = doc.add_paragraph()
    h6_run = h6.add_run('第六条 违约责任')
    h6_run.font.size = Pt(14)
    h6_run.font.bold = True
    h6_run.font.color.rgb = RGBColor(0, 0, 0)
    
    breach_para1 = doc.add_paragraph()
    add_colored_run(breach_para1, '6.1 任何一方未履行或未完全履行本合同项下的义务，即构成违约', 'black')
    
    breach_para2 = doc.add_paragraph()
    add_colored_run(breach_para2, '6.2 违约金标准：', 'black', bold=True)
    breach_para2.add_run('\n')
    add_colored_run(breach_para2, '（1）甲方逾期付款的，每逾期一日，应按逾期金额的', 'black')
    add_colored_run(breach_para2, '0.5%', 'black')
    add_colored_run(breach_para2, '向乙方支付违约金', 'black')
    breach_para2.add_run('\n')
    add_colored_run(breach_para2, '（2）乙方逾期交付工作成果的，每逾期一日，应按合同总额的', 'black')
    add_colored_run(breach_para2, '0.5%', 'black')
    add_colored_run(breach_para2, '向甲方支付违约金', 'black')
    breach_para2.add_run('\n')
    add_colored_run(breach_para2, '（3）任何一方严重违约导致合同目的无法实现的，守约方有权解除合同，违约方应向守约方支付合同总额', 'black')
    add_colored_run(breach_para2, '20%', 'red', bold=True)
    add_colored_run(breach_para2, '的违约金（', 'black')
    add_colored_run(breach_para2, '原为 10%', 'red')
    add_colored_run(breach_para2, '），并赔偿由此造成的全部损失', 'black')
    
    doc.add_paragraph()  # 空行
    
    # ========== 第七条 争议解决 ==========
    h7 = doc.add_paragraph()
    h7_run = h7.add_run('第七条 争议解决')
    h7_run.font.size = Pt(14)
    h7_run.font.bold = True
    h7_run.font.color.rgb = RGBColor(0, 0, 0)
    
    dispute_para = doc.add_paragraph()
    add_colored_run(dispute_para, '7.1 因本合同引起的或与本合同有关的任何争议，双方应首先通过友好协商解决', 'black')
    dispute_para.add_run('\n')
    add_colored_run(dispute_para, '7.2 协商不成的，任何一方均有权向', 'black')
    add_colored_run(dispute_para, '甲方所在地人民法院', 'red', bold=True)
    add_colored_run(dispute_para, '提起诉讼（', 'black')
    add_colored_run(dispute_para, '原为"乙方所在地"', 'red')
    add_colored_run(dispute_para, '）', 'black')
    dispute_para.add_run('\n')
    add_colored_run(dispute_para, '7.3 在争议解决期间，除争议事项外，双方应继续履行本合同的其他条款', 'black')
    
    doc.add_paragraph()  # 空行
    
    # ========== 第八条 合同期限及解除 ==========
    h8 = doc.add_paragraph()
    h8_run = h8.add_run('第八条 合同期限及解除')
    h8_run.font.size = Pt(14)
    h8_run.font.bold = True
    h8_run.font.color.rgb = RGBColor(0, 0, 0)
    
    term_para1 = doc.add_paragraph()
    add_colored_run(term_para1, '8.1 本合同有效期自', 'black')
    add_colored_run(term_para1, '[起始日期]', 'black')
    add_colored_run(term_para1, '至', 'black')
    add_colored_run(term_para1, '[终止日期]', 'black')
    add_colored_run(term_para1, '止', 'black')
    
    term_para2 = doc.add_paragraph()
    add_colored_run(term_para2, '8.2 合同解除条件：', 'black', bold=True)
    term_para2.add_run('\n')
    add_colored_run(term_para2, '（1）经双方协商一致，可以解除本合同', 'black')
    term_para2.add_run('\n')
    add_colored_run(term_para2, '（2）任何一方严重违约，守约方有权单方解除合同', 'black')
    term_para2.add_run('\n')
    add_colored_run(term_para2, '（3）因不可抗力导致合同无法继续履行的，双方均可解除合同', 'black')
    
    term_para3 = doc.add_paragraph()
    add_colored_run(term_para3, '8.3 合同解除后，乙方应在', 'black')
    add_colored_run(term_para3, '10 个工作日', 'black')
    add_colored_run(term_para3, '内向甲方移交全部工作成果和相关资料', 'black')
    
    doc.add_paragraph()  # 空行
    
    # ========== 第九条 其他条款 ==========
    h9 = doc.add_paragraph()
    h9_run = h9.add_run('第九条 其他条款')
    h9_run.font.size = Pt(14)
    h9_run.font.bold = True
    h9_run.font.color.rgb = RGBColor(0, 0, 0)
    
    other_para1 = doc.add_paragraph()
    add_colored_run(other_para1, '9.1 本合同一式', 'black')
    add_colored_run(other_para1, '贰', 'black')
    add_colored_run(other_para1, '份，甲乙双方各执', 'black')
    add_colored_run(other_para1, '壹', 'black')
    add_colored_run(other_para1, '份，具有同等法律效力', 'black')
    
    other_para2 = doc.add_paragraph()
    add_colored_run(other_para2, '9.2 本合同自双方签字盖章之日起生效', 'black')
    
    other_para3 = doc.add_paragraph()
    add_colored_run(other_para3, '9.3 本合同未尽事宜，双方可另行签订补充协议。补充协议与本合同具有同等法律效力', 'black')
    
    doc.add_paragraph()  # 空行
    doc.add_paragraph()  # 空行
    
    # 签字栏
    sign_title = doc.add_paragraph()
    sign_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sign_title_run = sign_title.add_run('（以下无正文，为签字盖章页）')
    sign_title_run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()  # 空行
    doc.add_paragraph()  # 空行
    
    # 甲方签字
    party_a = doc.add_paragraph()
    add_colored_run(party_a, '甲方（盖章）：____________________\n', 'black')
    add_colored_run(party_a, '法定代表人/授权代表（签字）：____________________\n', 'black')
    add_colored_run(party_a, '日期：______年____月____日\n', 'black')
    
    doc.add_paragraph()  # 空行
    doc.add_paragraph()  # 空行
    
    # 乙方签字
    party_b = doc.add_paragraph()
    add_colored_run(party_b, '乙方（盖章）：____________________\n', 'black')
    add_colored_run(party_b, '法定代表人/授权代表（签字）：____________________\n', 'black')
    add_colored_run(party_b, '日期：______年____月____日\n', 'black')
    
    return doc

def create_modification_summary():
    """创建修改说明清单（作为单独文档）"""
    doc = Document()
    
    # 标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('合作协议修改说明清单')
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()  # 空行
    
    # 说明文字
    intro = doc.add_paragraph()
    intro_run = intro.add_run('本文档列出了合作协议中所有用红色字体标注的修改内容，便于快速识别和审阅。')
    intro_run.font.color.rgb = RGBColor(0, 0, 0)
    
    doc.add_paragraph()  # 空行
    
    # 创建表格
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    # 表头
    header_cells = table.rows[0].cells
    headers = ['条款', '修改内容', '原内容', '修改后内容']
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    # 数据行
    modifications = [
        ('第二条 2.2(3)', '响应时间', '48 小时', '24 小时'),
        ('第三条 3.2(2)', '付款周期', '30 天', '15 个工作日'),
        ('第四条 4.1', '知识产权归属', '双方共有', '归甲方所有'),
        ('第五条 5.2', '保密期限', '2 年', '5 年'),
        ('第六条 6.2(3)', '违约金比例', '10%', '20%'),
        ('第七条 7.2', '争议解决管辖', '乙方所在地法院', '甲方所在地法院'),
    ]
    
    for mod in modifications:
        row = table.add_row()
        cells = row.cells
        for i, text in enumerate(mod):
            cells[i].text = text
    
    doc.add_paragraph()  # 空行
    
    # 备注
    note = doc.add_paragraph()
    note_run = note.add_run('备注：以上修改均为保护甲方权益的重要条款，建议重点审阅。')
    note_run.font.color.rgb = RGBColor(0, 0, 0)
    note_run.font.italic = True
    
    return doc

if __name__ == '__main__':
    # 生成合作协议
    agreement_doc = create_cooperation_agreement()
    agreement_path = '/home/admin/.openclaw/workspace-chief-agent/合作协议（红色标注版）.docx'
    agreement_doc.save(agreement_path)
    print(f'✓ 合作协议已生成：{agreement_path}')
    
    # 生成修改说明清单
    summary_doc = create_modification_summary()
    summary_path = '/home/admin/.openclaw/workspace-chief-agent/合作协议修改说明清单.docx'
    summary_doc.save(summary_path)
    print(f'✓ 修改说明清单已生成：{summary_path}')
    
    print('\n✅ 全部文档生成完成！')
