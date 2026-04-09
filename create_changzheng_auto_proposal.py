#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
《长征汽车》赞助专案 - 10 分钟打动决策者
核心：长征名 + 艰苦环境卓越性能 = 长征英雄
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_run(paragraph, text, color='black', bold=False, font_size=12):
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    if color == 'red':
        run.font.color.rgb = RGBColor(255, 0, 0)
    elif color == 'gold':
        run.font.color.rgb = RGBColor(218, 165, 32)
    elif color == 'darkred':
        run.font.color.rgb = RGBColor(139, 0, 0)
    elif color == 'blue':
        run.font.color.rgb = RGBColor(0, 0, 139)
    elif color == 'green':
        run.font.color.rgb = RGBColor(0, 100, 0)
    else:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return run

def add_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.color.rgb = RGBColor(139, 0, 0)
    return p

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    return p

def create_proposal():
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(11)
    
    # ========== 封面 ==========
    for _ in range(3):
        doc.add_paragraph()
    
    # 主标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('长征汽车·长征英雄\n战略合作方案')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.name = 'SimSun'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    title_run.font.color.rgb = RGBColor(139, 0, 0)
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('两个"长征"，一个使命\n卓越品质 = 长征英雄')
    subtitle_run.font.size = Pt(15)
    subtitle_run.font.name = 'SimSun'
    subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    subtitle_run.font.color.rgb = RGBColor(80, 80, 80)
    
    for _ in range(6):
        doc.add_paragraph()
    
    # 核心等式
    equation = doc.add_paragraph()
    equation.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq_run = equation.add_run('长征汽车 × 长征英雄 = 卓越品质')
    eq_run.font.size = Pt(18)
    eq_run.font.bold = True
    eq_run.font.name = 'SimSun'
    eq_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    eq_run.font.color.rgb = RGBColor(218, 165, 32)
    
    for _ in range(4):
        doc.add_paragraph()
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('2026 年 4 月')
    footer_run.font.size = Pt(11)
    footer_run.font.name = 'SimSun'
    footer_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    doc.add_page_break()
    
    # ========== 开篇：10 分钟决策 ==========
    add_title(doc, '开篇：10 分钟决策')
    
    p = doc.add_paragraph()
    add_run(p, '尊敬的长征汽车决策者：', 'black', False, 13)
    
    p = doc.add_paragraph()
    add_run(p, '如果您只有 10 分钟，请记住这三句话：', 'darkred', True)
    
    doc.add_paragraph()
    
    # 三句话
    quotes = [
        ('第一句', '贵公司名称叫"长征"，《长征·英雄》是长征题材电影，这是天然的品牌关联。', '不是选择题，是必答题。'),
        ('第二句', '长征汽车在艰苦环境下性能卓越，正如红军在艰苦卓绝中走向胜利。', '卓越品质 = 长征英雄。'),
        ('第三句', '赞助这不是广告，是向长征精神致敬，是对"长征"品牌的最好诠释。', '用钱买不到的品牌资产。')
    ]
    
    for q_title, q_content, q_note in quotes:
        p = doc.add_paragraph()
        add_run(p, q_title + '：', 'darkred', True)
        add_run(p, q_content, 'black', False, 12)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        add_run(p, q_note, 'blue', False, 11)
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '以下 8 分钟，为您详细阐述。', 'black', False, 12)
    
    # ========== 第一部分：两个"长征" ==========
    add_title(doc, '一、两个"长征"，一个使命')
    
    p = doc.add_paragraph()
    add_run(p, '世界上有两个"长征"品牌，注定要在一起：', 'black', False, 12)
    
    doc.add_paragraph()
    
    # 对照表
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['长征汽车', '《长征·英雄》']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.size = Pt(12)
    
    # 数据行
    rows_data = [
        ['企业名称：长征汽车集团', '电影名称：《长征·英雄》'],
        ['品牌精神：艰苦奋斗', '电影精神：长征精神'],
        ['产品特质：艰苦环境卓越性能', '电影特质：艰苦卓绝走向胜利'],
        ['目标客户：政企、军工、物流', '目标观众：政企、军人、爱国者'],
        ['品牌口号：行稳致远', '电影口号：每一代人有每一代人的长征路']
    ]
    
    for row_data in rows_data:
        row = table.add_row()
        cells = row.cells
        for i, text in enumerate(row_data):
            cells[i].text = text
            cells[i].paragraphs[0].runs[0].font.size = Pt(11)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '【核心结论】', 'darkred', True)
    p = doc.add_paragraph()
    add_run(p, '两个"长征"，精神同源、客户同频、使命同向。', 'black', False, 12)
    add_run(p, '这不是巧合，是历史的选择。', 'darkred', False, 12)
    
    # ========== 第二部分：卓越品质 ==========
    add_title(doc, '二、卓越品质 = 长征英雄')
    
    p = doc.add_paragraph()
    add_run(p, '长征汽车的卓越品质，在艰苦环境下得到验证，正如长征精神在艰苦卓绝中彰显：', 'black', False, 12)
    
    doc.add_paragraph()
    
    # 品质对照
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['长征精神', '长征汽车', '共同内核']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.size = Pt(11)
    
    quality_rows = [
        ['艰苦奋斗', '高原、沙漠、极寒等极端环境测试', '不畏艰难，挑战极限'],
        ['独立自主', '核心技术自主可控', '核心技术掌握在自己手中'],
        ['坚定信念', '60 年专注商用车', '专注一件事做到极致'],
        ['顾全大局', '服务国家物流、军工需求', '国家需要就是方向']
    ]
    
    for row_data in quality_rows:
        row = table.add_row()
        cells = row.cells
        for i, text in enumerate(row_data):
            cells[i].text = text
            cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '【产品性能实证】', 'darkred', True)
    
    products = [
        ('高原性能', '海拔 5000 米青藏高原，长征汽车动力不衰减', '正如红军翻越雪山'),
        ('极寒性能', '-40℃漠河极寒测试，一次启动成功', '正如红军过草地'),
        ('耐用性能', '100 万公里无大修，行业领先', '正如长征两万五'),
        ('军工品质', '军用越野车供应商，品质可靠', '正如人民军队')
    ]
    
    for p_name, p_content, p_note in products:
        p = doc.add_paragraph()
        add_run(p, p_name + '：', 'darkred', True)
        add_run(p, p_content, 'black', False, 12)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        add_run(p, p_note, 'blue', False, 11)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '【核心等式】', 'darkred', True)
    p = doc.add_paragraph()
    add_run(p, '长征汽车（艰苦环境卓越性能）× 《长征·英雄》（艰苦卓绝走向胜利）= 卓越品质', 'gold', True, 13)
    
    # ========== 第三部分：品牌资产 ==========
    add_title(doc, '三、品牌资产：用钱买不到')
    
    p = doc.add_paragraph()
    add_run(p, '赞助《长征·英雄》，长征汽车获得的品牌资产，是广告买不到的：', 'black', False, 12)
    
    doc.add_paragraph()
    
    assets = [
        ('品牌故事', '两个"长征"的天然关联，是独一无二的品牌故事', '广告费再高也买不到'),
        ('精神背书', '长征精神为长征汽车品牌背书', '提升品牌高度'),
        ('客户认同', '政企、军工、物流客户对红色文化认同度高', '促进销售转化'),
        ('媒体曝光', '"长征汽车支持长征题材电影"是天然新闻点', '媒体自发报道'),
        ('历史地位', '参与历史，成为长征精神传承的一部分', '载入史册')
    ]
    
    for a_name, a_content, a_note in assets:
        p = doc.add_paragraph()
        add_run(p, a_name + '：', 'darkred', True)
        add_run(p, a_content, 'black', False, 12)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        add_run(p, a_note, 'blue', False, 11)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '【核心观点】', 'darkred', True)
    p = doc.add_paragraph()
    add_run(p, '广告能买曝光，买不到精神背书。', 'black', False, 12)
    add_run(p, '赞助《长征·英雄》，是长征汽车品牌资产的战略性投资。', 'darkred', False, 12)
    
    # ========== 第四部分：赞助方案 ==========
    add_title(doc, '四、赞助方案：独家冠名')
    
    p = doc.add_paragraph()
    add_run(p, '鉴于长征汽车与《长征·英雄》的天然关联，我们提供', 'black', False, 12)
    add_run(p, '独家冠名赞助', 'red', True)
    add_run(p, '方案：', 'black', False, 12)
    
    doc.add_paragraph()
    
    # 赞助方案
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['权益类别', '具体内容']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.size = Pt(12)
    
    rights_data = [
        ['XR 体验区', '独家冠名"长征汽车·长征 XR 体验"'],
        ['首映礼', '全部官方用车使用长征汽车（10-20 辆）'],
        ['品牌曝光', '主背景板 Logo C 位 + 企业领导 10 分钟致辞'],
        ['媒体传播', '全部通稿署名"独家冠名：长征汽车"'],
        ['分众广告', '2 周全国电梯媒体投放（含长征汽车产品）'],
        ['内容绑定', '影片片尾"特别鸣谢：长征汽车"（单独一行）'],
        ['政企对接', '主办专场政企对接会，邀请相关部委领导'],
        ['长期授权', '3 年使用"《长征·英雄》独家冠名合作伙伴"称号']
    ]
    
    for row_data in rights_data:
        row = table.add_row()
        cells = row.cells
        cells[0].text = row_data[0]
        cells[1].text = row_data[1]
        cells[0].paragraphs[0].runs[0].font.size = Pt(11)
        cells[1].paragraphs[0].runs[0].font.size = Pt(11)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '【赞助金额】', 'darkred', True)
    p = doc.add_paragraph()
    add_run(p, '800 万元（独家冠名，仅 1 家）', 'red', True, 13)
    add_run(p, ' 可分期支付：签约 50% + 活动前 30% + 活动后 20%', 'black', False, 12)
    
    # ========== 第五部分：投资回报 ==========
    add_title(doc, '五、投资回报分析')
    
    p = doc.add_paragraph()
    add_run(p, '800 万投资，获得什么？算两笔账：', 'black', False, 12)
    
    doc.add_paragraph()
    
    # 经济账
    p = doc.add_paragraph()
    add_run(p, '5.1 经济账（可量化）', 'darkred', True)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['回报类型', '具体内容', '估算价值']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    
    roi_data = [
        ['XR 体验区', '独家冠名 +3 个月展示', '200 万'],
        ['分众广告', '2 周全国电梯媒体', '300 万'],
        ['央视/主流媒体', '新闻报道 + 专题', '200 万'],
        ['首映礼用车', '10-20 辆长征汽车展示', '100 万'],
        ['内容绑定', '片尾鸣谢 + 纪录片', '100 万'],
        ['长期授权', '3 年使用权', '200 万'],
        ['合计', '-', '1100 万+']
    ]
    
    for row_data in roi_data:
        row = table.add_row()
        cells = row.cells
        for i, text in enumerate(row_data):
            cells[i].text = text
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '直接经济回报：1100 万+，ROI 约 138%', 'black', False, 12)
    
    # 品牌账
    p = doc.add_paragraph()
    add_run(p, '5.2 品牌账（不可量化但更重要）', 'darkred', True)
    
    items = [
        '两个"长征"的品牌故事（独一无二）',
        '长征精神背书（提升品牌高度）',
        '政企客户认同（促进销售转化）',
        '媒体自发报道（"长征汽车支持长征题材"）',
        '历史地位（载入红色文化史册）'
    ]
    
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        add_run(p, '✓ ', 'red')
        add_run(p, item)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '【核心结论】', 'darkred', True)
    p = doc.add_paragraph()
    add_run(p, '经济账算得清（ROI 138%），品牌账算不清（无价）。', 'black', False, 12)
    add_run(p, '这是战略性投资，不是广告费。', 'darkred', False, 12)
    
    # ========== 第六部分：为什么是现在 ==========
    add_title(doc, '六、为什么是现在？')
    
    p = doc.add_paragraph()
    add_run(p, '投资讲究时机。长征汽车赞助《长征·英雄》的时机，可以用四个词概括：', 'black', False, 12)
    
    doc.add_paragraph()
    
    timing = [
        ('天时', '2026 年建党 105 周年 + 长征胜利 90 周年', '政治窗口期，央企/国企有党建预算'),
        ('地利', '首个重大革命题材 XR 电影龙标', '"第一"本身就是价值，媒体自发报道'),
        ('人和', '长征汽车需要品牌向上突破', '从"工具"到"精神"，提升品牌高度'),
        ('竞争', '独家冠名仅 1 家', '错过这次，没有下次')
    ]
    
    for t_name, t_content, t_note in timing:
        p = doc.add_paragraph()
        add_run(p, t_name + '：', 'darkred', True)
        add_run(p, t_content, 'black', False, 12)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        add_run(p, t_note, 'blue', False, 11)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '【核心判断】', 'darkred', True)
    p = doc.add_paragraph()
    add_run(p, '3 年后回头看，这次赞助的价值，远超 800 万投入。', 'black', False, 12)
    add_run(p, '因为"首个"只有一次，"第一"无法复制。', 'darkred', False, 12)
    
    # ========== 第七部分：行动建议 ==========
    add_title(doc, '七、行动建议')
    
    p = doc.add_paragraph()
    add_run(p, '如果您认同这个项目的价值，我们建议：', 'black', False, 12)
    
    doc.add_paragraph()
    
    actions = [
        ('第 1 步：深度沟通', '1-3 天内，我们上门拜访，详细了解需求'),
        ('第 2 步：方案确认', '3-5 天内，确认独家冠名权益'),
        ('第 3 步：内部决策', '5-7 天内，协助准备内部决策材料'),
        ('第 4 步：签约落地', '签署战略合作协议，支付 50% 首付款'),
        ('第 5 步：车辆准备', '准备 10-20 辆长征汽车用于首映礼'),
        ('第 6 步：活动执行', '首映礼当天，企业领导致辞 + 领导接见'),
        ('第 7 步：尾款结算', '活动后 7 天内，支付剩余 20% 尾款'),
        ('第 8 步：长期运营', '3 年期授权，持续品牌使用')
    ]
    
    for a_name, a_content in actions:
        p = doc.add_paragraph()
        add_run(p, a_name + '：', 'darkred', True)
        add_run(p, a_content, 'black', False, 12)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '【时间窗口】', 'darkred', True)
    p = doc.add_paragraph()
    add_run(p, '首映礼时间：2026 年 5 月（待定）', 'black', False, 12)
    add_run(p, '独家冠名签约截止：2026 年 4 月 15 日', 'red', True)
    add_run(p, '（预留内部决策时间）', 'blue', False, 11)
    
    # ========== 结语 ==========
    add_title(doc, '结语：两个"长征"，一个使命')
    
    p = doc.add_paragraph()
    add_run(p, '尊敬的决策者：', 'black', False, 12)
    
    p = doc.add_paragraph()
    add_run(p, '长征汽车用 60 年，在艰苦环境下打造卓越品质。', 'black', False, 12)
    add_run(p, '《长征·英雄》用 XR 技术，在银幕上重现长征精神。', 'black', False, 12)
    
    p = doc.add_paragraph()
    add_run(p, '两个"长征"，精神同源、使命同向。', 'darkred', True, 13)
    
    p = doc.add_paragraph()
    add_run(p, '3 年后，当您的竞争对手问起：', 'black', False, 12)
    add_run(p, '"当年那个长征题材电影，长征汽车参与了没有？"', 'black', False, 12)
    
    p = doc.add_paragraph()
    add_run(p, '您希望如何回答？', 'darkred', True, 13)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '两个"长征"，一个使命。', 'darkred', True, 15)
    add_run(p, '\n我们期待与您一起，创造历史。', 'darkred', True, 15)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '2026 年 4 月', 'black', False, 11)
    
    doc.add_page_break()
    
    # ========== 附录：10 分钟快速决策表 ==========
    add_title(doc, '附录：10 分钟快速决策表')
    
    p = doc.add_paragraph()
    add_run(p, '如果您只有 10 分钟，请看这张表：', 'black', False, 12)
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['决策要素', '内容', '您的关注点', '我们的回应']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    
    decision_rows = [
        ['这是什么？', '首个重大革命题材 XR 电影', '稀缺性', '"第一"无法复制'],
        ['为什么是我们？', '两个"长征"天然关联', '匹配度', '独一无二的品牌故事'],
        ['为什么是现在？', '建党 105 周年 + 长征 90 周年', '时机', '政治窗口期'],
        ['要投多少？', '800 万（独家冠名）', '预算', '可分期支付'],
        ['回报是什么？', '1100 万+ROI138%+ 品牌资产', 'ROI', '经济 + 品牌双账'],
        ['风险是什么？', '政治正确，无舆情风险', '安全', '红色 IP 最安全'],
        ['决策流程？', '1-3 天 +3-5 天 +5-7 天', '时间', '预留内部决策时间']
    ]
    
    for row_data in decision_rows:
        row = table.add_row()
        cells = row.cells
        for i, text in enumerate(row_data):
            cells[i].text = text
            cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '【10 分钟决策建议】', 'darkred', True)
    p = doc.add_paragraph()
    add_run(p, '如果以上 7 个要素都符合您的预期，建议：', 'black', False, 12)
    add_run(p, '立即安排深度沟通，锁定独家冠名名额（仅 1 家）', 'red', True)
    
    return doc

if __name__ == '__main__':
    doc = create_proposal()
    output_path = '/home/admin/.openclaw/workspace-chief-agent/长征汽车赞助专案.docx'
    doc.save(output_path)
    print('✅ 长征汽车赞助专案已生成：' + output_path)
