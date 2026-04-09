#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
《长征·英雄》决策者版本 - 10 分钟打动赞助商
参考：央视招商会开场演讲模式
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_run(paragraph, text, color='black', bold=False, font_size=12):
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    if color == 'red':
        run.font.color.rgb = RGBColor(255, 0, 0)
    elif color == 'gold':
        run.font.color.rgb = RGBColor(218, 165, 32)
    elif color == 'darkred':
        run.font.color.rgb = RGBColor(139, 0, 0)
    elif color == 'blue':
        run.font.color.rgb = RGBColor(0, 0, 139)
    else:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return run

def add_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.color.rgb = RGBColor(139, 0, 0)
    return p

def add_quote(doc, text, source=''):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(1)
    p.paragraph_format.right_indent = Inches(1)
    run = p.add_run('"' + text + '"')
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.color.rgb = RGBColor(100, 100, 100)
    if source:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(1)
        p.paragraph_format.right_indent = Inches(1)
        run = p.add_run('—— ' + source)
        run.font.size = Pt(10)
        run.font.name = 'SimSun'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
        run.font.color.rgb = RGBColor(150, 150, 150)

def create_decision_maker_version():
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(11)
    
    # ========== 封面 ==========
    for _ in range(2):
        doc.add_paragraph()
    
    # 主标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('致企业决策者的一封信')
    title_run.font.size = Pt(20)
    title_run.font.bold = True
    title_run.font.name = 'SimSun'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    title_run.font.color.rgb = RGBColor(139, 0, 0)
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('—— 关于长征精神与您的企业战略')
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.name = 'SimSun'
    subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    subtitle_run.font.color.rgb = RGBColor(80, 80, 80)
    
    for _ in range(4):
        doc.add_paragraph()
    
    # 引用
    quote = doc.add_paragraph()
    quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote_run = quote.add_run('"每一代人有每一代人的长征路"\n"每一代企业有每一代企业的使命担当"')
    quote_run.font.size = Pt(13)
    quote_run.font.italic = True
    quote_run.font.name = 'SimSun'
    quote_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    quote_run.font.color.rgb = RGBColor(100, 100, 100)
    
    for _ in range(6):
        doc.add_paragraph()
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('2026 年 4 月')
    footer_run.font.size = Pt(11)
    footer_run.font.name = 'SimSun'
    footer_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    doc.add_page_break()
    
    # ========== 开篇：三个问题 ==========
    add_title(doc, '开篇：三个问题')
    
    p = doc.add_paragraph()
    add_run(p, '尊敬的决策者：', 'black', False, 12)
    
    p = doc.add_paragraph()
    add_run(p, '在您继续阅读这份方案之前，我想请您先思考三个问题：', 'black', False, 12)
    
    doc.add_paragraph()
    
    questions = [
        ('问题一', '2026 年是建党 105 周年，也是"十四五"收官之年。您的企业如何向党和国家的生日献礼？', '这是政治任务，也是展示企业担当的历史机遇。'),
        ('问题二', '2026 年是"十五五"规划谋篇布局之年。您的企业如何在下一个五年中，与国家发展战略同频共振？', '这关乎企业未来五年的政策红利和发展空间。'),
        ('问题三', '在当前的经济环境下，您的企业如何找到既能完成政治任务、又能获得商业回报、还能提升品牌价值的"三赢"机会？', '这不是选择题，而是必答题。')
    ]
    
    for q_title, q_content, q_note in questions:
        p = doc.add_paragraph()
        add_run(p, q_title + '：', 'darkred', True)
        add_run(p, q_content, 'black', False, 12)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        add_run(p, q_note, 'blue', False, 11)
        doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '这三个问题，正是我们今天向您推荐《长征·英雄》项目的初衷。', 'black', False, 12)
    
    # ========== 第一部分：长征精神 ==========
    add_title(doc, '一、长征精神与企业家精神')
    
    p = doc.add_paragraph()
    add_run(p, '长征，不仅仅是一段历史，更是一种精神。这种精神，与企业家精神高度契合：', 'black', False, 12)
    
    doc.add_paragraph()
    
    # 精神对照表
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['长征精神', '企业家精神', '共同内核']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.size = Pt(11)
    
    # 数据行
    spirit_rows = [
        ['坚定信念', '愿景驱动', '心中有信仰，脚下有力量'],
        ['不怕牺牲', '敢于投入', '看准方向，坚定不移'],
        ['独立自主', '自主创新', '核心技术必须掌握在自己手中'],
        ['实事求是', '务实经营', '不盲目扩张，稳健发展'],
        ['顾全大局', '社会责任', '企业不仅要赚钱，更要担当']
    ]
    
    for row_data in spirit_rows:
        row = table.add_row()
        cells = row.cells
        for i, text in enumerate(row_data):
            cells[i].text = text
            cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '【核心观点】', 'darkred', True)
    p = doc.add_paragraph()
    add_run(p, '长征精神不是历史，而是现实。它不是口号，而是行动。它不是负担，而是力量。', 'black', False, 12)
    add_run(p, '您的企业一路走来，何尝不是一次"长征"？', 'black', False, 12)
    
    # ========== 第二部分：战略连接 ==========
    add_title(doc, '二、长征精神与您的企业战略')
    
    p = doc.add_paragraph()
    add_run(p, '我们深入研究了不同类型企业的战略诉求，发现长征精神可以与每个企业的战略找到连接点：', 'black', False, 12)
    
    doc.add_paragraph()
    
    # 央企/国企
    p = doc.add_paragraph()
    add_run(p, '2.1 央企/国企：政治任务 + 十五五规划', 'darkred', True, 13)
    
    p = doc.add_paragraph()
    add_run(p, '【您的战略诉求】', 'black', True)
    
    items = [
        '完成国资委党建考核任务',
        '在"十五五"规划中争取更多政策支持',
        '展示央企政治担当，提升领导政绩',
        '维护政府关系，获取政策信息'
    ]
    
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        add_run(p, '• ', 'red')
        add_run(p, item)
    
    p = doc.add_paragraph()
    add_run(p, '【长征·英雄的连接点】', 'black', True)
    
    items = [
        '参与首个重大革命题材龙标影片，是',
        '政治任务',
        '（国资委考核加分项）',
        '首映礼政企对接会，可直接对接相关部委领导',
        '党建活动可组织员工观看，作为主题教育素材',
        '企业参与红色文化工程，可纳入"十五五"规划社会责任章节'
    ]
    
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        if '政治任务' in item or '十五五' in item:
            add_run(p, item, 'red', True)
        else:
            add_run(p, item)
    
    doc.add_paragraph()
    
    # 科技公司
    p = doc.add_paragraph()
    add_run(p, '2.2 科技公司：G 端业务 + 技术背书', 'darkred', True, 13)
    
    p = doc.add_paragraph()
    add_run(p, '【您的战略诉求】', 'black', True)
    
    items = [
        '拓展 G 端（政府）业务',
        '在政府项目中获得技术背书',
        '提前布局"十五五"数字经济、文化数字化政策',
        '打造标杆案例，用于投标加分'
    ]
    
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        add_run(p, '• ', 'red')
        add_run(p, item)
    
    p = doc.add_paragraph()
    add_run(p, '【长征·英雄的连接点】', 'black', True)
    
    items = [
        'XR 体验区独家冠名，展示企业技术实力',
        '政府领导参观时，企业技术人员可现场讲解',
        '提供"红色文化工程技术支持单位"证书，可用于政府项目投标',
        '文化数字化是"十五五"重点方向，提前布局占位'
    ]
    
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        if 'XR' in item or '十五五' in item or '投标' in item:
            add_run(p, item, 'red', True)
        else:
            add_run(p, item)
    
    doc.add_paragraph()
    
    # 消费品牌
    p = doc.add_paragraph()
    add_run(p, '2.3 消费品牌：品牌安全 + 渠道下沉', 'darkred', True, 13)
    
    p = doc.add_paragraph()
    add_run(p, '【您的战略诉求】', 'black', True)
    
    items = [
        '品牌安全，无舆情风险',
        '三四线城市渠道下沉',
        '政企团购资源开发',
        '品牌美誉度提升'
    ]
    
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        add_run(p, '• ', 'red')
        add_run(p, item)
    
    p = doc.add_paragraph()
    add_run(p, '【长征·英雄的连接点】', 'black', True)
    
    items = [
        '红色 IP，政治正确，品牌安全系数最高',
        '分众媒体覆盖全国一二三线城市，助力渠道下沉',
        '三四线城市消费者对红色 IP 认可度更高',
        '可借势开展"红色文化 + 产品"促销活动',
        '对接参会央企/国企，开发政企团购资源'
    ]
    
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        if '红色' in item or '渠道' in item or '政企' in item:
            add_run(p, item, 'red', True)
        else:
            add_run(p, item)
    
    # ========== 第三部分：历史机遇 ==========
    add_title(doc, '三、为什么是现在？')
    
    p = doc.add_paragraph()
    add_run(p, '投资讲究时机。《长征·英雄》项目的时机，可以用四个词概括：', 'black', False, 12)
    
    doc.add_paragraph()
    
    timing = [
        ('天时', '2026 年建党 105 周年，政治窗口期', '央企/国企有党建预算，政府有宣传需求'),
        ('地利', '首个重大革命题材龙标，稀缺性', '"第一"本身就是价值，媒体会自发报道'),
        ('人和', 'XR 技术成熟，可商业化运营', '2016 年 VR 刚起步，2026 年 XR 可落地'),
        ('政策', '"十五五"规划文化数字化', '提前布局，占位政策红利')
    ]
    
    for t_name, t_content, t_note in timing:
        p = doc.add_paragraph()
        add_run(p, t_name + '：', 'darkred', True)
        add_run(p, t_content, 'black', False, 12)
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        add_run(p, t_note, 'blue', False, 11)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '【核心判断】', 'darkred', True)
    p = doc.add_paragraph()
    add_run(p, '【核心判断】', 'darkred', True)
    p = doc.add_paragraph()
    add_run(p, '这个项目的价值，3 年后回头看，会远超今天的投入。', 'black', False, 12)
    add_run(p, '因为"首个"只有一次，"第一"无法复制。', 'black', False, 12)
    
    # ========== 第四部分：我们提供什么 ==========
    add_title(doc, '四、我们提供什么？')
    
    p = doc.add_paragraph()
    add_run(p, '我们提供的不是"赞助权益"，而是：', 'black', False, 12)
    
    doc.add_paragraph()
    
    offerings = [
        ('一个政治资产', '首个重大革命题材龙标影片的合作伙伴身份，可用于企业长期背书'),
        ('一个政企通道', '首映礼政企对接会，可与相关部委、北京市领导面对面交流'),
        ('一个技术场景', 'XR 体验区，展示企业技术实力，打造 G 端业务标杆案例'),
        ('一个品牌背书', '红色 IP，政治正确，品牌安全，可用于企业宣传长期使用'),
        ('一个历史机遇', '参与历史，创造历史，成为"首个"的见证者和参与者')
    ]
    
    for o_name, o_content in offerings:
        p = doc.add_paragraph()
        add_run(p, '✓ ', 'red')
        add_run(p, o_name + '：', 'darkred', True)
        add_run(p, o_content, 'black', False, 12)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '【核心差异】', 'darkred', True)
    p = doc.add_paragraph()
    add_run(p, '传统赞助：买广告位，换曝光', 'black', False, 12)
    add_run(p, '我们的方案：共建红色文化工程，获得政治资产 + 政企通道 + 品牌背书', 'red', False, 12)
    
    # ========== 第五部分：投资回报 ==========
    add_title(doc, '五、投资回报分析')
    
    p = doc.add_paragraph()
    add_run(p, '我们不谈虚的，只算两笔账：', 'black', False, 12)
    
    doc.add_paragraph()
    
    # 经济账
    p = doc.add_paragraph()
    add_run(p, '5.1 经济账（可量化）', 'darkred', True)
    
    p = doc.add_paragraph()
    add_run(p, '以战略合作伙伴（500 万元）为例：', 'black', False, 12)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['回报类型', '具体内容', '估算价值']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    
    roi_data = [
        ['分众广告', '2 周全国电梯媒体', '300 万'],
        ['央视/主流媒体', '新闻报道 + 专题', '200 万'],
        ['XR 体验区', '独家冠名 + 展示', '150 万'],
        ['内容绑定', '片尾鸣谢 + 纪录片', '100 万'],
        ['长期授权', '3 年使用权', '150 万'],
        ['合计', '-', '900 万+']
    ]
    
    for row_data in roi_data:
        row = table.add_row()
        cells = row.cells
        for i, text in enumerate(row_data):
            cells[i].text = text
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '直接经济回报：900 万+，ROI 约 180%', 'black', False, 12)
    
    # 政治账
    p = doc.add_paragraph()
    add_run(p, '5.2 政治账（不可量化但更重要）', 'darkred', True)
    
    items = [
        '国资委党建考核加分（央企/国企）',
        '政府关系建立与维护（所有企业）',
        'G 端业务背书（科技公司）',
        '品牌安全背书（消费品牌）',
        '企业历史中的"红色篇章"（长期价值）'
    ]
    
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        add_run(p, '✓ ', 'red')
        add_run(p, item)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '【核心结论】', 'darkred', True)
    p = doc.add_paragraph()
    add_run(p, '经济账算得清，政治账算不清。但政治账的价值，往往远超经济账。', 'black', False, 12)
    
    # ========== 第六部分：行动建议 ==========
    add_title(doc, '六、行动建议')
    
    p = doc.add_paragraph()
    add_run(p, '如果您认同这个项目的价值，我们建议：', 'black', False, 12)
    
    doc.add_paragraph()
    
    actions = [
        ('第 1 步：深度沟通', '1-3 天内，我们上门拜访，详细了解您的战略诉求'),
        ('第 2 步：定制方案', '根据您的需求，定制专属合作方案'),
        ('第 3 步：内部决策', '协助您准备内部决策材料（党建考核、预算申请等）'),
        ('第 4 步：签约落地', '签署战略合作协议，启动权益执行'),
        ('第 5 步：长期运营', '不仅是首映礼，更是 3 年期的战略合作')
    ]
    
    for a_name, a_content in actions:
        p = doc.add_paragraph()
        add_run(p, a_name + '：', 'darkred', True)
        add_run(p, a_content, 'black', False, 12)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '【时间窗口】', 'darkred', True)
    p = doc.add_paragraph()
    add_run(p, '首映礼时间：2026 年 5 月（待定）', 'black', False, 12)
    add_run(p, '战略合作伙伴签约截止：2026 年 4 月 15 日', 'red', True)
    add_run(p, '（预留内部决策时间）', 'blue', False, 11)
    
    # ========== 结语 ==========
    add_title(doc, '结语：一起创造历史')
    
    p = doc.add_paragraph()
    add_run(p, '尊敬的决策者：', 'black', False, 12)
    
    p = doc.add_paragraph()
    add_run(p, '长征，是一段历史。但长征精神，是现实的力量。', 'black', False, 12)
    
    p = doc.add_paragraph()
    add_run(p, '《长征·英雄》不仅仅是一部电影，它是一个红色文化工程，是一个历史机遇，是一个让您的企业载入史册的机会。', 'black', False, 12)
    
    p = doc.add_paragraph()
    add_run(p, '3 年后，当您的竞争对手问起："当年那个首个重大革命题材龙标影片，你参与了没有？"', 'black', False, 12)
    
    p = doc.add_paragraph()
    add_run(p, '您希望如何回答？', 'darkred', True, 13)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '我们期待与您一起，创造历史。', 'darkred', True, 14)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '2026 年 4 月', 'black', False, 11)
    
    doc.add_page_break()
    
    # ========== 附录：快速决策表 ==========
    add_title(doc, '附录：10 分钟快速决策表')
    
    p = doc.add_paragraph()
    add_run(p, '如果您只有 10 分钟，请看这张表：', 'black', False, 12)
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['决策要素', '内容', '您的关注点', '我们的回应']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    
    decision_rows = [
        ['这是什么？', '首个重大革命题材龙标影片', '稀缺性', '"第一"无法复制'],
        ['为什么现在？', '建党 105 周年 + 十五五规划', '时机', '政治窗口期'],
        ['对我有什么用？', '政治资产 + 政企通道 + 品牌背书', '价值', '满足隐性需求'],
        ['要投多少？', '500-800 万（战略合作伙伴）', '预算', '可分期支付'],
        ['回报是什么？', '经济回报 900 万+ROI180%+ 隐性价值', 'ROI', '经济 + 政治双账'],
        ['风险是什么？', '政治正确，无舆情风险', '安全', '红色 IP 最安全'],
        ['决策流程？', '1-3 天沟通 +3-5 天定制 +5-7 天签约', '时间', '预留内部决策时间']
    ]
    
    for row_data in decision_rows:
        row = table.add_row()
        cells = row.cells
        for i, text in enumerate(row_data):
            cells[i].text = text
            cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '【10 分钟决策建议】', 'darkred', True)
    p = doc.add_paragraph()
    add_run(p, '如果以上 7 个要素都符合您的预期，建议：', 'black', False, 12)
    add_run(p, '立即安排深度沟通，锁定战略合作伙伴名额（仅 1-2 家）', 'red', True)
    
    return doc

if __name__ == '__main__':
    doc = create_decision_maker_version()
    output_path = '/home/admin/.openclaw/workspace-chief-agent/长征英雄赞助招商方案（决策者版本）.docx'
    doc.save(output_path)
    print('✅ 决策者版本方案已生成：' + output_path)
