#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
《长征·英雄》首映礼赞助招商方案 - 央视模式升级版
参考：央视《辉煌中国》《大国崛起》《开学第一课》招商模式
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_run(paragraph, text, color='black', bold=False, font_size=11):
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    if color == 'red':
        run.font.color.rgb = RGBColor(255, 0, 0)
    elif color == 'blue':
        run.font.color.rgb = RGBColor(0, 0, 128)
    elif color == 'darkred':
        run.font.color.rgb = RGBColor(128, 0, 0)
    else:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return run

def add_title(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.color.rgb = RGBColor(192, 0, 0)
    return p

def add_subtitle(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    return p

def create_proposal():
    doc = Document()
    
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(11)
    
    # ========== 封面 ==========
    for _ in range(3):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('《长征·英雄》\n红色文化工程战略合作方案')
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.name = 'SimSun'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    title_run.font.color.rgb = RGBColor(192, 0, 0)
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('首个重大革命题材龙标影片\nXR 沉浸式体验 + 红色 IP 运营')
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.name = 'SimSun'
    subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    subtitle_run.font.color.rgb = RGBColor(80, 80, 80)
    
    for _ in range(6):
        doc.add_paragraph()
    
    # 参考央视
    ref = doc.add_paragraph()
    ref.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ref_run = ref.add_run('参考央视《辉煌中国》《大国崛起》《开学第一课》招商模式')
    ref_run.font.size = Pt(10)
    ref_run.font.name = 'SimSun'
    ref_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    ref_run.font.color.rgb = RGBColor(100, 100, 100)
    ref_run.font.italic = True
    
    for _ in range(4):
        doc.add_paragraph()
    
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('2026 年 4 月')
    footer_run.font.size = Pt(11)
    footer_run.font.name = 'SimSun'
    footer_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    doc.add_page_break()
    
    # ========== 核心洞察 ==========
    add_title(doc, '一、赞助商核心需求洞察')
    
    p = doc.add_paragraph()
    add_run(p, '我们深入研究了央视主旋律项目招商模式，发现赞助商真正想要的不是"广告位"，而是：', 'black', False, 11)
    
    doc.add_paragraph()
    
    # 隐性需求 vs 显性需求
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['显性需求（表面）', '隐性需求（核心）']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].runs[0].font.size = Pt(11)
    
    # 央企/国企
    row = table.add_row()
    cells = row.cells
    cells[0].text = '• 品牌曝光\n• 媒体传播\n• 社会责任形象'
    cells[1].text = '• 完成政治任务（国资委考核）\n• 领导政绩展示\n• 政府关系维护\n• 获取政策信息'
    cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    
    # 科技公司
    row = table.add_row()
    cells = row.cells
    cells[0].text = '• 技术展示\n• 品牌知名度\n• 获客'
    cells[1].text = '• G 端业务拓展（政府项目）\n• 政府背书（投标加分）\n• 领导参观（展示实力）\n• 政策解读（提前布局）'
    cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    
    # 消费品牌
    row = table.add_row()
    cells = row.cells
    cells[0].text = '• 品牌美誉度\n• 销售转化'
    cells[1].text = '• 品牌安全（无舆情风险）\n• 渠道下沉（三四线城市）\n• 政企团购资源\n• 员工福利采购'
    cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '【核心结论】', 'darkred', True)
    p = doc.add_paragraph()
    add_run(p, '央企要政治、科技要 G 端、消费要安全。', 'black', False, 11)
    add_run(p, '我们的方案必须满足这三类企业的隐性需求，而不仅仅是给广告位。', 'black', False, 11)
    
    # ========== 央视模式借鉴 ==========
    add_title(doc, '二、央视主旋律项目招商模式借鉴')
    
    p = doc.add_paragraph()
    add_run(p, '我们研究了央视近年主旋律项目的招商模式，核心经验如下：', 'black', False, 11)
    
    doc.add_paragraph()
    
    # 央视案例
    cases = [
        {
            'name': '《辉煌中国》',
            'type': '纪录片',
            'sponsor': '中车、国家电网、中国移动',
            'mode': '特约播出 + 深度绑定',
            'key': '企业故事融入内容，不是简单贴片'
        },
        {
            'name': '《大国崛起》',
            'type': '纪录片',
            'sponsor': '中石化、宝钢',
            'mode': '联合制作 + 资源置换',
            'key': '企业领导出镜，讲述行业贡献'
        },
        {
            'name': '《开学第一课》',
            'type': '公益节目',
            'sponsor': '伊利、安踏、华为',
            'mode': '独家冠名 + 场景植入',
            'key': '产品使用场景自然植入'
        },
        {
            'name': '《长征大会师》',
            'type': '电视剧',
            'sponsor': '茅台、五粮液',
            'mode': '片头冠名 + 线下活动',
            'key': '白酒与红色文化高度契合'
        }
    ]
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    headers = ['项目名称', '类型', '赞助商', '合作模式', '核心亮点']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    
    for case in cases:
        row = table.add_row()
        cells = row.cells
        for i, key in enumerate(['name', 'type', 'sponsor', 'mode', 'key']):
            cells[i].text = case[key]
            cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '【可借鉴经验】', 'darkred', True)
    
    experiences = [
        '内容深度绑定：企业故事融入影片/活动，不是简单贴片广告',
        '领导出镜：企业领导可出席活动、致辞、接受采访',
        '场景植入：产品/技术在使用场景中自然展示',
        '长期授权：赞助称号可用于企业长期宣传',
        '政企对接：借活动建立政府关系，获取项目信息'
    ]
    
    for exp in experiences:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        add_run(p, '✓ ', 'red')
        add_run(p, exp)
    
    # ========== 我们的方案 ==========
    add_title(doc, '三、《长征·英雄》战略合作方案')
    
    p = doc.add_paragraph()
    add_run(p, '基于央视模式和赞助商隐性需求，我们提供', 'black', False, 11)
    add_run(p, '三级合作模式', 'red', True)
    add_run(p, '，不只是赞助，而是', 'black', False, 11)
    add_run(p, '红色文化工程战略合作', 'red', True)
    add_run(p, '：', 'black', False, 11)
    
    doc.add_paragraph()
    
    # 三级合作模式
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['战略合作伙伴', '联合合作伙伴', '支持单位']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(11)
    
    # 金额
    row = table.add_row()
    cells = row.cells
    cells[0].text = '500-800 万'
    cells[1].text = '200-300 万'
    cells[2].text = '50-100 万'
    
    # 名额
    row = table.add_row()
    cells = row.cells
    cells[0].text = '1-2 家'
    cells[1].text = '3-5 家'
    cells[2].text = '5-10 家'
    
    # 核心权益
    row = table.add_row()
    cells = row.cells
    cells[0].text = '内容深度绑定\n领导致辞 10 分钟\n政企对接会主办\nXR 体验区独家冠名\n分众 2 周全国'
    cells[1].text = '场景植入\n领导致辞 5 分钟\n政企对接会参与\nXR 体验区联合展示\n分众 1 周北京'
    cells[2].text = '名单展示\n政企对接会观摩\n分众 3 天北京'
    
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    # ========== 核心权益设计 ==========
    add_title(doc, '四、核心权益设计（满足隐性需求）')
    
    add_subtitle(doc, '4.1 政企对接权益（央企/国企最关心）')
    
    rights_gov = [
        '【首映礼政企对接会】',
        '• 战略合作伙伴：可主办专场对接会，邀请相关部委/北京市领导出席',
        '• 联合合作伙伴：可参加对接会，与领导面对面交流',
        '• 支持单位：可观摩对接会',
        '',
        '【政策信息获取】',
        '• 优先获取电影局、文资中心政策信息',
        '• 可参与红色文化项目政策研讨',
        '',
        '【领导政绩展示】',
        '• 企业参与红色文化工程可纳入党建考核',
        '• 提供证书、奖牌，可用于企业展厅展示',
        '• 媒体通稿中突出企业社会责任'
    ]
    
    for item in rights_gov:
        p = doc.add_paragraph()
        if item.startswith('【'):
            add_run(p, item, 'darkred', True)
        elif item.startswith('•'):
            p.paragraph_format.left_indent = Inches(0.5)
            add_run(p, item)
    
    doc.add_paragraph()
    
    add_subtitle(doc, '4.2 G 端业务拓展权益（科技公司最关心）')
    
    rights_g = [
        '【XR 体验区展示】',
        '• 战略合作伙伴：XR 体验区独家冠名"XXX 企业·长征 XR 体验"',
        '• 可展示企业 XR 技术、产品、解决方案',
        '• 政府领导参观时，企业技术人员可讲解',
        '',
        '【政府项目背书】',
        '• 提供"红色文化工程技术支持单位"证书',
        '• 可用于政府项目投标加分',
        '• 媒体通稿中突出企业技术实力',
        '',
        '【标杆案例打造】',
        '• XR 体验区可作为企业标杆案例',
        '• 可组织潜在客户参观体验',
        '• 长期授权 3 年使用"长征 XR 技术支持"称号'
    ]
    
    for item in rights_g:
        p = doc.add_paragraph()
        if item.startswith('【'):
            add_run(p, item, 'darkred', True)
        elif item.startswith('•'):
            p.paragraph_format.left_indent = Inches(0.5)
            add_run(p, item)
    
    doc.add_paragraph()
    
    add_subtitle(doc, '4.3 品牌安全与渠道权益（消费品牌最关心）')
    
    rights_c = [
        '【品牌安全】',
        '• 红色 IP，政治正确，无舆情风险',
        '• 可用于品牌长期背书',
        '',
        '【渠道下沉】',
        '• 分众媒体覆盖全国一二三线城市',
        '• 可借势开展"红色文化 + 产品"促销活动',
        '• 三四线城市消费者对红色 IP 认可度高',
        '',
        '【政企团购资源】',
        '• 可对接参会央企/国企采购资源',
        '• 企业福利采购、商务礼品采购',
        '• 借活动建立 B 端客户关系'
    ]
    
    for item in rights_c:
        p = doc.add_paragraph()
        if item.startswith('【'):
            add_run(p, item, 'darkred', True)
        elif item.startswith('•'):
            p.paragraph_format.left_indent = Inches(0.5)
            add_run(p, item)
    
    # ========== 内容深度绑定 ==========
    add_title(doc, '五、内容深度绑定（央视模式核心）')
    
    p = doc.add_paragraph()
    add_run(p, '借鉴央视《辉煌中国》《大国崛起》模式，我们提供内容深度绑定，不是简单贴片广告：', 'black', False, 11)
    
    doc.add_paragraph()
    
    content_rights = [
        '【影片片尾鸣谢】',
        '• 战略合作伙伴：片尾"特别鸣谢：XXX 企业"（单独一行）',
        '• 联合合作伙伴：片尾"联合支持：XXX 企业"（名单展示）',
        '',
        '【首映礼内容植入】',
        '• 企业领导可出席首映礼并致辞（5-10 分钟）',
        '• 可安排企业技术人员讲解 XR 体验',
        '• 媒体采访可安排企业代表',
        '',
        '【纪录片/花絮】',
        '• 制作《长征·英雄》幕后纪录片',
        '• 战略合作伙伴可融入企业参与红色文化工程的故事',
        '• 用于企业展厅、宣传片长期展示',
        '',
        '【XR 体验区内容定制】',
        '• 可根据企业需求定制 XR 体验内容',
        '• 如：科技企业可展示技术应用场景',
        '• 消费品牌可植入产品使用场景'
    ]
    
    for item in content_rights:
        p = doc.add_paragraph()
        if item.startswith('【'):
            add_run(p, item, 'darkred', True)
        elif item.startswith('•'):
            p.paragraph_format.left_indent = Inches(0.5)
            add_run(p, item)
    
    # ========== 投资回报 ==========
    add_title(doc, '六、投资回报分析')
    
    p = doc.add_paragraph()
    add_run(p, '以战略合作伙伴（500 万元）为例：', 'black', False, 11)
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['回报类型', '具体内容', '估算价值']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    
    roi_data = [
        ['分众广告', '2 周全国电梯媒体', '300 万'],
        ['央视/主流媒体', '新闻报道 + 专题', '200 万'],
        ['XR 体验区', '独家冠名 + 展示（3 个月）', '150 万'],
        ['政企对接', '政府关系建立', '无法量化'],
        ['G 端背书', '政府项目投标加分', '无法量化'],
        ['内容绑定', '片尾鸣谢 + 纪录片', '100 万'],
        ['长期授权', '3 年使用权', '150 万'],
        ['合计', '-', '900 万+']
    ]
    
    for row_data in roi_data:
        row = table.add_row()
        cells = row.cells
        for i, text in enumerate(row_data):
            cells[i].text = text
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '【ROI 结论】', 'darkred', True)
    p = doc.add_paragraph()
    add_run(p, '直接曝光价值约 900 万元，投资回报率 180%+。', 'black', False, 11)
    add_run(p, '政企关系、G 端背书等隐性价值无法量化，但对企业长期发展至关重要。', 'black', False, 11)
    
    # ========== 目标客户 ==========
    add_title(doc, '七、目标客户精准匹配')
    
    p = doc.add_paragraph()
    add_run(p, '基于隐性需求分析，我们精准匹配目标客户：', 'black', False, 11)
    
    doc.add_paragraph()
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    headers = ['客户类型', '代表企业', '核心需求', '匹配权益']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        table.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    
    targets = [
        ['央企/国企', '中石化、中国移动、工商银行', '政治任务、政府关系', '政企对接会、领导政绩展示'],
        ['科技公司', '华为、小米、字节', 'G 端业务、政府背书', 'XR 体验区、技术支持证书'],
        ['消费品牌', '茅台、伊利、安踏', '品牌安全、渠道下沉', '红色 IP 背书、分众媒体'],
        ['本地企业', '北京国企、文旅集团', '本地资源、政商关系', '政企对接、本地媒体曝光']
    ]
    
    for target in targets:
        row = table.add_row()
        cells = row.cells
        for i, text in enumerate(target):
            cells[i].text = text
            cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    
    # ========== 合作流程 ==========
    add_title(doc, '八、合作流程')
    
    steps = [
        '第 1 步：深度沟通（1-3 天）',
        '  - 了解企业核心需求（政治/G 端/品牌）',
        '  - 定制权益方案',
        '',
        '第 2 步：方案确认（3-5 天）',
        '  - 确认合作级别和权益',
        '  - 确认内容绑定方式',
        '',
        '第 3 步：合同签署（5-7 天）',
        '  - 法务审核',
        '  - 签署战略合作协议',
        '',
        '第 4 步：首付款（3 天内）',
        '  - 支付 50% 首付款',
        '  - 启动权益执行',
        '',
        '第 5 步：内容制作（首映礼前）',
        '  - 片尾鸣谢制作',
        '  - XR 体验区搭建',
        '  - 纪录片拍摄',
        '',
        '第 6 步：政企对接（首映礼前）',
        '  - 安排对接会',
        '  - 邀请政府领导',
        '',
        '第 7 步：活动落地（首映礼当天）',
        '  - 企业领导致辞',
        '  - 领导接见',
        '  - 媒体采访',
        '',
        '第 8 步：尾款结算（活动后 7 天）',
        '  - 支付 50% 尾款',
        '  - 提交执行报告',
        '',
        '第 9 步：长期运营（活动后）',
        '  - 授权企业使用赞助称号',
        '  - 持续政企对接服务'
    ]
    
    for step in steps:
        p = doc.add_paragraph()
        if step.startswith('第'):
            add_run(p, step, 'blue', True)
        elif step.startswith('  -'):
            p.paragraph_format.left_indent = Inches(0.5)
            add_run(p, step)
    
    # ========== 联系我们 ==========
    add_title(doc, '九、联系我们')
    
    p = doc.add_paragraph()
    add_run(p, '【项目总负责人】', 'darkred', True)
    p = doc.add_paragraph()
    add_run(p, '姓名：请填写\n电话：请填写\n邮箱：请填写')
    
    p = doc.add_paragraph()
    add_run(p, '【招商总监（央企/国企对接）】', 'darkred', True)
    p = doc.add_paragraph()
    add_run(p, '姓名：请填写\n电话：请填写\n邮箱：请填写')
    
    p = doc.add_paragraph()
    add_run(p, '【招商总监（科技公司对接）】', 'darkred', True)
    p = doc.add_paragraph()
    add_run(p, '姓名：请填写\n电话：请填写\n邮箱：请填写')
    
    p = doc.add_paragraph()
    add_run(p, '【媒体公关】', 'darkred', True)
    p = doc.add_paragraph()
    add_run(p, '姓名：请填写\n电话：请填写\n邮箱：请填写')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '不是卖广告位，是共建红色文化工程！', 'darkred', True, 13)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '2026 年 4 月', 'black', False, 10)
    
    return doc

if __name__ == '__main__':
    doc = create_proposal()
    output_path = '/home/admin/.openclaw/workspace-chief-agent/长征英雄赞助招商方案（央视模式升级版）.docx'
    doc.save(output_path)
    print('✅ 央视模式赞助方案已生成：' + output_path)
