#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
《长征·英雄》首映礼赞助方案 - 生成 Word 文档
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT

def add_run(paragraph, text, color='black', bold=False, font_size=12):
    """添加带颜色的文本片段"""
    run = paragraph.add_run(text)
    run.font.size = Pt(font_size)
    run.font.bold = bold
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    if color == 'red':
        run.font.color.rgb = RGBColor(255, 0, 0)
    elif color == 'blue':
        run.font.color.rgb = RGBColor(0, 0, 255)
    else:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return run

def add_title(doc, text):
    """添加标题"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.color.rgb = RGBColor(192, 0, 0)  # 深红色
    return p

def add_subtitle(doc, text):
    """添加副标题"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = 'SimSun'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def create_sponsorship_proposal():
    """创建赞助方案文档"""
    doc = Document()
    
    # 设置全局样式
    style = doc.styles['Normal']
    style.font.name = 'SimSun'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    style.font.size = Pt(11)
    
    # ========== 封面 ==========
    # 空几行
    for _ in range(5):
        doc.add_paragraph()
    
    # 主标题
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run('《长征·英雄》首映礼\n赞助招商方案')
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.name = 'SimSun'
    title_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    title_run.font.color.rgb = RGBColor(192, 0, 0)
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run('首个重大革命题材龙标影片\nXR 技术 + 红色 IP 创新呈现')
    subtitle_run.font.size = Pt(14)
    subtitle_run.font.name = 'SimSun'
    subtitle_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    subtitle_run.font.color.rgb = RGBColor(100, 100, 100)
    
    # 空行
    for _ in range(8):
        doc.add_paragraph()
    
    # 底部信息
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('2026 年 4 月\n联系方式：请填写')
    footer_run.font.size = Pt(12)
    footer_run.font.name = 'SimSun'
    footer_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
    
    # 分页
    doc.add_page_break()
    
    # ========== 目录 ==========
    add_title(doc, '目录')
    
    toc = [
        '一、项目背景',
        '二、核心价值',
        '三、首映礼概况',
        '四、赞助方案',
        '五、权益明细',
        '六、投资回报',
        '七、合作流程',
        '八、联系我们'
    ]
    
    for i, item in enumerate(toc, 1):
        p = doc.add_paragraph()
        add_run(p, f'{i}. {item}', 'blue', False, 12)
    
    doc.add_page_break()
    
    # ========== 一、项目背景 ==========
    add_title(doc, '一、项目背景')
    
    p = doc.add_paragraph()
    add_run(p, '《长征·英雄》', 'red', True)
    add_run(p, '已通过国家电影局审片，成为')
    add_run(p, '首个重大革命题材龙标影片', 'red', True)
    add_run(p, '。这是中国电影史上的里程碑事件，具有极高的政治价值、文化价值和商业价值。')
    
    p = doc.add_paragraph()
    add_run(p, '【核心亮点】', 'black', True)
    
    items = [
        '首个重大革命题材龙标——稀缺性、唯一性',
        'XR 技术呈现——沉浸式体验，科技 + 红色创新',
        '2026 年建党 105 周年——政策窗口期',
        '政府背书——中宣部、电影局支持',
        '分众传媒加持——可量化曝光效果'
    ]
    
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        add_run(p, '✓ ', 'red')
        add_run(p, item)
    
    # ========== 二、核心价值 ==========
    add_title(doc, '二、核心价值')
    
    p = doc.add_paragraph()
    add_run(p, '赞助商获得的不仅是品牌曝光，更是：', 'black', False, 12)
    
    # 创建表格
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    headers = ['政治资产', '品牌安全', '商业价值']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 数据行
    row = table.add_row()
    cells = row.cells
    cells[0].text = '政府背书\n政治正确\n国企可投'
    cells[1].text = '红色 IP\n正面形象\n无舆情风险'
    cells[2].text = '分众曝光\n可量化\n可转化'
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '【赞助商核心收益】', 'black', True)
    
    benefits = [
        '政治正确：响应国家号召，展示企业社会责任感',
        '品牌安全：红色 IP，正面形象，无舆情风险',
        '政府关系：与出席领导建立联系，获取政府资源',
        '媒体曝光：分众电梯媒体 + 主流媒体通稿',
        '技术展示：XR 体验区冠名，展示企业科技实力',
        '长期价值：可用于企业展厅、宣传片长期使用'
    ]
    
    for benefit in benefits:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        add_run(p, '• ', 'red')
        add_run(p, benefit)
    
    # ========== 三、首映礼概况 ==========
    add_title(doc, '三、首映礼概况')
    
    info_items = [
        ('活动名称', '《长征·英雄》首映礼暨 XR 体验展'),
        ('活动时间', '2026 年 5 月（待定）'),
        ('活动地点', '北京人民大会堂/国家大剧院（待定）'),
        ('活动规模', '500-800 人'),
        ('出席嘉宾', '政府领导、电影主创、赞助商代表、媒体记者'),
        ('媒体支持', '央视、新华社、人民日报、分众传媒')
    ]
    
    for label, value in info_items:
        p = doc.add_paragraph()
        add_run(p, f'{label}：', 'black', True)
        add_run(p, value)
    
    # ========== 四、赞助方案 ==========
    add_title(doc, '四、赞助方案')
    
    p = doc.add_paragraph()
    add_run(p, '我们提供三级赞助方案，满足不同企业需求：', 'black', False, 12)
    
    doc.add_paragraph()
    
    # 创建赞助方案表格
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # 表头
    headers = ['权益', '冠名赞助', '联合赞助', '支持单位']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 数据
    data = [
        ['赞助金额', '500 万元', '200 万元', '50 万元'],
        ['名额', '1 家', '3 家', '5 家'],
        ['首映礼 Logo', '主背景板 C 位', '背景板展示', '名单展示'],
        ['领导致辞', '5 分钟', '3 分钟', '-'],
        ['媒体曝光', '全部通稿署名', '部分通稿', '名单'],
        ['分众广告', '2 周全国投放', '1 周北京投放', '3 天北京投放'],
        ['XR 体验区', '冠名权', '联合展示', '-'],
        ['政府对接', '优先安排', '可参加', '-']
    ]
    
    for row_data in data:
        row = table.add_row()
        cells = row.cells
        for i, text in enumerate(row_data):
            cells[i].text = text
            cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    
    doc.add_paragraph()
    
    # ========== 五、权益明细 ==========
    add_title(doc, '五、权益明细')
    
    add_subtitle(doc, '5.1 冠名赞助权益（500 万元，1 家）')
    
    rights = [
        '【首映礼权益】',
        '• 主背景板 Logo C 位展示',
        '• 企业代表 5 分钟致辞',
        '• 与政府领导合影优先安排',
        '• 贵宾席前排就座',
        '',
        '【媒体曝光权益】',
        '• 全部媒体通稿署名"冠名赞助：XXX 企业"',
        '• 分众电梯媒体 2 周全国投放（含 Logo）',
        '• 央视新闻报道口播感谢',
        '',
        '【XR 体验区权益】',
        '• XR 体验区冠名"XXX 企业·长征 XR 体验"',
        '• 体验区内企业产品展示',
        '• 可安排技术讲解员',
        '',
        '【政府对接权益】',
        '• 优先安排与出席领导见面',
        '• 可参与后续政府项目对接',
        '',
        '【长期使用权】',
        '• 可在企业展厅、宣传片使用"《长征·英雄》冠名赞助企业"称号',
        '• 授权期限：3 年'
    ]
    
    for item in rights:
        p = doc.add_paragraph()
        if item.startswith('【'):
            add_run(p, item, 'red', True)
        elif item.startswith('•'):
            p.paragraph_format.left_indent = Inches(0.5)
            add_run(p, item)
        elif item == '':
            pass  # 空行
    
    doc.add_paragraph()
    
    add_subtitle(doc, '5.2 联合赞助权益（200 万元，3 家）')
    
    rights_2 = [
        '【首映礼权益】',
        '• 背景板 Logo 展示（次于冠名）',
        '• 企业代表 3 分钟致辞',
        '• 贵宾席就座',
        '',
        '【媒体曝光权益】',
        '• 部分媒体通稿署名',
        '• 分众电梯媒体 1 周北京投放',
        '',
        '【XR 体验区权益】',
        '• XR 体验区联合展示',
        '',
        '【长期使用权】',
        '• 可在企业宣传中使用"《长征·英雄》联合赞助企业"称号',
        '• 授权期限：2 年'
    ]
    
    for item in rights_2:
        p = doc.add_paragraph()
        if item.startswith('【'):
            add_run(p, item, 'red', True)
        elif item.startswith('•'):
            p.paragraph_format.left_indent = Inches(0.5)
            add_run(p, item)
    
    doc.add_paragraph()
    
    add_subtitle(doc, '5.3 支持单位权益（50 万元，5 家）')
    
    rights_3 = [
        '【首映礼权益】',
        '• 现场名单展示',
        '• 普通席位',
        '',
        '【媒体曝光权益】',
        '• 支持单位名单通稿',
        '• 分众电梯媒体 3 天北京投放',
        '',
        '【长期使用权】',
        '• 可在企业宣传中使用"《长征·英雄》支持单位"称号',
        '• 授权期限：1 年'
    ]
    
    for item in rights_3:
        p = doc.add_paragraph()
        if item.startswith('【'):
            add_run(p, item, 'red', True)
        elif item.startswith('•'):
            p.paragraph_format.left_indent = Inches(0.5)
            add_run(p, item)
    
    # ========== 六、投资回报 ==========
    add_title(doc, '六、投资回报分析')
    
    p = doc.add_paragraph()
    add_run(p, '以冠名赞助（500 万元）为例：', 'black', False, 12)
    
    doc.add_paragraph()
    
    # ROI 表格
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    headers = ['回报类型', '具体内容', '估算价值']
    header_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        header_cells[i].paragraphs[0].runs[0].font.bold = True
    
    roi_data = [
        ['分众广告', '2 周全国电梯媒体投放', '300 万元'],
        ['媒体曝光', '央视 + 主流媒体报道', '200 万元'],
        ['XR 体验区', '冠名 + 展示（3 个月）', '100 万元'],
        ['政府对接', '政企关系建立', '无法量化'],
        ['长期授权', '3 年使用权', '100 万元'],
        ['合计', '-', '700 万元+']
    ]
    
    for row_data in roi_data:
        row = table.add_row()
        cells = row.cells
        for i, text in enumerate(row_data):
            cells[i].text = text
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    add_run(p, '【ROI 结论】', 'red', True)
    p = doc.add_paragraph()
    add_run(p, '直接曝光价值约 700 万元，投资回报率 140%+。', 'black', False, 12)
    add_run(p, '政府关系、品牌安全等隐性价值无法量化。', 'black', False, 12)
    
    # ========== 七、合作流程 ==========
    add_title(doc, '七、合作流程')
    
    steps = [
        '第 1 步：意向沟通（1-3 天）',
        '  - 双方初步沟通，了解合作意向',
        '  - 发送完整赞助方案',
        '',
        '第 2 步：方案定制（3-5 天）',
        '  - 根据企业需求定制权益方案',
        '  - 确定赞助级别和金额',
        '',
        '第 3 步：合同签署（5-7 天）',
        '  - 法务审核合同',
        '  - 双方签署赞助协议',
        '',
        '第 4 步：首付款（3 天内）',
        '  - 支付赞助金额 50% 首付款',
        '  - 启动权益执行',
        '',
        '第 5 步：活动执行（首映礼前）',
        '  - 分众广告投放',
        '  - 媒体通稿准备',
        '  - XR 体验区搭建',
        '',
        '第 6 步：活动落地（首映礼当天）',
        '  - 企业代表出席',
        '  - 领导接见',
        '  - 媒体采访',
        '',
        '第 7 步：尾款结算（活动后 7 天）',
        '  - 支付剩余 50% 尾款',
        '  - 提交执行报告',
        '',
        '第 8 步：长期授权（活动后）',
        '  - 授权企业使用赞助称号',
        '  - 持续跟踪服务'
    ]
    
    for step in steps:
        p = doc.add_paragraph()
        if step.startswith('第'):
            add_run(p, step, 'blue', True)
        elif step.startswith('  -'):
            p.paragraph_format.left_indent = Inches(0.5)
            add_run(p, step)
    
    # ========== 八、联系我们 ==========
    add_title(doc, '八、联系我们')
    
    p = doc.add_paragraph()
    add_run(p, '【项目总负责人】', 'red', True)
    p = doc.add_paragraph()
    add_run(p, '姓名：请填写\n电话：请填写\n邮箱：请填写')
    
    p = doc.add_paragraph()
    add_run(p, '【招商总监】', 'red', True)
    p = doc.add_paragraph()
    add_run(p, '姓名：请填写\n电话：请填写\n邮箱：请填写')
    
    p = doc.add_paragraph()
    add_run(p, '【媒体公关】', 'red', True)
    p = doc.add_paragraph()
    add_run(p, '姓名：请填写\n电话：请填写\n邮箱：请填写')
    
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '期待与您携手，共创红色 IP 商业价值！', 'red', True, 14)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(p, '2026 年 4 月', 'black', False, 10)
    
    return doc

if __name__ == '__main__':
    doc = create_sponsorship_proposal()
    output_path = '/home/admin/.openclaw/workspace-chief-agent/长征英雄赞助招商方案.docx'
    doc.save(output_path)
    print('✅ 赞助方案已生成：' + output_path)
