#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成高校思政 VR 百校复制方案 Word 文档
"""

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

def create_word_document():
    doc = Document()
    
    # 标题样式
    title = doc.add_heading('高校思政 VR 百校复制方案', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 版本信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('版本号：V1.0\n').bold = True
    info.add_run('编制日期：2026 年 4 月\n')
    info.add_run('编制单位：[公司名称]')
    
    doc.add_page_break()
    
    # 目录
    doc.add_heading('目录', level=1)
    sections = [
        '1. 执行摘要',
        '2. 项目背景与政策风口',
        '3. 桂林学院标杆案例',
        '4. 三档合作方案',
        '5. 财务测算与投资回报',
        '6. 百校复制路线图',
        '7. 政策支持与申报指南',
        '8. 交付标准与服务承诺',
        '9. 常见问题 FAQ',
        '10. 附录'
    ]
    for section in sections:
        doc.add_paragraph(section)
    
    doc.add_page_break()
    
    # 1. 执行摘要
    doc.add_heading('1. 执行摘要', level=1)
    
    doc.add_heading('1.1 项目定位', level=2)
    doc.add_paragraph('本项目是全国首个高校思政 VR 轻资产运营标杆项目，以桂林学院为起点，计划 3 年内复制推广至 100 所高校，打造沉浸式思政教育新生态。')
    
    doc.add_heading('1.2 核心优势', level=2)
    advantages = [
        '政策强力支持：教育部大力推进"大思政课"建设，VR+ 教育是重点支持方向',
        '商业模式验证：桂林学院项目已验证可行性，毛利率达 42.5%',
        '轻资产运营：设备租赁 + 内容订阅模式，降低院校一次性投入压力',
        '标准化交付：单校部署周期≤30 天，可快速复制'
    ]
    for adv in advantages:
        doc.add_paragraph(adv, style='List Bullet')
    
    doc.add_heading('1.3 投资规模与回报', level=2)
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '方案档次'
    hdr_cells[1].text = '投资金额'
    hdr_cells[2].text = 'VR 设备数'
    hdr_cells[3].text = '投资回收期'
    
    data = [
        ('标配', '100 万元', '20 台', '18-24 个月'),
        ('高配', '200 万元', '40 台', '20-26 个月'),
        ('顶配', '500 万元', '100 台', '24-30 个月')
    ]
    for i, row_data in enumerate(data):
        row_cells = table.rows[i+1].cells
        for j, cell_text in enumerate(row_data):
            row_cells[j].text = cell_text
    
    doc.add_heading('1.4 三年目标', level=2)
    goals = [
        '2026 年：完成 10 所高校部署，实现收入 1000 万元',
        '2027 年：完成 50 所高校部署，实现收入 5000 万元',
        '2028 年：完成 100 所高校部署，实现收入 1 亿元'
    ]
    for goal in goals:
        doc.add_paragraph(goal, style='List Bullet')
    
    doc.add_page_break()
    
    # 2. 项目背景与政策风口
    doc.add_heading('2. 项目背景与政策风口', level=1)
    
    doc.add_heading('2.1 思政教育现状与痛点', level=2)
    doc.add_paragraph('传统思政教育存在教学方式单一、内容抽象难懂、实践环节薄弱、效果评估困难等问题。Z 世代大学生作为数字原住民，更习惯沉浸式、互动式学习方式。')
    
    doc.add_heading('2.2 VR 技术在教育中的应用优势', level=2)
    doc.add_paragraph('VR 技术具有沉浸式体验、互动性强、场景还原、安全可控、数据追踪等技术优势，能够提升学习兴趣、加深理解记忆、突破时空限制、实现规模化覆盖。')
    
    doc.add_heading('2.3 政策支持力度', level=2)
    policies = [
        '《全面推进"大思政课"建设的工作方案》（教育部等十部门，2022 年）',
        '《虚拟现实与行业应用融合发展行动计划（2022-2026 年）》（工信部等五部门）',
        '《新时代学校思想政治理论课改革创新实施方案》（教育部，2020 年）'
    ]
    for policy in policies:
        doc.add_paragraph(policy, style='List Bullet')
    
    doc.add_page_break()
    
    # 3. 桂林学院标杆案例
    doc.add_heading('3. 桂林学院标杆案例', level=1)
    
    doc.add_heading('3.1 项目概况', level=2)
    doc.add_paragraph('合作院校：桂林学院 | 签约时间：2025 年 9 月 | 正式运营：2025 年 12 月')
    doc.add_paragraph('投资金额：100 万元 | 设备规模：20 台 VR 一体机 | 场地面积：80 平方米')
    doc.add_paragraph('服务师生：年均 3000+ 人次')
    
    doc.add_heading('3.2 建设内容', level=2)
    doc.add_paragraph('硬件配置：20 台 PICO 4 Enterprise VR 一体机、管理主机、显示设备、网络设备、充电柜等')
    doc.add_paragraph('软件内容：12 门 VR 思政精品课程，涵盖党史教育、国情教育、红色文化、价值观教育等模块')
    
    doc.add_heading('3.3 运营数据', level=2)
    doc.add_paragraph('日均使用时长：6 小时 | 周均体验人次：150 人 | 月均体验人次：600 人')
    doc.add_paragraph('教学效果：学习兴趣提升 41.5%，知识掌握提升 22.2%，课堂参与提升 62.1%，记忆留存提升 73.3%')
    
    doc.add_heading('3.4 财务表现', level=2)
    doc.add_paragraph('年服务收入：约 60 万元 | 年运营成本：约 34.5 万元 | 年毛利润：约 25.5 万元')
    doc.add_paragraph('毛利率：42.5% | 投资回收期：约 20 个月')
    
    doc.add_page_break()
    
    # 4. 三档合作方案
    doc.add_heading('4. 三档合作方案', level=1)
    
    doc.add_heading('4.1 标配方案（100 万元）', level=2)
    doc.add_paragraph('适用对象：首次尝试 VR 思政教育的院校、预算有限的普通本科院校')
    doc.add_paragraph('配置：20 台 VR 设备、12 门标准课程、80㎡VR 实训室、2 天师资培训、3 个月运营指导')
    doc.add_paragraph('交付周期：合同签订后 30 天内完成交付')
    
    doc.add_heading('4.2 高配方案（200 万元）', level=2)
    doc.add_paragraph('适用对象：重点马院建设院校、省级思政示范院校')
    doc.add_paragraph('配置：40 台 VR 设备、20 门课程 +2 门定制、150㎡VR 实训中心、5 天师资培训、6 个月驻校支持')
    doc.add_paragraph('交付周期：合同签订后 45 天内完成交付')
    
    doc.add_heading('4.3 顶配方案（500 万元）', level=2)
    doc.add_paragraph('适用对象：双一流高校、省级思政教育中心、区域示范标杆项目')
    doc.add_paragraph('配置：100 台 VR 设备、30 门课程 +5 门定制、300㎡VR 思政中心、10 天系统培训、12 个月驻校支持')
    doc.add_paragraph('交付周期：合同签订后 60 天内完成交付')
    
    doc.add_page_break()
    
    # 5. 财务测算与投资回报
    doc.add_heading('5. 财务测算与投资回报', level=1)
    
    doc.add_heading('5.1 收入模型', level=2)
    doc.add_paragraph('收入来源：设备租赁费（30%）、内容订阅费（15%）、运营服务费（10%）、定制开发费（按需）')
    doc.add_paragraph('单校年收入：标配 60 万、高配 120 万、顶配 300 万')
    
    doc.add_heading('5.2 成本模型', level=2)
    doc.add_paragraph('成本构成：设备折旧 43.5%、内容摊销 23.2%、运营人力 17.4%、维护费用 8.7%、其他费用 7.2%')
    doc.add_paragraph('单校年成本：标配 34.5 万、高配 69 万、顶配 172.5 万')
    
    doc.add_heading('5.3 盈利分析', level=2)
    doc.add_paragraph('单校年毛利：标配 25.5 万、高配 51 万、顶配 127.5 万')
    doc.add_paragraph('毛利率：42.5%（三档一致）')
    doc.add_paragraph('投资回收期：约 20 个月（三档一致）')
    
    doc.add_heading('5.4 百校规模预测', level=2)
    doc.add_paragraph('2026 年：10 所院校，收入 1000 万')
    doc.add_paragraph('2027 年：50 所院校，收入 5000 万')
    doc.add_paragraph('2028 年：100 所院校，收入 1 亿')
    doc.add_paragraph('三年累计收入：1.6 亿元，累计毛利：6800 万元')
    
    doc.add_page_break()
    
    # 6. 百校复制路线图
    doc.add_heading('6. 百校复制路线图', level=1)
    
    doc.add_heading('6.1 总体目标', level=2)
    doc.add_paragraph('3 年 100 校，打造全国高校思政 VR 教育第一品牌')
    
    doc.add_heading('6.2 阶段规划', level=2)
    doc.add_paragraph('第一阶段（2026 Q2-Q4）：试点验证期，完成 10 所高校部署，收入 1000 万')
    doc.add_paragraph('第二阶段（2027 全年）：快速扩张期，完成 50 所高校部署，收入 5000 万')
    doc.add_paragraph('第三阶段（2028 全年）：全面覆盖期，完成 100 所高校部署，收入 1 亿')
    
    doc.add_heading('6.3 区域布局', level=2)
    doc.add_paragraph('华东 25 所、华北 20 所、华南 15 所、华中 15 所、西南 15 所、西北 10 所')
    
    doc.add_page_break()
    
    # 7. 政策支持与申报指南
    doc.add_heading('7. 政策支持与申报指南', level=1)
    
    doc.add_heading('7.1 可申报的政策项目', level=2)
    doc.add_paragraph('国家级：大思政课示范项目（50-100 万）、虚拟仿真实验教学项目（30-80 万）')
    doc.add_paragraph('省级：思政工作精品项目（20-50 万）、智慧马院建设（30-100 万）、虚拟仿真基地（50-200 万）')
    
    doc.add_heading('7.2 申报策略', level=2)
    doc.add_paragraph('最佳时机：项目落地后 3-6 个月（有运营数据支撑）')
    doc.add_paragraph('申报材料：项目申报书、可行性研究报告、合作协议、运营数据报告等')
    
    doc.add_page_break()
    
    # 8. 交付标准与服务承诺
    doc.add_heading('8. 交付标准与服务承诺', level=1)
    
    doc.add_heading('8.1 交付标准', level=2)
    doc.add_paragraph('硬件交付：VR 设备 100% 开机正常，网络延迟≤20ms')
    doc.add_paragraph('软件交付：内容平台正常运行，课程完整可用，账号创建完成')
    doc.add_paragraph('场地交付：空间布局合理，装修质量合格，安全设施完备')
    
    doc.add_heading('8.2 服务承诺', level=2)
    doc.add_paragraph('培训服务：基础培训 1 天、教学培训 1-4 天、认证培训 5-10 天')
    doc.add_paragraph('运维服务：电话咨询即时响应、远程支持≤30 分钟、现场服务≤4 小时')
    doc.add_paragraph('内容更新：季度常规更新、重大节日专题更新、校本定制开发')
    
    doc.add_heading('8.3 质量保证', level=2)
    doc.add_paragraph('质保期限：硬件 2 年、软件 1 年、装修 1 年、内容合同期内')
    doc.add_paragraph('服务 SLA：系统可用性≥99%、故障响应≤30 分钟、故障解决≤24 小时')
    
    doc.add_page_break()
    
    # 9. 常见问题 FAQ
    doc.add_heading('9. 常见问题 FAQ', level=1)
    
    faqs = [
        ('Q1：VR 设备会不会让学生头晕？', 'A：选用高分辨率（3664×1920）和高刷新率（90Hz）设备，95% 以上学生无眩晕感。'),
        ('Q2：网络条件要求高吗？', 'A：本地内容无需网络，在线更新需要稳定网络。建议带宽≥100Mbps。'),
        ('Q3：教师没有 VR 经验怎么办？', 'A：提供系统化培训，教师经过 2 天培训即可独立开展 VR 教学。'),
        ('Q4：投资是一次性还是分期？', 'A：支持一次性付款和分期付款。分期付款通常为 30% 首付 +40% 交付 +30% 验收。'),
        ('Q5：能否申请政府专项资金？', 'A：可以。提供全套申报材料支持，可申报 20-200 万元不等的专项资金。')
    ]
    for question, answer in faqs:
        doc.add_paragraph(question, style='List Bullet')
        doc.add_paragraph(answer)
    
    doc.add_page_break()
    
    # 10. 附录
    doc.add_heading('10. 附录', level=1)
    
    doc.add_heading('10.1 设备清单（标配方案）', level=2)
    doc.add_paragraph('VR 一体机 20 台、管理主机 1 台、显示设备 2 台、充电存储柜 1 台、网络设备 1 套等')
    
    doc.add_heading('10.2 课程内容目录', level=2)
    doc.add_paragraph('党史教育 5 门、国情教育 4 门、红色文化 4 门、价值观教育 3 门、校史教育 1 门、实践实训 3 门')
    
    doc.add_heading('10.3 项目申报材料清单', level=2)
    doc.add_paragraph('项目申报书、可行性研究报告、合作协议、学校配套资金承诺函、案例材料、运营数据报告等 14 项')
    
    doc.add_page_break()
    
    # 结尾信息
    end_info = doc.add_paragraph()
    end_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_info.add_run('\n编制单位：[公司名称]\n').bold = True
    end_info.add_run('联系人：[待填写]\n')
    end_info.add_run('联系电话：[待填写]\n')
    end_info.add_run('电子邮箱：[待填写]\n')
    end_info.add_run('编制日期：2026 年 4 月\n')
    
    # 保存文档
    doc.save('百校复制方案 - 完整版.docx')
    print('Word 文档生成成功：百校复制方案 - 完整版.docx')

if __name__ == '__main__':
    create_word_document()
