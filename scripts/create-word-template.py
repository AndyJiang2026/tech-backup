#!/usr/bin/env python3
"""
🐴 强国小马 - 创建专业 Word 模板
用途：生成标准化 Word 文档模板
"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

def create_standard_report_template(output_path):
    """创建标准报告模板"""
    
    doc = Document()
    
    # 设置页面边距（标准公文边距）
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(3.17)
        section.bottom_margin = Cm(3.17)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
    
    # 添加标题
    title = doc.add_heading('报告标题', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.font.name = 'Microsoft YaHei'
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    
    # 添加副标题
    subtitle = doc.add_heading('副标题（可选）', level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.name = 'Microsoft YaHei'
    subtitle.runs[0].font.size = Pt(16)
    
    # 添加一级标题
    h1 = doc.add_heading('一、一级标题', level=1)
    h1.runs[0].font.name = 'Microsoft YaHei'
    h1.runs[0].font.size = Pt(16)
    h1.runs[0].font.bold = True
    
    # 添加正文
    p = doc.add_paragraph()
    run = p.add_run('这是正文内容示例。中文使用宋体或微软雅黑，英文使用 Times New Roman。行间距设置为 1.5 倍，段后间距 6pt。')
    run.font.name = 'SimSun'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    
    # 添加二级标题
    h2 = doc.add_heading('（一）二级标题', level=2)
    h2.runs[0].font.name = 'Microsoft YaHei'
    h2.runs[0].font.size = Pt(14)
    h2.runs[0].font.bold = True
    
    # 添加列表
    doc.add_paragraph('• 列表项 1', style='List Bullet')
    doc.add_paragraph('• 列表项 2', style='List Bullet')
    doc.add_paragraph('• 列表项 3', style='List Bullet')
    
    # 添加三级标题
    h3 = doc.add_heading('1. 三级标题', level=3)
    h3.runs[0].font.name = 'Microsoft YaHei'
    h3.runs[0].font.size = Pt(14)
    h3.runs[0].font.bold = True
    
    # 添加表格示例
    doc.add_heading('表格示例', level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # 表头
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '序号'
    hdr_cells[1].text = '项目'
    hdr_cells[2].text = '备注'
    
    # 设置表头样式
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.name = 'Microsoft YaHei'
        cell.paragraphs[0].runs[0].font.size = Pt(12)
    
    # 添加数据
    data = [
        ('1', '示例项目 1', '备注 1'),
        ('2', '示例项目 2', '备注 2'),
        ('3', '示例项目 3', '备注 3'),
    ]
    
    for i, row_data in enumerate(data):
        row_cells = table.rows[i + 1].cells
        for j, cell_data in enumerate(row_data):
            row_cells[j].text = cell_data
            row_cells[j].paragraphs[0].runs[0].font.name = 'SimSun'
            row_cells[j].paragraphs[0].runs[0].font.size = Pt(12)
    
    # 添加页脚（页码）
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('第 1 页')
    footer_run.font.name = 'SimSun'
    footer_run.font.size = Pt(9)
    
    # 保存模板
    doc.save(output_path)
    print(f"✅ 模板已保存：{output_path}")
    return output_path

if __name__ == "__main__":
    import sys
    output = sys.argv[1] if len(sys.argv) > 1 else "/home/admin/.openclaw/workspace-chief-agent/templates/word/标准报告模板.docx"
    create_standard_report_template(output)
    print("🎉 模板创建完成！")
