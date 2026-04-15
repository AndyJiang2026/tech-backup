#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成《长征·英雄》3-5 分钟 Demo 销售工具方案 Word 文档
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_word_document():
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 标题
    title = doc.add_heading('《长征·英雄》3-5 分钟 Demo 销售工具方案', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run('关键销售工具补充｜学习强国 + 龙标双重背书\n').italic = True
    subtitle.add_run('版本号：V1.0｜编制日期：2026 年 4 月')
    
    doc.add_paragraph()
    
    # 执行摘要
    doc.add_heading('执行摘要', level=1)
    doc.add_paragraph('3-5 分钟 Demo 视频是高校思政 VR 项目销售的标配工具，用于首次拜访、招标比稿、预算审批、教师培训等关键场景。本方案基于曼恒等行业头部厂商的销售经验验证，提供完整的内容设计、制作方案和预算建议。', style='Intense Quote')
    
    # 一、Demo 战略价值
    doc.add_heading('一、Demo 战略价值分析', level=1)
    
    doc.add_heading('1.1 为什么 Demo 如此重要？', level=2)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '销售场景'
    hdr_cells[1].text = '痛点'
    hdr_cells[2].text = 'Demo 解决方案'
    
    scenarios = [
        ('首次拜访', '校领导难以想象 VR 效果', '3 分钟直观展示，降低理解成本'),
        ('招标比稿', '多家竞品方案雷同', 'Demo 差异化呈现，脱颖而出'),
        ('预算审批', '财务/资产处质疑效果', '实证视频，增强决策信心'),
        ('教师培训', '老师担心不会用', 'Demo 展示操作流程，降低顾虑')
    ]
    for i, (scene, pain, solution) in enumerate(scenarios):
        table.rows[i+1].cells[0].text = scene
        table.rows[i+1].cells[1].text = pain
        table.rows[i+1].cells[2].text = solution
    
    doc.add_heading('1.2 行业经验验证', level=2)
    doc.add_paragraph('曼恒作为 LBE 大空间头部厂商，其销售经验验证了：', style='List Bullet')
    doc.add_paragraph('✅ Demo 是标配，不是选配', style='List Bullet')
    doc.add_paragraph('✅ 3-5 分钟是黄金时长（太短说不清，太长没耐心）', style='List Bullet')
    doc.add_paragraph('✅ 面向高校领导需要专门设计（不是产品说明书）', style='List Bullet')
    
    doc.add_page_break()
    
    # 二、Demo 内容设计
    doc.add_heading('二、Demo 内容设计建议', level=1)
    
    doc.add_heading('2.1 核心结构（3-5 分钟）', level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '时间段'
    hdr_cells[1].text = '内容'
    hdr_cells[2].text = '目标'
    
    structure = [
        ('0-30 秒', '开场：政策 + 痛点', '引起共鸣'),
        ('30-90 秒', '产品亮相：桂林学院实景', '建立信任'),
        ('90-180 秒', '核心体验：VR 课程片段', '展示效果'),
        ('180-240 秒', '学生反馈 + 数据', '验证效果'),
        ('240-300 秒', '合作方案 + 联系方式', '促成行动')
    ]
    for i, (time, content, goal) in enumerate(structure):
        table.rows[i+1].cells[0].text = time
        table.rows[i+1].cells[1].text = content
        table.rows[i+1].cells[2].text = goal
    
    doc.add_heading('2.2 Demo 形式建议', level=2)
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '形式'
    hdr_cells[1].text = '优点'
    hdr_cells[2].text = '缺点'
    hdr_cells[3].text = '建议'
    
    forms = [
        ('实拍视频', '真实可信', '需要场地和设备', '✅ 首选'),
        ('录屏演示', '成本低', '缺乏真实感', '辅助使用'),
        ('3D 动画', '视觉效果佳', '成本高', '暂不推荐'),
        ('混合剪辑', '兼顾真实与效果', '需要专业剪辑', '✅ 推荐')
    ]
    for i, (form, pros, cons, suggest) in enumerate(forms):
        table.rows[i+1].cells[0].text = form
        table.rows[i+1].cells[1].text = pros
        table.rows[i+1].cells[2].text = cons
        table.rows[i+1].cells[3].text = suggest
    
    doc.add_page_break()
    
    # 三、Demo 脚本框架
    doc.add_heading('三、Demo 脚本框架', level=1)
    
    doc.add_heading('3.1 开场：政策 + 痛点（0-30 秒）', level=2)
    doc.add_paragraph('画面：新闻联播片段 + 课堂实拍', style='Intense Quote')
    doc.add_paragraph('配音："教育部大力推进大思政课建设，但传统思政教育存在参与度低、场景感知弱、情感共鸣不足等痛点..."', style='Intense Quote')
    doc.add_paragraph('字幕：政策驱动 × 教育痛点')
    
    doc.add_heading('3.2 产品亮相：桂林学院实景（30-90 秒）', level=2)
    doc.add_paragraph('画面：桂林学院 VR 实训室外景→内景→学生体验', style='Intense Quote')
    doc.add_paragraph('配音："全国首个高校思政 VR 轻资产运营标杆项目——桂林学院，投资 100 万，20 台 VR 设备，年均服务 3000+ 人次..."', style='Intense Quote')
    doc.add_paragraph('字幕：桂林学院标杆案例')
    
    doc.add_heading('3.3 核心体验：VR 课程片段（90-180 秒）', level=2)
    doc.add_paragraph('画面：第一视角 VR 体验（长征场景）+ 学生反应', style='Intense Quote')
    doc.add_paragraph('配音："学生戴上 VR 设备，亲历长征场景，从\'观看历史\'到\'成为历史\'，学习兴趣提升 41.5%，记忆留存提升 73.3%..."', style='Intense Quote')
    doc.add_paragraph('字幕：沉浸式体验 × 数据验证')
    
    doc.add_heading('3.4 学生反馈 + 数据（180-240 秒）', level=2)
    doc.add_paragraph('画面：学生采访片段 + 数据图表', style='Intense Quote')
    doc.add_paragraph('配音："桂林学院学生反馈：\'第一次觉得思政课这么有意思\'，课程满意度 96 分，90% 学生主动推荐..."', style='Intense Quote')
    doc.add_paragraph('字幕：学生认可 × 效果显著')
    
    doc.add_heading('3.5 合作方案 + 联系方式（240-300 秒）', level=2)
    doc.add_paragraph('画面：三档方案对比 + 联系方式', style='Intense Quote')
    doc.add_paragraph('配音："标配 100 万、高配 200 万、顶配 500 万，交付周期 30-60 天，内容订阅持续服务...欢迎垂询，共创思政 VR 新生态！"', style='Intense Quote')
    doc.add_paragraph('字幕：三档方案 × 联系方式')
    
    doc.add_page_break()
    
    # 四、Demo 制作方案
    doc.add_heading('四、Demo 制作方案', level=1)
    
    doc.add_heading('4.1 方案 A：快速版', level=2)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    quick_info = [
        ('周期', '1-2 周'),
        ('拍摄方式', '用手机/相机拍摄桂林学院实景'),
        ('内容', '录屏 VR 课程体验片段'),
        ('后期', '简单剪辑 + 配音 + 字幕'),
        ('预算', '1-2 万'),
        ('适用', '初期销售急需')
    ]
    for i, (item, value) in enumerate(quick_info):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = value
    
    doc.add_heading('4.2 方案 B：专业版', level=2)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    pro_info = [
        ('周期', '3-4 周'),
        ('拍摄方式', '专业团队拍摄 + 灯光 + 收音'),
        ('内容', '第一视角 VR 录屏（高画质）'),
        ('后期', '专业剪辑 + 配乐 + 配音'),
        ('预算', '5-8 万'),
        ('适用', '正式推广')
    ]
    for i, (item, value) in enumerate(pro_info):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = value
    
    doc.add_heading('4.3 方案 C：旗舰版', level=2)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    premium_info = [
        ('周期', '6-8 周'),
        ('拍摄方式', '多校取景（桂林 + 其他合作校）'),
        ('内容', '学生/教师采访纪录片风格'),
        ('后期', '3D 动画辅助说明技术原理'),
        ('预算', '15-20 万'),
        ('适用', '品牌宣传 + 大型展会')
    ]
    for i, (item, value) in enumerate(premium_info):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = value
    
    doc.add_page_break()
    
    # 五、制作方案对比
    doc.add_heading('五、制作方案对比', level=1)
    table = doc.add_table(rows=4, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '方案'
    hdr_cells[1].text = '周期'
    hdr_cells[2].text = '预算'
    hdr_cells[3].text = '效果'
    hdr_cells[4].text = '推荐度'
    
    plans = [
        ('快速版', '1-2 周', '1-2 万', '★★★', '⭐⭐⭐ 初期急需'),
        ('专业版', '3-4 周', '5-8 万', '★★★★', '⭐⭐⭐⭐⭐ 正式推广'),
        ('旗舰版', '6-8 周', '15-20 万', '★★★★★', '⭐⭐⭐⭐ 品牌宣传')
    ]
    for i, (plan, time, budget, quality, recommend) in enumerate(plans):
        table.rows[i+1].cells[0].text = plan
        table.rows[i+1].cells[1].text = time
        table.rows[i+1].cells[2].text = budget
        table.rows[i+1].cells[3].text = quality
        table.rows[i+1].cells[4].text = recommend
    
    doc.add_heading('推荐方案', level=2)
    doc.add_paragraph('先出快速版（1-2 周），满足初期销售急需；再迭代专业版（3-4 周），用于正式推广。', style='Intense Quote')
    
    # 六、待确认事项
    doc.add_heading('六、待确认事项', level=1)
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '事项'
    hdr_cells[1].text = '建议'
    hdr_cells[2].text = '待确认'
    
    items = [
        ('制作时机', '先出快速版，再迭代专业版', '✅/❌'),
        ('拍摄场地', '桂林学院（已有案例）', '是否可协调？'),
        ('出镜人员', '学生 + 教师 + 校领导', '是否可安排？'),
        ('VR 课程内容', '《长征路上的抉择》', '是否有可用素材？'),
        ('制作预算', '快速版 1-2 万', '预算范围？'),
        ('制作方式', '内部团队 or 外包', '偏好？')
    ]
    for i, (item, suggest, confirm) in enumerate(items):
        table.rows[i+1].cells[0].text = item
        table.rows[i+1].cells[1].text = suggest
        table.rows[i+1].cells[2].text = confirm
    
    doc.add_page_break()
    
    # 七、核心背书
    doc.add_heading('七、Demo 核心背书', level=1)
    doc.add_paragraph('Demo 视频需突出以下核心背书：', style='List Bullet')
    doc.add_paragraph('🎬 国家电影局 首个重大革命题材 VR"龙标"电影', style='List Bullet')
    doc.add_paragraph('🔴 中央宣传部宣传舆情研究中心', style='List Bullet')
    doc.add_paragraph('🔴 "学习强国"学习平台 领衔出品', style='List Bullet')
    doc.add_paragraph('🏆 2026 年广西高校思政工作精品项目', style='List Bullet')
    
    doc.add_heading('八、运营数据支撑', level=1)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    data = [
        ('日均体验人次', '150 人'),
        ('课程满意度', '96 分'),
        ('学习兴趣提升', '+41.5%'),
        ('记忆留存提升', '+73.3%'),
        ('媒体报道', '8 家媒体（广西日报、桂林电视台等）'),
        ('参观接待', '15 批次区内外高校参观团')
    ]
    for i, (item, value) in enumerate(data):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = value
    
    # 结尾信息
    doc.add_paragraph()
    doc.add_paragraph()
    end_info = doc.add_paragraph()
    end_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_info.add_run('编制单位：[公司名称]\n').bold = True
    end_info.add_run('版本号：V1.0\n')
    end_info.add_run('更新日期：2026 年 4 月\n')
    end_info.add_run('\n本方案版权归 [学习强国学习平台有限责任公司 未来新视界科技 (北京) 有限公司] 所有，未经许可不得外传').italic = True
    
    # 保存文档
    doc.save('/home/admin/.openclaw/workspace-chief-agent/长征 - 英雄 -3-5 分钟 Demo 销售工具方案 V1.0.docx')
    print('✅ Word 文档生成成功：长征 - 英雄 -3-5 分钟 Demo 销售工具方案 V1.0.docx')

if __name__ == '__main__':
    create_word_document()
