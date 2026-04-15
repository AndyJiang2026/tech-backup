#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成《长征·英雄》城市合伙人合作模式优化方案 V2.0 Word 文档
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_word_document():
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 标题
    title = doc.add_heading('《长征·英雄》城市合伙人\n合作模式优化方案 V2.0', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 版本信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('编制日期：2026 年 4 月 13 日\n')
    info.add_run('优化背景：福建/咸阳/内蒙古三地测算反馈"难以回本"\n')
    info.add_run('核心目标：降低门槛 + 多元收入 + 快速回本')
    
    doc.add_paragraph()
    
    # 执行摘要
    doc.add_heading('执行摘要', level=1)
    
    doc.add_heading('一、问题诊断', level=2)
    doc.add_paragraph('原合作模式（50 万授权费 +30 元门票分账）经福建海峡出版、咸阳文旅集团、内蒙古广电三家客户测算后，反馈"难以回本"，最终导致目标客户先后放弃合作。', style='Intense Quote')
    
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '客户'
    hdr_cells[1].text = '原授权费'
    hdr_cells[2].text = '回本周期'
    hdr_cells[3].text = '决策'
    
    clients = [
        ('福建海峡出版', '50 万', '3.5 年', '❌ 放弃'),
        ('咸阳文旅集团', '50 万', '3 年', '❌ 放弃'),
        ('内蒙古广电', '50 万', '4 年', '❌ 放弃')
    ]
    for i, (client, fee, payback, decision) in enumerate(clients):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = client
        row_cells[1].text = fee
        row_cells[2].text = payback
        row_cells[3].text = decision
    
    doc.add_heading('核心问题', level=3)
    problems = [
        '授权费过高（50 万 upfront，客户心理门槛高）',
        '收入来源单一（仅靠门票分账，抗风险能力弱）',
        '客群定位模糊（党员/红培/国防教育未差异化运营）',
        '增值服务缺失（无培训/认证/申报等增值收入）'
    ]
    for problem in problems:
        doc.add_paragraph(problem, style='List Bullet')
    
    doc.add_heading('二、优化后合作模式（V2.0）', level=2)
    doc.add_paragraph('三层合作模式设计，满足不同客户需求', style='Intense Quote')
    
    table = doc.add_table(rows=4, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '版本'
    hdr_cells[1].text = '授权费'
    hdr_cells[2].text = '门票分账'
    hdr_cells[3].text = '增值服务'
    hdr_cells[4].text = '回本周期'
    
    versions = [
        ('轻量版', '15 万/年', '30 元/人次', '基础服务', '1.5 年'),
        ('标准版⭐', '30 万/年', '25 元/人次', '培训 + 申报', '1 年'),
        ('旗舰版', '80 万/3 年', '20 元/人次', '全部服务', '8 个月')
    ]
    for i, (version, fee, ticket, service, payback) in enumerate(versions):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = version
        row_cells[1].text = fee
        row_cells[2].text = ticket
        row_cells[3].text = service
        row_cells[4].text = payback
    
    doc.add_heading('2.1 轻量版（15 万/年）', level=3)
    doc.add_paragraph('适合试探性合作、客流不确定客户')
    details = [
        '授权费：15 万元/年（降低 70% 门槛）',
        '授权范围：单城市单场馆',
        '门票分账：30 元/人次',
        '基础服务：内容更新 + 平台维护 + 远程运维',
        '回本周期：5000 人次/年（约 1.5 年回本）'
    ]
    for detail in details:
        doc.add_paragraph(detail, style='List Bullet')
    
    doc.add_heading('2.2 标准版（30 万/年）⭐推荐', level=3)
    doc.add_paragraph('适合主流客户（福建/咸阳/内蒙古级别）')
    details = [
        '授权费：30 万元/年',
        '授权范围：单城市单场馆',
        '门票分账：25 元/人次（优惠 17%）',
        '增值服务：党员培训认证 + 红培研学课程包',
        '政策申报：协助申报省级思政/研学基地（20-50 万补贴）',
        '回本周期：8000 人次/年 + 增值服务（约 1 年回本）'
    ]
    for detail in details:
        doc.add_paragraph(detail, style='List Bullet')
    
    doc.add_heading('2.3 旗舰版（80 万/3 年）', level=3)
    doc.add_paragraph('适合省级出版集团/大型文旅集团')
    details = [
        '授权费：80 万元/3 年（相当于 26.7 万/年）',
        '授权范围：单城市独家 + 优先续约权',
        '门票分账：20 元/人次（优惠 33%）',
        '增值服务：全部增值服务（培训 + 认证 + 申报 + 定制）',
        '定制内容：本地红色资源 VR 化（1 个项目）',
        '回本周期：1 万人次/年 + 增值服务（约 8 个月回本）'
    ]
    for detail in details:
        doc.add_paragraph(detail, style='List Bullet')
    
    doc.add_page_break()
    
    # 多元收入结构
    doc.add_heading('三、多元收入结构设计', level=1)
    
    doc.add_heading('3.1 收入来源对比', level=2)
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '收入类型'
    hdr_cells[1].text = '原模式'
    hdr_cells[2].text = 'V2.0 优化'
    
    income_types = [
        ('门票分账', '✅ 30 元/人次', '✅ 20-30 元/人次（分层）'),
        ('党员培训', '❌ 无', '✅ 500 元/人次（认证费）'),
        ('红培研学', '❌ 无', '✅ 200 元/学生（课程包）'),
        ('国防教育', '❌ 无', '✅ 300 元/学生（军训拓展）'),
        ('政策申报', '❌ 无', '✅ 20-50 万/项目（补贴）'),
        ('定制内容', '❌ 无', '✅ 10-30 万/项目')
    ]
    for i, (income_type, old, new) in enumerate(income_types):
        table.rows[i+1].cells[0].text = income_type
        table.rows[i+1].cells[1].text = old
        table.rows[i+1].cells[2].text = new
    
    doc.add_heading('3.2 标准版收入测算（以福建海峡出版为例）', level=2)
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '收入来源'
    hdr_cells[1].text = '单价'
    hdr_cells[2].text = '年人次/项目'
    hdr_cells[3].text = '年收入'
    
    income_details = [
        ('门票分账', '25 元', '8000 人次', '20 万'),
        ('党员培训', '500 元', '200 人次', '10 万'),
        ('红培研学', '200 元', '1000 学生', '20 万'),
        ('国防教育', '300 元', '300 学生', '9 万'),
        ('政策申报', '30 万', '1 项目/年', '30 万'),
        ('合计', '-', '-', '89 万')
    ]
    for i, (source, price, volume, total) in enumerate(income_details):
        table.rows[i].cells[0].text = source
        table.rows[i].cells[1].text = price
        table.rows[i].cells[2].text = volume
        table.rows[i].cells[3].text = total
    
    doc.add_paragraph()
    doc.add_paragraph('成本：30 万授权费 + 20 万运营 = 50 万/年', style='Intense Quote')
    doc.add_paragraph('毛利：89 万 - 50 万 = 39 万/年', style='Intense Quote')
    doc.add_paragraph('回本周期：第一年即盈利 ✅', style='Intense Quote')
    
    doc.add_page_break()
    
    # 分客群运营策略
    doc.add_heading('四、分客群运营策略', level=1)
    
    doc.add_heading('4.1 党员培训（高毛利）', level=2)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    party_strategy = [
        ('目标客户', '机关单位、国企、事业单位'),
        ('产品设计', '主题党日 VR 体验 + 党课认证'),
        ('定价', '500 元/人次（含认证证书）'),
        ('渠道', '组织部、党校、机关工委'),
        ('目标', '200 人次/年（10 万收入）')
    ]
    for i, (item, value) in enumerate(party_strategy):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = value
    
    doc.add_heading('4.2 红培研学（规模化）', level=2)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    study_strategy = [
        ('目标客户', '中小学、高校、教育机构'),
        ('产品设计', '研学课程包 + VR 体验 + 实践证书'),
        ('定价', '200 元/学生（批量优惠）'),
        ('渠道', '教育局、学校、研学机构'),
        ('目标', '1000 学生/年（20 万收入）')
    ]
    for i, (item, value) in enumerate(study_strategy):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = value
    
    doc.add_heading('4.3 国防教育（政策支持）', level=2)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    defense_strategy = [
        ('目标客户', '高校、军训基地、退役军人'),
        ('产品设计', '军事历史 VR + 战术体验 + 国防认证'),
        ('定价', '300 元/学生'),
        ('渠道', '国防办、武装部、高校武装部'),
        ('目标', '300 学生/年（9 万收入）')
    ]
    for i, (item, value) in enumerate(defense_strategy):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = value
    
    doc.add_page_break()
    
    # 政策申报支持
    doc.add_heading('五、政策申报支持', level=1)
    
    doc.add_heading('5.1 可申报项目', level=2)
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '项目'
    hdr_cells[1].text = '主管部门'
    hdr_cells[2].text = '支持额度'
    hdr_cells[3].text = '申报时间'
    
    projects = [
        ('思政教育示范基地', '省委教育工委', '20-50 万', '每年 3-5 月'),
        ('研学实践教育基地', '省教育厅', '10-30 万', '每年 6-8 月'),
        ('国防教育示范基地', '省国防办', '10-20 万', '每年 4-6 月'),
        ('文旅融合示范项目', '省文旅厅', '30-100 万', '每年 5-7 月')
    ]
    for i, (project, dept, amount, time) in enumerate(projects):
        table.rows[i+1].cells[0].text = project
        table.rows[i+1].cells[1].text = dept
        table.rows[i+1].cells[2].text = amount
        table.rows[i+1].cells[3].text = time
    
    doc.add_heading('5.2 申报服务', level=2)
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    services = [
        ('基础申报', '材料准备 + 申报指导'),
        ('深度申报', '全程代办 + 专家评审'),
        ('成功后', '基地挂牌 + 宣传推广')
    ]
    for i, (service, content) in enumerate(services):
        table.rows[i].cells[0].text = service
        table.rows[i].cells[1].text = content
    
    doc.add_page_break()
    
    # 回本周期对比
    doc.add_heading('六、回本周期对比', level=1)
    
    doc.add_heading('6.1 原模式 vs V2.0', level=2)
    table = doc.add_table(rows=5, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '模式'
    hdr_cells[1].text = '授权费'
    hdr_cells[2].text = '年收入'
    hdr_cells[3].text = '年成本'
    hdr_cells[4].text = '回本周期'
    
    models = [
        ('原模式', '50 万', '24 万（仅门票）', '20 万', '3.5 年 ❌'),
        ('轻量版', '15 万/年', '30 万', '20 万', '1.5 年 ✅'),
        ('标准版', '30 万/年', '89 万（多元）', '50 万', '1 年 ✅'),
        ('旗舰版', '26.7 万/年', '120 万', '60 万', '8 个月 ✅')
    ]
    for i, (model, fee, income, cost, payback) in enumerate(models):
        table.rows[i].cells[0].text = model
        table.rows[i].cells[1].text = fee
        table.rows[i].cells[2].text = income
        table.rows[i].cells[3].text = cost
        table.rows[i].cells[4].text = payback
    
    doc.add_heading('6.2 客户价值提升', level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '客户'
    hdr_cells[1].text = '原模式决策'
    hdr_cells[2].text = 'V2.0 决策'
    
    clients = [
        ('福建海峡出版', '❌ 放弃', '✅ 标准版（1 年回本）'),
        ('咸阳文旅集团', '❌ 放弃', '✅ 标准版 + 政策申报'),
        ('内蒙古广电', '❌ 放弃', '✅ 轻量版试水')
    ]
    for i, (client, old, new) in enumerate(clients):
        table.rows[i+1].cells[0].text = client
        table.rows[i+1].cells[1].text = old
        table.rows[i+1].cells[2].text = new
    
    doc.add_page_break()
    
    # 销售话术优化
    doc.add_heading('七、销售话术优化', level=1)
    
    doc.add_heading('7.1 原话术（已验证无效）', level=2)
    doc.add_paragraph('"50 万授权费，30 元门票分账，客流做起来就回本了"', style='Intense Quote')
    doc.add_paragraph('问题：客户测算后认为难以回本', style='List Bullet')
    
    doc.add_heading('7.2 新话术（V2.0）', level=2)
    doc.add_paragraph('"15 万起合作，多元收入结构，1 年回本"', style='Intense Quote')
    doc.add_paragraph('话术要点：', style='List Bullet')
    doc.add_paragraph('1. 低门槛：15 万/年起，降低决策风险', style='List Bullet')
    doc.add_paragraph('2. 多元收入：门票 + 培训 + 研学 + 申报，不依赖单一客流', style='List Bullet')
    doc.add_paragraph('3. 政策补贴：协助申报 20-50 万补贴，实际成本更低', style='List Bullet')
    doc.add_paragraph('4. 回本保障：标准版 1 年回本，旗舰版 8 个月回本', style='List Bullet')
    
    doc.add_heading('7.3 分客群话术', level=2)
    doc.add_paragraph('党员培训客户：', style='List Bullet')
    doc.add_paragraph('"500 元/人次认证费，200 人次就是 10 万收入，加上门票分账，1 年回本没问题"', style='Intense Quote')
    doc.add_paragraph('红培研学客户：', style='List Bullet')
    doc.add_paragraph('"200 元/学生，寒暑假各 500 学生就是 20 万，还不算门票和党员培训"', style='Intense Quote')
    doc.add_paragraph('文旅集团客户：', style='List Bullet')
    doc.add_paragraph('"协助申报文旅融合示范项目，成功就是 30-100 万补贴，授权费直接覆盖"', style='Intense Quote')
    
    doc.add_page_break()
    
    # 实施计划
    doc.add_heading('八、实施计划', level=1)
    
    doc.add_heading('8.1 时间表', level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '时间'
    hdr_cells[1].text = '任务'
    hdr_cells[2].text = '状态'
    
    timeline = [
        ('2026-04-15', '完成合作模式 V2.0 文档', '✅ 已完成'),
        ('2026-04-20', '福建海峡出版重新谈判', '⏳ 待执行'),
        ('2026-04-25', '咸阳文旅集团重新谈判', '⏳ 待执行'),
        ('2026-04-30', '内蒙古广电重新谈判', '⏳ 待执行'),
        ('2026-05-15', '签约首家标准版客户', '⏳ 待执行')
    ]
    for i, (time, task, status) in enumerate(timeline):
        table.rows[i].cells[0].text = time
        table.rows[i].cells[1].text = task
        table.rows[i].cells[2].text = status
    
    doc.add_heading('8.2 核心结论', level=2)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '维度'
    hdr_cells[1].text = '原模式'
    hdr_cells[2].text = 'V2.0 优化'
    
    conclusions = [
        ('授权费', '50 万一次性', '15-30 万/年（分层）'),
        ('收入结构', '单一门票', '门票 + 培训 + 研学 + 申报'),
        ('回本周期', '3.5 年', '8 个月 -1.5 年'),
        ('客户决策', '难', '易（低门槛 + 快回本）')
    ]
    for i, (dimension, old, new) in enumerate(conclusions):
        table.rows[i+1].cells[0].text = dimension
        table.rows[i+1].cells[1].text = old
        table.rows[i+1].cells[2].text = new
    
    # 结尾信息
    doc.add_paragraph()
    doc.add_paragraph()
    end_info = doc.add_paragraph()
    end_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_info.add_run('编制单位：[公司名称]\n').bold = True
    end_info.add_run('版本号：V2.0\n')
    end_info.add_run('更新日期：2026 年 4 月 13 日\n')
    end_info.add_run('\n本方案版权归 [公司名称] 所有，未经许可不得外传').italic = True
    
    # 保存文档
    doc.save('/home/admin/.openclaw/workspace-chief-agent/长征 - 英雄 - 城市合伙人合作模式优化方案 V2.0.docx')
    print('✅ Word 文档生成成功：长征 - 英雄 - 城市合伙人合作模式优化方案 V2.0.docx')

if __name__ == '__main__':
    create_word_document()
