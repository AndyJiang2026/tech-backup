#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成《长征·英雄》VR 大空间体验馆故事梗概 Word 文档 简化版（一页纸）
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_word_document():
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 标题
    title = doc.add_heading('《长征·英雄》VR 大空间体验馆\n故事梗概（简化版）', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 官方背书（醒目展示）
    doc.add_paragraph()
    highlights = doc.add_paragraph()
    highlights.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = highlights.add_run('🎬 国家电影局 首个重大革命题材 VR"龙标"电影\n')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(200, 0, 0)
    
    run2 = highlights.add_run('🔴 "学习强国"学习平台 领衔出品\n')
    run2.bold = True
    run2.font.size = Pt(14)
    
    doc.add_paragraph()
    doc.add_paragraph('【核心亮点】全国首个｜政治安全｜教育合规｜任务驱动交互', style='Intense Quote')
    doc.add_paragraph()
    
    # 基本信息
    doc.add_heading('一、基本信息', level=1)
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    info = [
        ('体验时长', '共 50 分钟（上篇 25 分钟 + 下篇 25 分钟）'),
        ('体验形式', 'VR 大空间沉浸式体验（任务驱动交互、多人同时参与）'),
        ('适合客群', '党员培训｜红培研学｜国防教育｜高校思政')
    ]
    for i, (item, value) in enumerate(info):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = value
    
    # 故事主线
    doc.add_heading('二、故事主线', level=1)
    doc.add_paragraph('观众化身红军战士，在红军班长"红伢子"的引导下，重走人类历史上的战争奇迹——长征。在任务驱动中"行军打仗"，在互动体验中历经血与火的洗礼，切身感受艰苦卓绝，理解"革命理想高于天"的英雄之旅。', style='Intense Quote')
    
    # 上篇
    doc.add_heading('三、上篇（25 分钟）', level=2)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '章节'
    hdr_cells[1].text = '大场景'
    hdr_cells[2].text = '情感目标'
    
    chapters = [
        ('第一章', '中央苏区', '居者有其屋、耕者有其田'),
        ('第二章', '踏上征程', '红军与亲人难舍难分'),
        ('第三章', '湘江血战', '感受牺牲与担当'),
        ('第四章', '遵义会议', '理解转折与智慧')
    ]
    for i, (chapter, scene, emotion) in enumerate(chapters):
        table.rows[i+1].cells[0].text = chapter
        table.rows[i+1].cells[1].text = scene
        table.rows[i+1].cells[2].text = emotion
    
    # 下篇
    doc.add_heading('四、下篇（25 分钟）', level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '章节'
    hdr_cells[1].text = '大场景'
    hdr_cells[2].text = '情感目标'
    
    chapters = [
        ('第五章', '四渡赤水', '领悟军事智慧'),
        ('第六章', '飞夺泸定桥', '体验勇敢与协作'),
        ('第七章', '雪山之巅', '感受坚韧与信念'),
        ('第八章', '死亡草地', '体验团结与互助'),
        ('第九章', '胜利会师', '升华信仰与传承')
    ]
    for i, (chapter, scene, emotion) in enumerate(chapters):
        table.rows[i+1].cells[0].text = chapter
        table.rows[i+1].cells[1].text = scene
        table.rows[i+1].cells[2].text = emotion
    
    # 核心互动
    doc.add_heading('五、核心互动', level=1)
    tasks = [
        '🎯 任务驱动：协同行军路线，实时生成歼敌人数',
        '🎯 营救行动：协作穿越沼泽，战友遭遇"陷落"体感',
        '🤖 动态剧情：体验后生成《歼敌人数报告》'
    ]
    for task in tasks:
        doc.add_paragraph(task, style='List Bullet')
    
    # 荣誉资质
    doc.add_heading('六、荣誉资质', level=1)
    honors = [
        '🎬 国家电影局 首个重大革命题材 VR"龙标"电影',
        '🔴 "学习强国"学习平台 领衔出品',
        '📰 学习强国平台全网传播，品牌曝光亿级'
    ]
    for honor in honors:
        doc.add_paragraph(honor, style='List Bullet')
    
    # 结尾信息（官方版权）
    doc.add_paragraph()
    end_info = doc.add_paragraph()
    end_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_info.add_run('编制单位：[公司名称]\n').bold = True
    end_info.add_run('版本号：简化版（一页纸）\n')
    end_info.add_run('更新日期：2026 年 4 月\n')
    end_info.add_run('\n本故事梗概版权归 [学习强国学习平台有限责任公司 未来新视界科技 (北京) 有限公司] 所有\n').italic = True
    end_info.add_run('《长征·英雄》内容经国家电影局龙标认证，政治安全 100% 保障').bold = True
    
    # 保存文档
    doc.save('/home/admin/.openclaw/workspace-chief-agent/长征 - 英雄 - 故事梗概简化版（一页纸）.docx')
    print('✅ Word 文档简化版生成成功：长征 - 英雄 - 故事梗概简化版（一页纸）.docx')

if __name__ == '__main__':
    create_word_document()
