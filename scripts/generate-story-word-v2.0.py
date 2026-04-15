#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成《长征·英雄》VR 大空间体验馆故事梗概 Word 文档 V2.0
基于官方版本对齐（9 章结构）
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_word_document():
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 标题
    title = doc.add_heading('《长征·英雄》VR 大空间体验馆\n故事梗概（宣推版 V2.0）', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 官方背书（醒目展示）
    doc.add_paragraph()
    highlights = doc.add_paragraph()
    highlights.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = highlights.add_run('🎬 国家电影局 首个重大革命题材 VR"龙标"电影\n')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(200, 0, 0)
    
    run2 = highlights.add_run('🔴 中央宣传部宣传舆情研究中心\n')
    run2.bold = True
    run2.font.size = Pt(14)
    
    run3 = highlights.add_run('🔴 "学习强国"学习平台 领衔出品\n')
    run3.bold = True
    run3.font.size = Pt(14)
    
    doc.add_paragraph()
    doc.add_paragraph('【核心亮点】全国首个｜政治安全｜教育合规｜任务驱动交互', style='Intense Quote')
    doc.add_paragraph()
    
    # 基本信息
    doc.add_heading('一、基本信息', level=1)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    info = [
        ('项目名称', '《长征·英雄》（上/下）'),
        ('体验时长', '共 50 分钟（上篇 25 分钟 + 下篇 25 分钟）'),
        ('体验形式', 'VR 大空间沉浸式体验（任务驱动交互、多人同时参与）'),
        ('引导角色', '红军班长"红伢子"（NPC 全程引导）'),
        ('适合客群', '党员培训｜红培研学｜国防教育｜高校思政'),
        ('推荐年龄', '12 岁 +（建议初中及以上）')
    ]
    for i, (item, value) in enumerate(info):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = value
    
    # 故事主线
    doc.add_heading('二、故事主线', level=1)
    doc.add_paragraph('核心叙事："让观众从\'观看历史\'到\'走进历史\'"', style='Intense Quote')
    doc.add_paragraph()
    doc.add_paragraph('观众化身红军战士，在红军班长"红伢子"的引导下，重走人类历史上的战争奇迹——长征。')
    doc.add_paragraph()
    doc.add_paragraph('在任务驱动中"行军打仗"，在互动体验中历经血与火的洗礼，切身感受艰苦卓绝，理解"革命理想高于天"的英雄之旅。')
    
    # 上篇（官方 4 章）
    doc.add_heading('三、上篇（25 分钟）', level=1)
    
    doc.add_heading('3.1 核心体验', level=2)
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '章节'
    hdr_cells[1].text = '大场景'
    hdr_cells[2].text = '体验内容'
    hdr_cells[3].text = '情感目标'
    
    chapters = [
        ('第一章', '中央苏区', '第五次"反围剿"失利，红军被动撤离', '居者有其屋、耕者有其田'),
        ('第二章', '踏上征程', '百姓、战友含泪相送，辎重如山', '红军与亲人难舍难分'),
        ('第三章', '湘江血战', '团队协同阻击敌人，炮火中领命任务', '感受牺牲与担当'),
        ('第四章', '遵义会议', '参与历史抉择，AI 生成动态剧情反馈', '理解转折与智慧')
    ]
    for i, (chapter, scene, content, emotion) in enumerate(chapters):
        table.rows[i+1].cells[0].text = chapter
        table.rows[i+1].cells[1].text = scene
        table.rows[i+1].cells[2].text = content
        table.rows[i+1].cells[3].text = emotion
    
    doc.add_heading('3.2 感官体验', level=2)
    senses = [
        '🔥 视觉：360°全景还原战场场景，炮弹呼啸由远及近',
        '🌬️ 触觉：手柄震动，射击敌军部队'
    ]
    for sense in senses:
        doc.add_paragraph(sense, style='List Bullet')
    
    doc.add_heading('3.3 教育价值（党员培训）', level=2)
    doc.add_paragraph('"每一代人有每一代人的长征路"', style='Intense Quote')
    values = [
        '✅ 理解初心使命：从湘江血战看共产党人的牺牲精神',
        '✅ 领悟转折智慧：遵义会议的历史意义与现实启示'
    ]
    for value in values:
        doc.add_paragraph(value, style='List Bullet')
    
    # 下篇（官方 5 章）
    doc.add_heading('四、下篇（25 分钟）', level=1)
    
    doc.add_heading('4.1 核心体验', level=2)
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '章节'
    hdr_cells[1].text = '大场景'
    hdr_cells[2].text = '体验内容'
    hdr_cells[3].text = '情感目标'
    
    chapters = [
        ('第五章', '四渡赤水', '现场聆听佯攻战术决策', '领悟军事智慧'),
        ('第六章', '飞夺泸定桥', '触觉手柄模拟汉阳造手枪射击', '体验勇敢与协作'),
        ('第七章', '雪山之巅', '感受"鸟儿飞不过，神仙也难攀"的全景壮阔', '感受坚韧与信念'),
        ('第八章', '死亡草地', '参与营救行动，团队协作穿越沼泽', '体验团结与互助'),
        ('第九章', '胜利会师', '见证历史时刻，集体宣誓仪式', '升华信仰与传承')
    ]
    for i, (chapter, scene, content, emotion) in enumerate(chapters):
        table.rows[i+1].cells[0].text = chapter
        table.rows[i+1].cells[1].text = scene
        table.rows[i+1].cells[2].text = content
        table.rows[i+1].cells[3].text = emotion
    
    doc.add_heading('4.2 感官体验', level=2)
    senses = [
        '🌊 四渡赤水：战术地图实时推演',
        '🏔️ 雪山场景：毫米级定位还原雪地行走阻力',
        '🌾 草地场景：模拟沼泽陷足，饥肠辘辘体感',
        '🎉 会师场景：交互敲鼓技术实现欢呼声环绕，集体仪式感'
    ]
    for sense in senses:
        doc.add_paragraph(sense, style='List Bullet')
    
    doc.add_heading('4.3 教育价值（红培研学/国防教育）', level=2)
    doc.add_paragraph('"长征精神永放光芒"', style='Intense Quote')
    values = [
        '✅ 锤炼意志品质：雪山草地的极限挑战',
        '✅ 培养团队意识：死亡草地的互助营救',
        '✅ 传承红色基因：胜利会师的信仰升华'
    ]
    for value in values:
        doc.add_paragraph(value, style='List Bullet')
    
    # 分客群价值主张
    doc.add_heading('五、分客群价值主张', level=1)
    
    doc.add_heading('5.1 党员培训', level=2)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    party_values = [
        ('价值点', '说明'),
        ('初心使命教育', '从长征精神看共产党人的初心'),
        ('斗争精神锤炼', '血战湘江的牺牲与担当'),
        ('转折智慧领悟', '遵义会议的历史启示'),
        ('推荐场景', '主题党日｜党校培训｜党性教育基地')
    ]
    for i, (item, value) in enumerate(party_values):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = value
    
    doc.add_heading('5.2 红培研学', level=2)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    study_values = [
        ('价值点', '说明'),
        ('沉浸式历史学习', '从"听故事"到"成为故事中人"'),
        ('情感共鸣设计', '学生反馈"对长征艰苦历程产生直观共情"'),
        ('团队协作培养', '多人协同任务，提升集体意识'),
        ('推荐场景', '青少年研学｜红色教育基地｜学校思政实践')
    ]
    for i, (item, value) in enumerate(study_values):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = value
    
    doc.add_heading('5.3 国防教育', level=2)
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    defense_values = [
        ('价值点', '说明'),
        ('军事历史教育', '真实还原长征战役场景'),
        ('战斗精神培育', '飞夺泸定桥的勇气与战术'),
        ('国防意识强化', '从长征看人民军队的初心'),
        ('推荐场景', '国防教育基地｜军训拓展｜退役军人培训')
    ]
    for i, (item, value) in enumerate(defense_values):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = value
    
    # 核心互动亮点
    doc.add_heading('六、核心互动亮点', level=1)
    
    doc.add_heading('6.1 任务驱动式体验', level=2)
    doc.add_paragraph('传统 VR：被动观看 → 《长征·英雄》：主动参与', style='Intense Quote')
    tasks = [
        '🎯 阻击敌人：协同行军路线，实时生成歼敌人数',
        '🎯 营救行动：协作穿越沼泽，同班战友遭遇"陷落"体感'
    ]
    for task in tasks:
        doc.add_paragraph(task, style='List Bullet')
    
    doc.add_heading('6.2 动态剧情', level=2)
    ai_features = [
        '🤖 个性化学习报告：体验后生成《歼敌人数报告》（团队协作、应对评分）'
    ]
    for feature in ai_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # 运营数据
    doc.add_heading('七、运营数据（标杆案例）', level=1)
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    data = [
        ('日均体验人次', '[待填写] 人'),
        ('课程满意度', '[待填写] 分'),
        ('学习兴趣提升', '+[待填写] %'),
        ('记忆留存提升', '+[待填写] %'),
        ('媒体报道', '[待填写] 家媒体'),
        ('参观接待', '[待填写] 批次')
    ]
    for i, (item, value) in enumerate(data):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = value
    
    # 教学闭环（官方时长）
    doc.add_heading('八、教学闭环设计', level=1)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '阶段'
    hdr_cells[1].text = '内容'
    hdr_cells[2].text = '时长'
    
    stages = [
        ('课前导学', '微课视频 + 任务卡（如"湘江战役歼敌人数"）', '10 分钟'),
        ('课中沉浸', '20 人分组完成"湘江突围""飞夺泸定桥"任务', '20 分钟'),
        ('课后复盘', 'AI 生成《歼敌人数报告》', '20 分钟'),
        ('社会实践', 'VR 勋章兑换长征英雄胜利徽章一枚', '持续')
    ]
    for i, (stage, content, time) in enumerate(stages):
        table.rows[i+1].cells[0].text = stage
        table.rows[i+1].cells[1].text = content
        table.rows[i+1].cells[2].text = time
    
    # 荣誉资质
    doc.add_heading('九、荣誉资质', level=1)
    honors = [
        '🎬 国家电影局 首个重大革命题材 VR"龙标"电影',
        '🔴 中央宣传部宣传舆情研究中心',
        '🔴 "学习强国"学习平台 领衔出品',
        '🏆 2026 年国央企思政工作精品项目',
        '📰 学习强国平台全网传播，品牌曝光亿级'
    ]
    for honor in honors:
        doc.add_paragraph(honor, style='List Bullet')
    
    # 现场展播话术
    doc.add_heading('十、现场展播话术', level=1)
    
    doc.add_heading('10.1 30 秒开场白', level=2)
    doc.add_paragraph('"各位党员同志/同学们，欢迎来到《长征·英雄》VR 体验馆。\n\n这是国家电影局首个重大革命题材 VR 龙标电影，由中宣部学习强国领衔出品。\n\n接下来 50 分钟，您将化身红军战士，在班长\'红伢子\'的引导下，亲历中央苏区、踏上征程、湘江血战、遵义会议、四渡赤水、飞夺泸定桥、翻雪山过草地，最终胜利会师。\n\n这不是看电影，而是成为历史中人。请准备好，我们的长征即将开始……"', style='Intense Quote')
    
    doc.add_heading('10.2 体验后总结语', level=2)
    doc.add_paragraph('"同志们/同学们，50 分钟的长征体验结束了。\n\n刚才您感受的，只是长征的万分之一。90 年前，红军战士用双脚丈量了 25000 里，用生命诠释了\'革命理想高于天\'。\n\n每一代人有每一代人的长征路。希望今天的体验，能让您带着长征精神，走好新时代的长征路。\n\n谢谢！"', style='Intense Quote')
    
    # 结尾信息（官方版权）
    doc.add_paragraph()
    doc.add_paragraph()
    end_info = doc.add_paragraph()
    end_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_info.add_run('编制单位：[公司名称]\n').bold = True
    end_info.add_run('版本号：V2.0（官方对齐版）\n')
    end_info.add_run('更新日期：2026 年 4 月\n')
    end_info.add_run('适用场景：党员培训｜红培研学｜国防教育｜高校思政\n').italic = True
    end_info.add_run('\n本故事梗概版权归 [学习强国学习平台有限责任公司 未来新视界科技 (北京) 有限公司] 所有，未经许可不得外传\n').italic = True
    end_info.add_run('《长征·英雄》内容经国家电影局龙标认证，政治安全 100% 保障').bold = True
    
    # 保存文档
    doc.save('/home/admin/.openclaw/workspace-chief-agent/长征 - 英雄 - 故事梗概展播版 V2.0.docx')
    print('✅ Word 文档 V2.0 生成成功：长征 - 英雄 - 故事梗概展播版 V2.0.docx')

if __name__ == '__main__':
    create_word_document()
