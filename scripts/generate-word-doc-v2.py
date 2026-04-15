#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成百校思政教育合作方案 Word 文档 V2.0
优化点：
1. 商业模式：设备销售 + 内容订阅
2. 课程体系：对标《中国近代史纲要》教材
3. 教学闭环：四阶教学（课前/课中/课后/实践）
4. 竞品分析：增加与曼恒等厂商对比
5. Demo 计划：新增章节
6. 排他协议：3 年独家合作条款
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

def create_word_document():
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 标题
    title = doc.add_heading('高校思政 VR 百校复制方案', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 版本信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('版本号：V2.0\n').bold = True
    info.add_run('编制日期：2026 年 4 月\n')
    info.add_run('编制单位：[公司名称]')
    
    doc.add_paragraph()  # 空行
    
    # 目录
    doc.add_heading('目录', level=1)
    sections = [
        '1. 执行摘要',
        '2. 项目背景与政策风口',
        '3. 桂林学院标杆案例',
        '4. 三档合作方案',
        '5. 课程体系与教材对标',
        '6. 四阶教学闭环设计',
        '7. 财务测算与投资回报',
        '8. 百校复制路线图',
        '9. 竞品分析与差异化',
        '10. Demo 演示方案',
        '11. 政策支持与申报指南',
        '12. 交付标准与服务承诺',
        '13. 常见问题 FAQ',
        '14. 附录'
    ]
    for section in sections:
        doc.add_paragraph(section, style='List Number')
    
    doc.add_page_break()
    
    # 第 1 章：执行摘要
    doc.add_heading('1. 执行摘要', level=1)
    
    doc.add_heading('1.1 项目定位', level=2)
    doc.add_paragraph('本项目是全国首个高校思政 VR 轻资产运营标杆项目，以桂林学院为起点，计划 3 年内复制推广至 100 所高校，打造沉浸式思政教育新生态。')
    
    doc.add_heading('1.2 核心优势', level=2)
    advantages = [
        '政策强力支持：教育部大力推进"大思政课"建设，VR+ 教育是重点支持方向',
        '商业模式验证：桂林学院项目已验证可行性，设备销售 + 内容订阅双轮驱动',
        '教材深度对标：课程体系与《中国近代史纲要（2023 版）》章节精准对应',
        '轻资产运营：设备一次性购买 + 第二年内容订阅，降低院校持续投入压力',
        '标准化交付：单校部署周期≤30 天，可快速复制'
    ]
    for adv in advantages:
        doc.add_paragraph(adv, style='List Bullet')
    
    doc.add_heading('1.3 投资规模与回报', level=2)
    table = doc.add_table(rows=4, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '方案档次'
    hdr_cells[1].text = '首年投资'
    hdr_cells[2].text = 'VR 设备数'
    hdr_cells[3].text = '订阅费/年'
    hdr_cells[4].text = '投资回收期'
    
    data = [
        ('标配', '100 万元', '20 台', '15 万元', '18-24 个月'),
        ('高配', '200 万元', '40 台', '30 万元', '20-26 个月'),
        ('顶配', '500 万元', '100 台', '75 万元', '24-30 个月')
    ]
    for i, (level, invest, devices, sub, payback) in enumerate(data):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = level
        row_cells[1].text = invest
        row_cells[2].text = devices
        row_cells[3].text = sub
        row_cells[4].text = payback
    
    doc.add_heading('1.4 三年目标', level=2)
    targets = [
        '2026 年：完成 10 所高校部署，实现收入 1000 万元',
        '2027 年：完成 50 所高校部署，实现收入 4150 万元（含订阅）',
        '2028 年：完成 100 所高校部署，实现收入 5750 万元（含订阅）'
    ]
    for target in targets:
        doc.add_paragraph(target, style='List Bullet')
    
    doc.add_page_break()
    
    # 第 2 章：项目背景与政策风口
    doc.add_heading('2. 项目背景与政策风口', level=1)
    
    doc.add_heading('2.1 思政教育现状与痛点', level=2)
    doc.add_paragraph('传统思政教育存在教学方式单一、内容抽象难懂、实践环节薄弱、效果评估困难等问题。Z 世代大学生是数字原住民，习惯沉浸式、互动式学习。')
    
    doc.add_heading('2.2 VR 技术优势', level=2)
    doc.add_paragraph('VR 技术具有沉浸式体验、互动性强、场景还原、安全可控、数据追踪等优势，可提升学习兴趣，加深理解记忆，突破时空限制。眩晕率<5%，达到行业标准。')
    
    doc.add_heading('2.3 政策支持', level=2)
    policies = [
        '《全面推进"大思政课"建设的工作方案》（教育部等十部门，2022 年）',
        '《虚拟现实与行业应用融合发展行动计划（2022-2026 年）》（工信部等五部门）',
        '《新时代学校思想政治理论课改革创新实施方案》（教育部，2020 年）',
        '《关于深化新时代学校思想政治理论课改革创新的若干意见》'
    ]
    for policy in policies:
        doc.add_paragraph(policy, style='List Bullet')
    
    doc.add_page_break()
    
    # 第 3 章：桂林学院标杆案例
    doc.add_heading('3. 桂林学院标杆案例', level=1)
    
    doc.add_heading('3.1 项目概况', level=2)
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    overview = [
        ('合作院校', '桂林学院'),
        ('签约时间', '2025 年 9 月'),
        ('正式运营', '2025 年 12 月'),
        ('投资金额', '100 万元（设备销售）'),
        ('设备规模', '20 台 VR 一体机'),
        ('场地面积', '80 平方米 VR 实训室'),
        ('服务师生', '年均 3000+ 人次')
    ]
    for i, (item, value) in enumerate(overview):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = value
    
    doc.add_heading('3.2 课程内容', level=2)
    doc.add_paragraph('包含 12 门 VR 思政精品课程，涵盖党史教育、国情教育、红色文化、价值观教育、校史教育、实践实训六大模块，与《中国近代史纲要（2023 版）》章节精准对应。')
    
    doc.add_heading('3.3 教学效果', level=2)
    effects = [
        '学习兴趣：65 分 → 92 分（+41.5%）',
        '知识掌握：72 分 → 88 分（+22.2%）',
        '课堂参与：58 分 → 94 分（+62.1%）',
        '记忆留存：45% → 78%（+73.3%）',
        '课程满意度：75 分 → 96 分（+28.0%）'
    ]
    for effect in effects:
        doc.add_paragraph(effect, style='List Bullet')
    
    doc.add_heading('3.4 商业模式', level=2)
    doc.add_paragraph('首年：100 万元设备销售（硬件 + 软件 + 内容打包）\n第二年起：15 万元/年内容订阅（课程更新 + 平台维护 + 运营服务）\n合同期限：3 年（含排他条款）')
    
    doc.add_heading('3.5 财务表现', level=2)
    doc.add_paragraph('首年毛利：约 40 万元（40%）\n订阅毛利：约 12 万元/年（80%）\n综合投资回收期：约 20 个月')
    
    doc.add_page_break()
    
    # 第 4 章：三档合作方案
    doc.add_heading('4. 三档合作方案', level=1)
    
    doc.add_heading('4.1 标配方案（100 万元）', level=2)
    doc.add_paragraph('适用对象：首次尝试 VR 思政教育的院校、普通本科院校、单校试点项目')
    doc.add_paragraph('配置：20 台 VR 一体机、12 门课程、80㎡VR 实训室、2 天师资培训、3 个月运营指导')
    doc.add_paragraph('交付周期：合同签订后 30 天')
    doc.add_paragraph('订阅服务：第二年起 15 万元/年')
    
    doc.add_heading('4.2 高配方案（200 万元）', level=2)
    doc.add_paragraph('适用对象：重点马院、省级思政示范院校、有一定 VR 应用基础的院校')
    doc.add_paragraph('配置：40 台 VR 一体机、20+2 定制课程、150㎡VR 实训中心、5 天师资培训、6 个月驻校支持')
    doc.add_paragraph('交付周期：合同签订后 45 天')
    doc.add_paragraph('订阅服务：第二年起 30 万元/年')
    
    doc.add_heading('4.3 顶配方案（500 万元）', level=2)
    doc.add_paragraph('适用对象：双一流高校、省级思政教育中心、区域示范标杆项目')
    doc.add_paragraph('配置：100 台 VR 一体机、30+5 定制课程、300㎡VR 思政中心、10 天系统培训、12 个月驻校支持')
    doc.add_paragraph('交付周期：合同签订后 60 天')
    doc.add_paragraph('订阅服务：第二年起 75 万元/年')
    
    doc.add_heading('4.4 排他性条款', level=2)
    doc.add_paragraph('合同期限：3 年\n排他范围：合作院校所在省份同类型 VR 思政产品独家合作\n续约优惠：连续订阅 3 年，第 4 年起 8 折优惠')
    
    doc.add_page_break()
    
    # 第 5 章：课程体系与教材对标
    doc.add_heading('5. 课程体系与教材对标', level=1)
    
    doc.add_heading('5.1 课程模块设计', level=2)
    modules = [
        '党史教育模块：《长征路上的抉择》《遵义会议》《开国大典》等',
        '国情教育模块：《脱贫攻坚伟大成就》《大国重器》等',
        '红色文化模块：《井冈山精神》《延安岁月》《西柏坡赶考》等',
        '价值观教育模块：《榜样的力量》《青春告白祖国》等',
        '校史教育模块：《学校发展史》（可定制）',
        '实践实训模块：《思政微课演练》《演讲比赛模拟》等'
    ]
    for module in modules:
        doc.add_paragraph(module, style='List Bullet')
    
    doc.add_heading('5.2 与《中国近代史纲要》对标', level=2)
    doc.add_paragraph('所有 VR 课程内容与《中国近代史纲要（2023 版）》教材章节精准对应，确保教学合规性和采购依据。')
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'VR 课程'
    hdr_cells[1].text = '对应教材章节'
    hdr_cells[2].text = '教材页码'
    
    mapping = [
        ('《长征路上的抉择》', '第五章 中国革命的新道路', 'P134-135'),
        ('《遵义会议》', '第五章 中国革命的新道路', 'P134-135'),
        ('《开国大典》', '第七章 为新中国而奋斗', '待确认'),
        ('《改革开放春潮》', '第八章 社会主义建设在探索中', '待确认'),
        ('《脱贫攻坚》', '第十一章 中国特色社会主义进入新时代', '待确认')
    ]
    for i in range(len(mapping)):
        course, chapter, page = mapping[i]
        table.rows[i+1].cells[0].text = course
        table.rows[i+1].cells[1].text = chapter
        table.rows[i+1].cells[2].text = page
    
    doc.add_page_break()
    
    # 第 6 章：四阶教学闭环设计
    doc.add_heading('6. 四阶教学闭环设计', level=1)
    
    doc.add_heading('6.1 教学流程', level=2)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '阶段'
    hdr_cells[1].text = '教学内容'
    hdr_cells[2].text = '技术支撑'
    
    stages = [
        ('课前导学', '微课视频 + 任务卡', '云平台推送预习资料'),
        ('课中沉浸', '20 人分组 VR 体验任务', '全感协同交互 + 数据采集'),
        ('课后复盘', 'AI 生成决策力报告', '数据可视化平台'),
        ('社会实践', 'VR 勋章兑换研学资格', '数字 - 实景联名认证')
    ]
    for i, (stage, content, tech) in enumerate(stages):
        table.rows[i+1].cells[0].text = stage
        table.rows[i+1].cells[1].text = content
        table.rows[i+1].cells[2].text = tech
    
    doc.add_heading('6.2 学习数据追踪', level=2)
    doc.add_paragraph('系统自动采集学生体验数据：体验时长、互动次数、答题正确率、团队协作效率、决策路径等，生成个性化学习报告。')
    
    doc.add_heading('6.3 师资培训认证', level=2)
    doc.add_paragraph('种子教师计划：每校培养 2-3 名认证讲师\n认证培训：5-10 天系统培训\n教师共创社区：线上互助 + 积分奖励')
    
    doc.add_page_break()
    
    # 第 7 章：财务测算与投资回报
    doc.add_heading('7. 财务测算与投资回报', level=1)
    
    doc.add_heading('7.1 收入模型', level=2)
    doc.add_paragraph('首年收入：设备销售（一次性）\n第二年起：内容订阅（设备额的 15%/年）\n增值服务：定制开发、师资培训、活动支持（按需）')
    
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '方案'
    hdr_cells[1].text = '首年销售'
    hdr_cells[2].text = '订阅费/年'
    hdr_cells[3].text = '毛利率'
    
    income_data = [
        ('标配', '100 万', '15 万', '40%/80%'),
        ('高配', '200 万', '30 万', '40%/80%'),
        ('顶配', '500 万', '75 万', '40%/80%')
    ]
    for i, row in enumerate(income_data):
        for j, cell in enumerate(row):
            table.rows[i+1].cells[j].text = cell
    
    doc.add_heading('7.2 三年收入预测', level=2)
    table = doc.add_table(rows=4, cols=5)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '年份'
    hdr_cells[1].text = '新增院校'
    hdr_cells[2].text = '设备销售'
    hdr_cells[3].text = '订阅收入'
    hdr_cells[4].text = '总收入'
    
    forecast = [
        ('2026', '10 所', '1000 万', '0', '1000 万'),
        ('2027', '40 所', '4000 万', '150 万', '4150 万'),
        ('2028', '50 所', '5000 万', '750 万', '5750 万')
    ]
    for i, row in enumerate(forecast):
        for j, cell in enumerate(row):
            table.rows[i+1].cells[j].text = cell
    
    doc.add_page_break()
    
    # 第 8 章：百校复制路线图
    doc.add_heading('8. 百校复制路线图', level=1)
    
    doc.add_heading('8.1 总体目标', level=2)
    doc.add_paragraph('3 年 100 校，打造全国高校思政 VR 教育第一品牌')
    
    doc.add_heading('8.2 阶段规划', level=2)
    stages = [
        '第一阶段（2026 Q2-Q4）：完成 10 所高校部署，收入 1000 万元',
        '第二阶段（2027 全年）：完成 50 所高校部署，收入 4150 万元',
        '第三阶段（2028 全年）：完成 100 所高校部署，收入 5750 万元'
    ]
    for stage in stages:
        doc.add_paragraph(stage, style='List Bullet')
    
    doc.add_heading('8.3 区域布局', level=2)
    regions = [
        '华东（江浙沪皖鲁）：25 所',
        '华北（京津冀晋）：20 所',
        '华南（粤桂琼）：15 所',
        '华中（鄂湘豫）：15 所',
        '西南（川渝云贵）：15 所',
        '西北（陕甘新）：10 所'
    ]
    for region in regions:
        doc.add_paragraph(region, style='List Bullet')
    
    doc.add_page_break()
    
    # 第 9 章：竞品分析与差异化
    doc.add_heading('9. 竞品分析与差异化', level=1)
    
    doc.add_heading('9.1 主要竞品', level=2)
    doc.add_paragraph('曼恒数字：LBE 大空间 VR 方案，高端路线，技术壁垒强\n其他厂商：单机 VR 方案，内容单薄，缺乏教学闭环')
    
    doc.add_heading('9.2 竞品对比', level=2)
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '维度'
    hdr_cells[1].text = '曼恒（高端）'
    hdr_cells[2].text = '其他厂商'
    hdr_cells[3].text = '我们（轻资产）'
    
    comparison = [
        ('技术类型', 'LBE 大空间（行走式）', '单机 VR', 'VR 一体机（定点式）'),
        ('场地要求', '≥200㎡', '50-100㎡', '80-300㎡'),
        ('投资门槛', '300 万+', '50-100 万', '100-500 万'),
        ('部署周期', '60-90 天', '15-30 天', '30-45 天'),
        ('运维复杂度', '专业团队', '简单', '1 名管理员即可'),
        ('核心优势', '技术壁垒', '低价', '快速复制 + 教材对标')
    ]
    for i, row in enumerate(comparison):
        for j, cell in enumerate(row):
            table.rows[i+1].cells[j].text = cell
    
    doc.add_heading('9.3 差异化定位', level=2)
    doc.add_paragraph('我们不走高端大空间路线，而是聚焦轻资产快速复制：\n- 目标客户：普通本科/职业院校（非双一流）\n- 投资门槛：100 万起（曼恒 1/3）\n- 部署周期：30 天（曼恒 1/2）\n- 教材对标：深度绑定《纲要》，竞品难以复制')
    
    doc.add_page_break()
    
    # 第 10 章：Demo 演示方案
    doc.add_heading('10. Demo 演示方案', level=1)
    
    doc.add_heading('10.1 Demo 战略价值', level=2)
    doc.add_paragraph('3-5 分钟 Demo 是高校销售标配，用于：\n- 首次拜访：直观展示 VR 效果，降低理解成本\n- 招标比稿：差异化呈现，脱颖而出\n- 预算审批：实证视频，增强决策信心\n- 教师培训：展示操作流程，降低顾虑')
    
    doc.add_heading('10.2 Demo 内容结构', level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '时间段'
    hdr_cells[1].text = '内容'
    hdr_cells[2].text = '目标'
    
    demo_flow = [
        ('0-30 秒', '开场：政策 + 痛点', '引起共鸣'),
        ('30-90 秒', '产品亮相：桂林学院实景', '建立信任'),
        ('90-180 秒', '核心体验：VR 课程片段', '展示效果'),
        ('180-240 秒', '学生反馈 + 数据', '验证效果'),
        ('240-300 秒', '合作方案 + 联系方式', '促成行动')
    ]
    for i, (time, content, goal) in enumerate(demo_flow):
        table.rows[i+1].cells[0].text = time
        table.rows[i+1].cells[1].text = content
        table.rows[i+1].cells[2].text = goal
    
    doc.add_heading('10.3 Demo 制作计划', level=2)
    doc.add_paragraph('快速版（1-2 周）：手机拍摄 + 简单剪辑，预算 1-2 万\n专业版（3-4 周）：专业团队拍摄，预算 5-8 万\n旗舰版（6-8 周）：多校取景 + 纪录片风格，预算 15-20 万')
    
    doc.add_page_break()
    
    # 第 11 章：政策支持与申报指南
    doc.add_heading('11. 政策支持与申报指南', level=1)
    
    doc.add_heading('11.1 可申报项目', level=2)
    projects = [
        '国家级：大思政课示范项目（50-100 万）、虚拟仿真实验教学项目（30-80 万）',
        '省级：思政工作精品项目（20-50 万）、智慧马院建设（30-100 万）、虚拟仿真基地（50-200 万）'
    ]
    for project in projects:
        doc.add_paragraph(project, style='List Bullet')
    
    doc.add_heading('11.2 申报策略', level=2)
    doc.add_paragraph('最佳时机：项目落地后 3-6 个月（有运营数据支撑）\n申报成功率提升技巧：突出创新性、数据支撑、示范效应、校政企协同')
    
    doc.add_page_break()
    
    # 第 12 章：交付标准与服务承诺
    doc.add_heading('12. 交付标准与服务承诺', level=1)
    
    doc.add_heading('12.1 交付标准', level=2)
    doc.add_paragraph('硬件交付：100% 开机正常，网络延迟≤20ms，眩晕率<5%\n软件交付：所有功能正常运行，课程内容完整可用\n场地交付：符合设计方案，安全设施完备')
    
    doc.add_heading('12.2 服务承诺', level=2)
    doc.add_paragraph('培训服务：基础培训 1 天、教学培训 1-4 天、认证培训 5-10 天\n运维服务：电话咨询即时响应、远程支持≤30 分钟（80% 故障远程解决）、现场服务≤4 小时\n内容更新：月度常规更新、季度大版本、重大节日专题更新')
    
    doc.add_heading('12.3 质量保证', level=2)
    doc.add_paragraph('硬件设备质保 2 年、软件系统质保 1 年、装修工程质保 1 年\n系统可用性≥99%、故障响应≤30 分钟、故障解决≤24 小时\n数据安全：国密级加密传输，私有云存储')
    
    doc.add_page_break()
    
    # 第 13 章：常见问题 FAQ
    doc.add_heading('13. 常见问题 FAQ', level=1)
    
    faqs = [
        ('Q1：VR 设备会不会让学生头晕？', 'A：选用 PICO 4 Enterprise 高分辨率高刷新率，眩晕率<5%，95% 以上学生无眩晕感。'),
        ('Q2：网络条件要求高吗？', 'A：本地内容无需网络，建议带宽≥100Mbps，提供离线内容包。'),
        ('Q3：教师没有 VR 经验怎么办？', 'A：提供系统化培训，2 天培训即可独立开展 VR 教学。种子教师计划每校培养 2-3 名认证讲师。'),
        ('Q4：投资是一次性还是分期？', 'A：支持一次性付款和分期付款，分期通常为 30% 首付 +40% 交付 +30% 验收。'),
        ('Q5：能否申请政府专项资金？', 'A：可以，提供全套申报材料支持，可申报 20-200 万元不等专项资金。'),
        ('Q6：内容如何更新？', 'A：月度常规更新，季度大版本，重大节日专题更新。第二年起订阅服务包含所有更新。'),
        ('Q7：合同期限多长？', 'A：标准合同期 3 年，含排他条款。期满可续签，续约享受优惠价格。')
    ]
    for question, answer in faqs:
        p = doc.add_paragraph()
        p.add_run(question).bold = True
        doc.add_paragraph(answer)
    
    doc.add_page_break()
    
    # 第 14 章：附录
    doc.add_heading('14. 附录', level=1)
    
    doc.add_heading('14.1 合同模板（摘要）', level=2)
    doc.add_paragraph('合作协议包含：合作内容、投资金额、交付周期、服务期限、双方权利义务、排他条款、违约责任、争议解决等。')
    
    doc.add_heading('14.2 设备清单（标配方案）', level=2)
    doc.add_paragraph('VR 一体机 20 台（PICO 4 Enterprise）、管理主机 1 台、显示设备 2 台、充电存储柜 1 台、网络设备 1 套等，硬件设备成本合计 20.9 万元。')
    
    doc.add_heading('14.3 课程内容目录', level=2)
    courses = [
        '党史教育：《长征路上的抉择》《遵义会议》《开国大典》《改革开放春潮》等',
        '国情教育：《脱贫攻坚伟大成就》《大国重器》《抗疫精神》等',
        '红色文化：《井冈山精神》《延安岁月》《西柏坡赶考》《红船启航》等',
        '价值观教育：《榜样的力量》《青春告白祖国》《职业道德》等',
        '校史教育：《学校发展史》（可定制）',
        '实践实训：《思政微课演练》《演讲比赛模拟》《情景剧排演》等'
    ]
    for course in courses:
        doc.add_paragraph(course, style='List Bullet')
    
    # 结尾信息
    doc.add_paragraph()
    end_info = doc.add_paragraph()
    end_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_info.add_run('编制单位：[公司名称]\n').bold = True
    end_info.add_run('编制日期：2026 年 4 月\n')
    end_info.add_run('本方案版权归 [公司名称] 所有，未经许可不得外传').italic = True
    
    # 保存文档
    doc.save('/home/admin/.openclaw/workspace-chief-agent/百校复制方案 - 完整版 V2.0.docx')
    print('✅ Word 文档 V2.0 生成成功：百校复制方案 - 完整版 V2.0.docx')

if __name__ == '__main__':
    create_word_document()
