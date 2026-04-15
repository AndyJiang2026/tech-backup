#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成百校思政教育合作方案 Word 文档 V3.0
核心：学习强国 + 龙标双重背书贯穿全文
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_word_document():
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 标题
    title = doc.add_heading('高校思政 VR 百校复制方案', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 版本信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('版本号：V3.0（终极版）\n').bold = True
    info.add_run('编制日期：2026 年 4 月\n')
    info.add_run('编制单位：[公司名称]')
    
    doc.add_paragraph()
    
    # 核心背书（醒目展示）
    highlights = doc.add_paragraph()
    highlights.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = highlights.add_run('🔴 中宣部学习强国领衔出品 | 🎬 国家电影局首个 VR 龙标认证\n')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(200, 0, 0)
    
    doc.add_paragraph()
    doc.add_paragraph('【核心优势】学习强国 + 龙标双重背书，政治安全 100% 保障，竞品 100% 无法复制').italic = True
    doc.add_paragraph()
    
    # 目录
    doc.add_heading('目录', level=1)
    sections = [
        '1. 执行摘要',
        '2. 独家资源优势（学习强国 + 龙标）',
        '3. 项目背景与政策风口',
        '4. 桂林学院标杆案例',
        '5. 三档合作方案',
        '6. 课程体系与教材对标',
        '7. 四阶教学闭环设计',
        '8. AIGC 技术服务',
        '9. 财务测算与投资回报',
        '10. 百校复制路线图',
        '11. 竞品分析与差异化',
        '12. Demo 演示方案',
        '13. 交付标准与服务承诺',
        '14. 常见问题 FAQ',
        '15. 附录'
    ]
    for section in sections:
        doc.add_paragraph(section, style='List Number')
    
    doc.add_page_break()
    
    # 第 1 章：执行摘要
    doc.add_heading('1. 执行摘要', level=1)
    
    doc.add_heading('1.1 项目定位', level=2)
    doc.add_paragraph('本项目是全国首个高校思政 VR 轻资产运营标杆项目，核心内容获中宣部学习强国领衔出品、国家电影局首个 VR 龙标认证。以桂林学院为起点，计划 3 年内复制推广至 100 所高校，打造沉浸式思政教育新生态。')
    
    doc.add_heading('1.2 核心优势（竞品无法复制）', level=2)
    advantages = [
        '🔴 独家背书：中宣部学习强国领衔出品 + 国家电影局首个 VR 龙标认证，政治安全 100% 保障',
        '🎬 内容审查：通过国家最高级别内容审查，高校采购零风险，审批无障碍',
        '📚 教材对标：课程体系与《中国近代史纲要（2023 版）》精准对应',
        '🤖 AIGC 赋能：校本课程 AI 生成、AI 教学助手、学习报告自动生成',
        '📊 商业验证：桂林学院项目已验证，设备销售 + 订阅+AIGC 三轮驱动'
    ]
    for adv in advantages:
        doc.add_paragraph(adv, style='List Bullet')
    
    doc.add_heading('1.3 投资规模与回报', level=2)
    table = doc.add_table(rows=4, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '方案'
    hdr_cells[1].text = '首年投资'
    hdr_cells[2].text = 'VR 设备'
    hdr_cells[3].text = '订阅/年'
    hdr_cells[4].text = 'AIGC/年'
    hdr_cells[5].text = '回收期'
    
    data = [
        ('标配', '100 万', '20 台', '15 万', '5-10 万', '18-24 月'),
        ('高配', '200 万', '40 台', '30 万', '10-20 万', '20-26 月'),
        ('顶配', '500 万', '100 台', '75 万', '20-40 万', '24-30 月')
    ]
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.rows[i+1].cells[j].text = cell
    
    doc.add_heading('1.4 三年目标', level=2)
    targets = [
        '2026 年：完成 10 所高校部署，实现收入 1000 万元',
        '2027 年：完成 50 所高校部署，实现收入 4250 万元（含订阅+AIGC）',
        '2028 年：完成 100 所高校部署，实现收入 6200 万元（含订阅+AIGC）'
    ]
    for target in targets:
        doc.add_paragraph(target, style='List Bullet')
    
    doc.add_page_break()
    
    # 第 2 章：独家资源优势
    doc.add_heading('2. 独家资源优势（学习强国 + 龙标）', level=1)
    
    doc.add_heading('2.1 中宣部学习强国领衔出品', level=2)
    doc.add_paragraph('🔴 政治背书：中宣部主管的国家级学习平台，政治安全性 100% 保障，高校采购无政治风险')
    doc.add_paragraph('🔴 品牌效应："学习强国"品牌认知度 95%+，高校认可度极高，无需额外教育成本')
    doc.add_paragraph('🔴 采购合规：高校采购零风险，财务/资产处审批无障碍，无顾虑')
    doc.add_paragraph('🔴 传播价值：学习强国平台全网传播，品牌曝光亿级，项目知名度快速提升')
    doc.add_paragraph('🔴 排他性：学习强国出品内容具有唯一性，竞品无法获得同等背书')
    
    doc.add_heading('2.2 国家电影局首个 VR"龙标"电影', level=2)
    doc.add_paragraph('🎬 行业首创：全国首个 VR 内容获龙标，历史意义，里程碑事件，媒体关注度高')
    doc.add_paragraph('🎬 审查通过：国家最高级别内容审查，政治安全 100% 保障，内容合规性无争议')
    doc.add_paragraph('🎬 稀缺性：重大革命题材 VR 唯一龙标，排他性，竞品 100% 无法复制')
    doc.add_paragraph('🎬 教育认可：龙标=教育内容合规认证，高校采购依据充分，教学应用无风险')
    doc.add_paragraph('🎬 品牌溢价：龙标内容具有收藏价值和学术价值，提升项目整体档次')
    
    doc.add_heading('2.3 竞品壁垒分析', level=2)
    doc.add_paragraph('曼恒数字：技术强，但无学习强国 + 龙标背书，政治安全性存疑，高校采购需谨慎')
    doc.add_paragraph('其他厂商：无官方背书，内容审查未通过，高校采购风险高，可能被叫停')
    doc.add_paragraph('我们的优势：学习强国 + 龙标双重背书，100% 排他性，竞品无法复制，高校采购放心')
    
    doc.add_heading('2.4 背书使用规范', level=2)
    doc.add_paragraph('学习强国 Logo 使用：需获得官方书面授权，按规范使用')
    doc.add_paragraph('龙标标识使用：按国家电影局规范使用，不得变形、变色')
    doc.add_paragraph('宣传话术：统一使用"中宣部学习强国领衔出品"、"国家电影局首个 VR 龙标认证"')
    
    doc.add_page_break()
    
    # 第 3 章：项目背景与政策风口
    doc.add_heading('3. 项目背景与政策风口', level=1)
    doc.add_heading('3.1 思政教育现状与痛点', level=2)
    doc.add_paragraph('传统思政教育存在教学方式单一、内容抽象难懂、实践环节薄弱、效果评估困难等问题。')
    
    doc.add_heading('3.2 政策支持', level=2)
    policies = [
        '《全面推进"大思政课"建设的工作方案》（教育部等十部门，2022 年）',
        '《虚拟现实与行业应用融合发展行动计划（2022-2026 年）》（工信部等五部门）',
        '《关于深化新时代学校思想政治理论课改革创新的若干意见》'
    ]
    for policy in policies:
        doc.add_paragraph(policy, style='List Bullet')
    
    doc.add_page_break()
    
    # 第 4 章：桂林学院标杆案例
    doc.add_heading('4. 桂林学院标杆案例', level=1)
    
    doc.add_heading('4.1 项目概况', level=2)
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    overview = [
        ('合作院校', '桂林学院'),
        ('签约时间', '2025 年 9 月'),
        ('正式运营', '2025 年 12 月'),
        ('投资金额', '100 万元（设备销售）'),
        ('设备规模', '20 台 VR 一体机'),
        ('场地面积', '80 平方米 VR 实训室'),
        ('服务师生', '年均 3000+ 人次')
    ]
    for i, (item, value) in enumerate(overview):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = value
    
    doc.add_heading('4.2 商业模式', level=2)
    doc.add_paragraph('首年：100 万元设备销售（硬件 + 软件 + 内容打包）\n第二年起：15 万元/年内容订阅 + 5-10 万元/年 AIGC 技术服务\n合同期限：3 年（含排他条款）')
    
    doc.add_heading('4.3 教学效果', level=2)
    effects = [
        '学习兴趣：65 分 → 92 分（+41.5%）',
        '知识掌握：72 分 → 88 分（+22.2%）',
        '课堂参与：58 分 → 94 分（+62.1%）',
        '记忆留存：45% → 78%（+73.3%）',
        '课程满意度：75 分 → 96 分（+28.0%）'
    ]
    for effect in effects:
        doc.add_paragraph(effect, style='List Bullet')
    
    doc.add_page_break()
    
    # 第 5 章：三档合作方案
    doc.add_heading('5. 三档合作方案', level=1)
    
    doc.add_heading('5.1 标配方案（100 万元）', level=2)
    doc.add_paragraph('适用对象：首次尝试 VR 思政教育的院校、普通本科院校')
    doc.add_paragraph('配置：20 台 VR 一体机、12 门课程（学习强国 + 龙标认证内容）、80㎡VR 实训室、2 天师资培训')
    doc.add_paragraph('交付周期：合同签订后 30 天')
    doc.add_paragraph('订阅服务：第二年起 15 万元/年（内容更新 + 运维）')
    doc.add_paragraph('AIGC 服务：第二年起 5-10 万元/年（校本课程 AI 生成等）')
    
    doc.add_heading('5.2 高配方案（200 万元）', level=2)
    doc.add_paragraph('适用对象：重点马院、省级思政示范院校')
    doc.add_paragraph('配置：40 台 VR 一体机、20+2 定制课程、150㎡VR 实训中心')
    doc.add_paragraph('交付周期：合同签订后 45 天')
    doc.add_paragraph('订阅服务：第二年起 30 万元/年')
    doc.add_paragraph('AIGC 服务：第二年起 10-20 万元/年')
    
    doc.add_heading('5.3 顶配方案（500 万元）', level=2)
    doc.add_paragraph('适用对象：双一流高校、省级思政教育中心')
    doc.add_paragraph('配置：100 台 VR 一体机、30+5 定制课程、300㎡VR 思政中心')
    doc.add_paragraph('交付周期：合同签订后 60 天')
    doc.add_paragraph('订阅服务：第二年起 75 万元/年')
    doc.add_paragraph('AIGC 服务：第二年起 20-40 万元/年')
    
    doc.add_heading('5.4 排他性条款', level=2)
    doc.add_paragraph('合同期限：3 年\n排他范围：合作院校所在省份同类型 VR 思政产品独家合作\n续约优惠：连续订阅 3 年，第 4 年起 8 折优惠')
    
    doc.add_page_break()
    
    # 第 6 章：课程体系与教材对标
    doc.add_heading('6. 课程体系与教材对标', level=1)
    
    doc.add_heading('6.1 课程模块设计', level=2)
    modules = [
        '党史教育模块：《长征路上的抉择》《遵义会议》《开国大典》等（学习强国 + 龙标认证）',
        '国情教育模块：《脱贫攻坚伟大成就》《大国重器》等',
        '红色文化模块：《井冈山精神》《延安岁月》《西柏坡赶考》等',
        '价值观教育模块：《榜样的力量》《青春告白祖国》等',
        '校史教育模块：《学校发展史》（可定制）',
        '实践实训模块：《思政微课演练》《演讲比赛模拟》等'
    ]
    for module in modules:
        doc.add_paragraph(module, style='List Bullet')
    
    doc.add_heading('6.2 与《中国近代史纲要》对标', level=2)
    doc.add_paragraph('所有 VR 课程内容与《中国近代史纲要（2023 版）》教材章节精准对应，确保教学合规性和采购依据。')
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'VR 课程'
    hdr_cells[1].text = '对应教材章节'
    hdr_cells[2].text = '教材页码'
    
    mapping = [
        ('《长征路上的抉择》', '第五章 中国革命的新道路', 'P134-135'),
        ('《遵义会议》', '第五章 中国革命的新道路', 'P134-135'),
        ('《开国大典》', '第七章 为新中国而奋斗', '待确认'),
        ('《改革开放春潮》', '第八章 社会主义建设在探索中', '待确认'),
        ('《脱贫攻坚》', '第十一章 中国特色社会主义进入新时代', '待确认')
    ]
    for i in range(len(mapping)):
        course, chapter, page = mapping[i]
        table.rows[i+1].cells[0].text = course
        table.rows[i+1].cells[1].text = chapter
        table.rows[i+1].cells[2].text = page
    
    doc.add_page_break()
    
    # 第 7 章：四阶教学闭环设计
    doc.add_heading('7. 四阶教学闭环设计', level=1)
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '阶段'
    hdr_cells[1].text = '教学内容'
    hdr_cells[2].text = '技术支撑'
    
    stages = [
        ('课前导学', '微课视频 + 任务卡', '云平台推送预习资料'),
        ('课中沉浸', '20 人分组 VR 体验任务', '全感协同交互 + 数据采集'),
        ('课后复盘', 'AI 生成决策力报告', '数据可视化平台'),
        ('社会实践', 'VR 勋章兑换研学资格', '数字 - 实景联名认证')
    ]
    for i, (stage, content, tech) in enumerate(stages):
        table.rows[i+1].cells[0].text = stage
        table.rows[i+1].cells[1].text = content
        table.rows[i+1].cells[2].text = tech
    
    doc.add_page_break()
    
    # 第 8 章：AIGC 技术服务
    doc.add_heading('8. AIGC 技术服务', level=1)
    
    doc.add_heading('8.1 服务清单', level=2)
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '服务项'
    hdr_cells[1].text = '说明'
    hdr_cells[2].text = '定价'
    hdr_cells[3].text = '毛利'
    
    aigc_services = [
        ('校本课程 AI 生成', '基于学校历史/地方红色资源生成定制课程', '5-10 万/门', '85%+'),
        ('本地红色资源数字化', '扫描/建模本地红色遗址、纪念馆', '10-30 万/项目', '80%+'),
        ('AI 教学助手', '自动生成教案、课件、测试题', '3-5 万/年', '90%+'),
        ('学习报告 AI 生成', '自动分析学生数据，生成个性化报告', '包含在订阅', '95%+'),
        ('AI 虚拟讲解员', '虚拟历史人物对话系统', '5-8 万/年', '85%+'),
        ('内容智能推荐', '个性化课程路径推荐', '包含在订阅', '95%+')
    ]
    for i, (service, desc, price, margin) in enumerate(aigc_services):
        table.rows[i+1].cells[0].text = service
        table.rows[i+1].cells[1].text = desc
        table.rows[i+1].cells[2].text = price
        table.rows[i+1].cells[3].text = margin
    
    doc.add_heading('8.2 AIGC 渗透率预测', level=2)
    doc.add_paragraph('第二年：50% 客户采购 AIGC 增值服务\n第三年：70% 客户采购 AIGC 增值服务\n续约客户：90% 以上续订 AIGC 服务')
    
    doc.add_page_break()
    
    # 第 9 章：财务测算与投资回报
    doc.add_heading('9. 财务测算与投资回报', level=1)
    
    doc.add_heading('9.1 三年收入预测', level=2)
    table = doc.add_table(rows=4, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '年份'
    hdr_cells[1].text = '新增院校'
    hdr_cells[2].text = '设备销售'
    hdr_cells[3].text = '订阅收入'
    hdr_cells[4].text = 'AIGC 收入'
    hdr_cells[5].text = '总收入'
    
    forecast = [
        ('2026', '10 所', '1000 万', '0', '0', '1000 万'),
        ('2027', '40 所', '4000 万', '150 万', '100 万', '4250 万'),
        ('2028', '50 所', '5000 万', '750 万', '450 万', '6200 万')
    ]
    for i, row in enumerate(forecast):
        for j, cell in enumerate(row):
            table.rows[i+1].cells[j].text = cell
    
    doc.add_page_break()
    
    # 第 10 章：百校复制路线图
    doc.add_heading('10. 百校复制路线图', level=1)
    doc.add_heading('10.1 总体目标', level=2)
    doc.add_paragraph('3 年 100 校，打造全国高校思政 VR 教育第一品牌')
    
    doc.add_heading('10.2 阶段规划', level=2)
    stages = [
        '第一阶段（2026 Q2-Q4）：完成 10 所高校部署，收入 1000 万元',
        '第二阶段（2027 全年）：完成 50 所高校部署，收入 4250 万元',
        '第三阶段（2028 全年）：完成 100 所高校部署，收入 6200 万元'
    ]
    for stage in stages:
        doc.add_paragraph(stage, style='List Bullet')
    
    doc.add_page_break()
    
    # 第 11 章：竞品分析与差异化
    doc.add_heading('11. 竞品分析与差异化', level=1)
    
    doc.add_heading('11.1 竞品对比', level=2)
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '维度'
    hdr_cells[1].text = '曼恒（高端）'
    hdr_cells[2].text = '其他厂商'
    hdr_cells[3].text = '我们（轻资产）'
    
    comparison = [
        ('官方背书', '无', '无', '✅ 学习强国 + 龙标双重背书'),
        ('内容审查', '未通过龙标', '未通过', '✅ 国家电影局龙标认证'),
        ('政治安全', '存疑', '高风险', '✅ 100% 安全保障'),
        ('教材对标', '一般', '无', '✅ 深度绑定《纲要》'),
        ('AIGC 能力', '基础 AI 对话', '无', '✅ 校本课程 AI 生成'),
        ('投资门槛', '300 万+', '50-100 万', '100-500 万'),
        ('部署周期', '60-90 天', '15-30 天', '30-45 天')
    ]
    for i, row in enumerate(comparison):
        for j, cell in enumerate(row):
            table.rows[i+1].cells[j].text = cell
    
    doc.add_heading('11.2 差异化定位', level=2)
    doc.add_paragraph('我们不走高端大空间路线，而是聚焦轻资产快速复制 + 独家背书：\n- 目标客户：普通本科/职业院校（非双一流）\n- 投资门槛：100 万起（曼恒 1/3）\n- 独家优势：学习强国 + 龙标，竞品 100% 无法复制')
    
    doc.add_page_break()
    
    # 第 12 章：Demo 演示方案
    doc.add_heading('12. Demo 演示方案', level=1)
    
    doc.add_heading('12.1 Demo 战略价值', level=2)
    doc.add_paragraph('3-5 分钟 Demo 是高校销售标配，用于：\n- 首次拜访：直观展示 VR 效果，降低理解成本\n- 招标比稿：差异化呈现，脱颖而出\n- 预算审批：实证视频，增强决策信心\n- 教师培训：展示操作流程，降低顾虑')
    
    doc.add_heading('12.2 Demo 内容结构', level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '时间段'
    hdr_cells[1].text = '内容'
    hdr_cells[2].text = '目标'
    
    demo_flow = [
        ('0-30 秒', '开场：政策 + 痛点', '引起共鸣'),
        ('30-90 秒', '产品亮相：桂林学院实景', '建立信任'),
        ('90-180 秒', '核心体验：VR 课程片段', '展示效果'),
        ('180-240 秒', '学生反馈 + 数据', '验证效果'),
        ('240-300 秒', '合作方案 + 联系方式', '促成行动')
    ]
    for i, (time, content, goal) in enumerate(demo_flow):
        table.rows[i+1].cells[0].text = time
        table.rows[i+1].cells[1].text = content
        table.rows[i+1].cells[2].text = goal
    
    doc.add_heading('12.3 Demo 制作计划', level=2)
    doc.add_paragraph('快速版（1-2 周）：手机拍摄 + 简单剪辑，预算 1-2 万\n专业版（3-4 周）：专业团队拍摄，预算 5-8 万\n旗舰版（6-8 周）：多校取景 + 纪录片风格，预算 15-20 万')
    
    doc.add_page_break()
    
    # 第 13 章：交付标准与服务承诺
    doc.add_heading('13. 交付标准与服务承诺', level=1)
    
    doc.add_heading('13.1 交付标准', level=2)
    doc.add_paragraph('硬件交付：100% 开机正常，网络延迟≤20ms，眩晕率<5%\n软件交付：所有功能正常运行，课程内容完整可用（学习强国 + 龙标认证内容）\n场地交付：符合设计方案，安全设施完备')
    
    doc.add_heading('13.2 服务承诺', level=2)
    doc.add_paragraph('培训服务：基础培训 1 天、教学培训 1-4 天、认证培训 5-10 天\n运维服务：电话咨询即时响应、远程支持≤30 分钟（80% 故障远程解决）、现场服务≤4 小时\n内容更新：月度常规更新、季度大版本、重大节日专题更新')
    
    doc.add_heading('13.3 质量保证', level=2)
    doc.add_paragraph('硬件设备质保 2 年、软件系统质保 1 年、装修工程质保 1 年\n系统可用性≥99%、故障响应≤30 分钟、故障解决≤24 小时\n数据安全：国密级加密传输，私有云存储')
    
    doc.add_page_break()
    
    # 第 14 章：常见问题 FAQ
    doc.add_heading('14. 常见问题 FAQ', level=1)
    
    faqs = [
        ('Q1：内容是否政治安全？', 'A：中宣部学习强国领衔出品 + 国家电影局龙标认证，双重背书，政治安全性 100% 保障。'),
        ('Q2：能否申请政府专项资金？', 'A：可以，学习强国 + 龙标背书，申报成功率更高，可申报 20-200 万元专项资金。'),
        ('Q3：AIGC 服务是否必须采购？', 'A：AIGC 为可选增值服务，基础订阅已包含课程更新和运维服务。'),
        ('Q4：龙标编号是多少？', 'A：[待补充具体龙标编号]'),
        ('Q5：学习强国如何展示？', 'A：[待补充学习强国平台链接或页面截图]'),
        ('Q6：竞品价格更低怎么办？', 'A：强调学习强国 + 龙标背书，政治安全无价，高校采购零风险是核心价值。'),
        ('Q7：合同期限多长？', 'A：标准合同期 3 年，含排他条款。期满可续签，续约享受优惠价格。')
    ]
    for question, answer in faqs:
        p = doc.add_paragraph()
        p.add_run(question).bold = True
        doc.add_paragraph(answer)
    
    doc.add_page_break()
    
    # 第 15 章：附录
    doc.add_heading('15. 附录', level=1)
    
    doc.add_heading('15.1 学习强国 + 龙标证明材料', level=2)
    doc.add_paragraph('- 学习强国平台页面截图\n- 龙标证书扫描件\n- 官方授权书（如适用）\n- 宣传话术规范')
    
    doc.add_heading('15.2 合同模板（摘要）', level=2)
    doc.add_paragraph('合作协议包含：合作内容、投资金额、交付周期、服务期限、排他条款、违约责任、争议解决等。')
    
    doc.add_heading('15.3 设备清单（标配方案）', level=2)
    doc.add_paragraph('VR 一体机 20 台（PICO 4 Enterprise）、管理主机 1 台、显示设备 2 台、充电存储柜 1 台、网络设备 1 套等。')
    
    doc.add_heading('15.4 课程内容目录', level=2)
    courses = [
        '党史教育：《长征路上的抉择》《遵义会议》《开国大典》《改革开放春潮》等（学习强国 + 龙标认证）',
        '国情教育：《脱贫攻坚伟大成就》《大国重器》《抗疫精神》等',
        '红色文化：《井冈山精神》《延安岁月》《西柏坡赶考》《红船启航》等',
        '价值观教育：《榜样的力量》《青春告白祖国》《职业道德》等',
        '校史教育：《学校发展史》（可定制）',
        '实践实训：《思政微课演练》《演讲比赛模拟》《情景剧排演》等'
    ]
    for course in courses:
        doc.add_paragraph(course, style='List Bullet')
    
    # 结尾信息
    doc.add_paragraph()
    end_info = doc.add_paragraph()
    end_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_info.add_run('编制单位：[公司名称]\n').bold = True
    end_info.add_run('编制日期：2026 年 4 月\n')
    end_info.add_run('🔴 中宣部学习强国领衔出品 | 🎬 国家电影局首个 VR 龙标认证\n').bold = True
    end_info.add_run('本方案版权归 [公司名称] 所有，未经许可不得外传').italic = True
    
    # 保存文档
    doc.save('/home/admin/.openclaw/workspace-chief-agent/百校复制方案 - 完整版 V3.0.docx')
    print('✅ Word 文档 V3.0 生成成功：百校复制方案 - 完整版 V3.0.docx')

if __name__ == '__main__':
    create_word_document()
