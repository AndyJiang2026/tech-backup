#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成百校思政教育合作方案 Word 文档 V3.0
核心优化：
1. 独家资源：学习强国 + 龙标双重背书
2. 商业模式：设备销售 + 内容订阅+AIGC 技术服务
3. 教材对标：《中国近代史纲要（2023 版）》
4. 教学闭环：四阶教学
5. 竞品分析：强化官方背书对比
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_word_document():
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 标题
    title = doc.add_heading('高校思政 VR 百校复制方案', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 版本信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('版本号：V3.0（终极版）\n').bold = True
    info.add_run('编制日期：2026 年 4 月\n')
    info.add_run('编制单位：[公司名称]')
    
    doc.add_paragraph()
    
    # 核心亮点摘要
    highlights = doc.add_paragraph()
    highlights.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = highlights.add_run('🔴 中宣部学习强国领衔出品 | 🎬 国家电影局首个 VR 龙标认证\n')
    run.bold = True
    run.font.color.rgb = RGBColor(200, 0, 0)
    
    doc.add_paragraph()
    
    # 目录
    doc.add_heading('目录', level=1)
    sections = [
        '1. 执行摘要',
        '2. 独家资源优势（学习强国 + 龙标）',
        '3. 项目背景与政策风口',
        '4. 桂林学院标杆案例',
        '5. 三档合作方案',
        '6. 课程体系与教材对标',
        '7. 四阶教学闭环设计',
        '8. AIGC 技术服务',
        '9. 财务测算与投资回报',
        '10. 百校复制路线图',
        '11. 竞品分析与差异化',
        '12. Demo 演示方案',
        '13. 政策支持与申报指南',
        '14. 交付标准与服务承诺',
        '15. 常见问题 FAQ',
        '16. 附录'
    ]
    for section in sections:
        doc.add_paragraph(section, style='List Number')
    
    doc.add_page_break()
    
    # 第 1 章：执行摘要
    doc.add_heading('1. 执行摘要', level=1)
    
    doc.add_heading('1.1 项目定位', level=2)
    doc.add_paragraph('本项目是全国首个高校思政 VR 轻资产运营标杆项目，以桂林学院为起点，计划 3 年内复制推广至 100 所高校。核心内容获中宣部学习强国领衔出品、国家电影局首个 VR 龙标认证，打造沉浸式思政教育新生态。')
    
    doc.add_heading('1.2 核心优势（竞品无法复制）', level=2)
    advantages = [
        '🔴 独家背书：中宣部学习强国领衔出品 + 国家电影局首个 VR 龙标认证',
        '📚 教材对标：课程体系与《中国近代史纲要（2023 版）》精准对应',
        '🤖 AIGC 赋能：校本课程 AI 生成、AI 教学助手、学习报告自动生成',
        '📊 商业验证：桂林学院项目已验证，设备销售 + 订阅+AIGC 三轮驱动',
        '⚡ 快速复制：单校部署周期≤30 天，轻资产运营模式'
    ]
    for adv in advantages:
        doc.add_paragraph(adv, style='List Bullet')
    
    doc.add_heading('1.3 投资规模与回报', level=2)
    table = doc.add_table(rows=4, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '方案'
    hdr_cells[1].text = '首年投资'
    hdr_cells[2].text = 'VR 设备'
    hdr_cells[3].text = '订阅/年'
    hdr_cells[4].text = 'AIGC/年'
    hdr_cells[5].text = '回收期'
    
    data = [
        ('标配', '100 万', '20 台', '15 万', '5-10 万', '18-24 月'),
        ('高配', '200 万', '40 台', '30 万', '10-20 万', '20-26 月'),
        ('顶配', '500 万', '100 台', '75 万', '20-40 万', '24-30 月')
    ]
    for i, row in enumerate(data):
        for j, cell in enumerate(row):
            table.rows[i+1].cells[j].text = cell
    
    doc.add_heading('1.4 三年目标', level=2)
    targets = [
        '2026 年：完成 10 所高校部署，实现收入 1000 万元',
        '2027 年：完成 50 所高校部署，实现收入 4500 万元（含订阅+AIGC）',
        '2028 年：完成 100 所高校部署，实现收入 6200 万元（含订阅+AIGC）'
    ]
    for target in targets:
        doc.add_paragraph(target, style='List Bullet')
    
    doc.add_page_break()
    
    # 第 2 章：独家资源优势
    doc.add_heading('2. 独家资源优势（学习强国 + 龙标）', level=1)
    
    doc.add_heading('2.1 中宣部学习强国领衔出品', level=2)
    doc.add_paragraph('🔴 政治背书：中宣部主管的国家级学习平台，政治安全性 100% 保障')
    doc.add_paragraph('🔴 品牌效应："学习强国"品牌认知度 95%+，高校认可度极高')
    doc.add_paragraph('🔴 采购合规：高校采购零风险，审批无障碍，财务/资产处无顾虑')
    doc.add_paragraph('🔴 传播价值：学习强国平台全网传播，品牌曝光亿级')
    
    doc.add_heading('2.2 国家电影局首个 VR"龙标"电影', level=2)
    doc.add_paragraph('🎬 行业首创：全国首个 VR 内容获龙标，历史意义，里程碑事件')
    doc.add_paragraph('🎬 审查通过：国家最高级别内容审查，政治安全 100% 保障')
    doc.add_paragraph('🎬 稀缺性：重大革命题材 VR 唯一龙标，排他性，竞品无法复制')
    doc.add_paragraph('🎬 教育认可：龙标=教育内容合规认证，高校采购依据充分')
    
    doc.add_heading('2.3 竞品壁垒分析', level=2)
    doc.add_paragraph('曼恒数字：技术强，但无学习强国 + 龙标背书，政治安全性存疑')
    doc.add_paragraph('其他厂商：无官方背书，内容审查未通过，高校采购风险高')
    doc.add_paragraph('我们的优势：学习强国 + 龙标双重背书，100% 排他性，竞品无法复制')
    
    doc.add_page_break()
    
    # 第 3 章：项目背景与政策风口
    doc.add_heading('3. 项目背景与政策风口', level=1)
    doc.add_paragraph('（内容同 V2.0，略）')
    doc.add_page_break()
    
    # 第 4 章：桂林学院标杆案例
    doc.add_heading('4. 桂林学院标杆案例', level=1)
    doc.add_heading('4.1 项目概况', level=2)
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    overview = [
        ('合作院校', '桂林学院'),
        ('签约时间', '2025 年 9 月'),
        ('正式运营', '2025 年 12 月'),
        ('投资金额', '100 万元（设备销售）'),
        ('设备规模', '20 台 VR 一体机'),
        ('场地面积', '80 平方米 VR 实训室'),
        ('服务师生', '年均 3000+ 人次')
    ]
    for i, (item, value) in enumerate(overview):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = value
    
    doc.add_heading('4.2 商业模式', level=2)
    doc.add_paragraph('首年：100 万元设备销售（硬件 + 软件 + 内容打包）\n第二年起：15 万元/年内容订阅 + 5-10 万元/年 AIGC 技术服务\n合同期限：3 年（含排他条款）')
    
    doc.add_page_break()
    
    # 第 5 章：三档合作方案
    doc.add_heading('5. 三档合作方案', level=1)
    
    doc.add_heading('5.1 标配方案（100 万元）', level=2)
    doc.add_paragraph('适用对象：首次尝试 VR 思政教育的院校、普通本科院校')
    doc.add_paragraph('配置：20 台 VR 一体机、12 门课程、80㎡VR 实训室、2 天师资培训')
    doc.add_paragraph('交付周期：合同签订后 30 天')
    doc.add_paragraph('订阅服务：第二年起 15 万元/年（内容更新 + 运维）')
    doc.add_paragraph('AIGC 服务：第二年起 5-10 万元/年（校本课程 AI 生成等）')
    
    doc.add_heading('5.2 高配方案（200 万元）', level=2)
    doc.add_paragraph('适用对象：重点马院、省级思政示范院校')
    doc.add_paragraph('配置：40 台 VR 一体机、20+2 定制课程、150㎡VR 实训中心')
    doc.add_paragraph('交付周期：合同签订后 45 天')
    doc.add_paragraph('订阅服务：第二年起 30 万元/年')
    doc.add_paragraph('AIGC 服务：第二年起 10-20 万元/年')
    
    doc.add_heading('5.3 顶配方案（500 万元）', level=2)
    doc.add_paragraph('适用对象：双一流高校、省级思政教育中心')
    doc.add_paragraph('配置：100 台 VR 一体机、30+5 定制课程、300㎡VR 思政中心')
    doc.add_paragraph('交付周期：合同签订后 60 天')
    doc.add_paragraph('订阅服务：第二年起 75 万元/年')
    doc.add_paragraph('AIGC 服务：第二年起 20-40 万元/年')
    
    doc.add_page_break()
    
    # 第 6 章：课程体系与教材对标
    doc.add_heading('6. 课程体系与教材对标', level=1)
    doc.add_paragraph('（内容同 V2.0，增加学习强国 + 龙标内容认证说明）')
    doc.add_page_break()
    
    # 第 7 章：四阶教学闭环设计
    doc.add_heading('7. 四阶教学闭环设计', level=1)
    doc.add_paragraph('（内容同 V2.0）')
    doc.add_page_break()
    
    # 第 8 章：AIGC 技术服务
    doc.add_heading('8. AIGC 技术服务', level=1)
    
    doc.add_heading('8.1 服务清单', level=2)
    table = doc.add_table(rows=7, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '服务项'
    hdr_cells[1].text = '说明'
    hdr_cells[2].text = '定价'
    hdr_cells[3].text = '毛利'
    
    aigc_services = [
        ('校本课程 AI 生成', '基于学校历史/地方红色资源生成定制课程', '5-10 万/门', '85%+'),
        ('本地红色资源数字化', '扫描/建模本地红色遗址、纪念馆', '10-30 万/项目', '80%+'),
        ('AI 教学助手', '自动生成教案、课件、测试题', '3-5 万/年', '90%+'),
        ('学习报告 AI 生成', '自动分析学生数据，生成个性化报告', '包含在订阅', '95%+'),
        ('AI 虚拟讲解员', '虚拟历史人物对话系统', '5-8 万/年', '85%+'),
        ('内容智能推荐', '个性化课程路径推荐', '包含在订阅', '95%+')
    ]
    for i, (service, desc, price, margin) in enumerate(aigc_services):
        table.rows[i+1].cells[0].text = service
        table.rows[i+1].cells[1].text = desc
        table.rows[i+1].cells[2].text = price
        table.rows[i+1].cells[3].text = margin
    
    doc.add_heading('8.2 AIGC 渗透率预测', level=2)
    doc.add_paragraph('第二年：50% 客户采购 AIGC 增值服务\n第三年：70% 客户采购 AIGC 增值服务\n续约客户：90% 以上续订 AIGC 服务')
    
    doc.add_page_break()
    
    # 第 9 章：财务测算与投资回报
    doc.add_heading('9. 财务测算与投资回报', level=1)
    
    doc.add_heading('9.1 三年收入预测', level=2)
    table = doc.add_table(rows=4, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '年份'
    hdr_cells[1].text = '新增院校'
    hdr_cells[2].text = '设备销售'
    hdr_cells[3].text = '订阅收入'
    hdr_cells[4].text = 'AIGC 收入'
    hdr_cells[5].text = '总收入'
    
    forecast = [
        ('2026', '10 所', '1000 万', '0', '0', '1000 万'),
        ('2027', '40 所', '4000 万', '150 万', '100 万', '4250 万'),
        ('2028', '50 所', '5000 万', '750 万', '450 万', '6200 万')
    ]
    for i, row in enumerate(forecast):
        for j, cell in enumerate(row):
            table.rows[i+1].cells[j].text = cell
    
    doc.add_page_break()
    
    # 第 10 章：百校复制路线图
    doc.add_heading('10. 百校复制路线图', level=1)
    doc.add_paragraph('（内容同 V2.0）')
    doc.add_page_break()
    
    # 第 11 章：竞品分析与差异化
    doc.add_heading('11. 竞品分析与差异化', level=1)
    
    doc.add_heading('11.1 竞品对比', level=2)
    table = doc.add_table(rows=8, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '维度'
    hdr_cells[1].text = '曼恒（高端）'
    hdr_cells[2].text = '其他厂商'
    hdr_cells[3].text = '我们（轻资产）'
    
    comparison = [
        ('官方背书', '无', '无', '✅ 学习强国 + 龙标双重背书'),
        ('内容审查', '未通过龙标', '未通过', '✅ 国家电影局龙标认证'),
        ('教材对标', '一般', '无', '✅ 深度绑定《纲要》'),
        ('AIGC 能力', '基础 AI 对话', '无', '✅ 校本课程 AI 生成'),
        ('技术类型', 'LBE 大空间', '单机 VR', 'VR 一体机'),
        ('投资门槛', '300 万+', '50-100 万', '100-500 万'),
        ('部署周期', '60-90 天', '15-30 天', '30-45 天')
    ]
    for i, row in enumerate(comparison):
        for j, cell in enumerate(row):
            table.rows[i+1].cells[j].text = cell
    
    doc.add_heading('11.2 差异化定位', level=2)
    doc.add_paragraph('我们不走高端大空间路线，而是聚焦轻资产快速复制 + 独家背书：\n- 目标客户：普通本科/职业院校（非双一流）\n- 投资门槛：100 万起（曼恒 1/3）\n- 独家优势：学习强国 + 龙标，竞品 100% 无法复制')
    
    doc.add_page_break()
    
    # 第 12 章：Demo 演示方案
    doc.add_heading('12. Demo 演示方案', level=1)
    doc.add_paragraph('（内容同 V2.0，增加学习强国 + 龙标展示环节）')
    doc.add_page_break()
    
    # 第 13-16 章
    doc.add_heading('13. 政策支持与申报指南', level=1)
    doc.add_paragraph('（内容同 V2.0）')
    doc.add_page_break()
    
    doc.add_heading('14. 交付标准与服务承诺', level=1)
    doc.add_paragraph('（内容同 V2.0，增加内容审查说明）')
    doc.add_page_break()
    
    doc.add_heading('15. 常见问题 FAQ', level=1)
    faqs = [
        ('Q1：内容是否政治安全？', 'A：中宣部学习强国领衔出品 + 国家电影局龙标认证，双重背书，政治安全性 100% 保障。'),
        ('Q2：能否申请政府专项资金？', 'A：可以，学习强国 + 龙标背书，申报成功率更高，可申报 20-200 万元专项资金。'),
        ('Q3：AIGC 服务是否必须采购？', 'A：AIGC 为可选增值服务，基础订阅已包含课程更新和运维服务。'),
        ('Q4：龙标编号是多少？', 'A：[待补充具体龙标编号]'),
        ('Q5：学习强国如何展示？', 'A：[待补充学习强国平台链接或页面截图]')
    ]
    for question, answer in faqs:
        p = doc.add_paragraph()
        p.add_run(question).bold = True
        doc.add_paragraph(answer)
    
    doc.add_page_break()
    
    # 第 16 章：附录
    doc.add_heading('16. 附录', level=1)
    doc.add_heading('16.1 学习强国 + 龙标证明材料', level=2)
    doc.add_paragraph('- 学习强国平台页面截图\n- 龙标证书扫描件\n- 官方授权书（如适用）')
    
    doc.add_heading('16.2 合同模板（摘要）', level=2)
    doc.add_paragraph('合作协议包含：合作内容、投资金额、交付周期、服务期限、排他条款、违约责任等。')
    
    # 结尾信息
    doc.add_paragraph()
    end_info = doc.add_paragraph()
    end_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_info.add_run('编制单位：[公司名称]\n').bold = True
    end_info.add_run('编制日期：2026 年 4 月\n')
    end_info.add_run('🔴 中宣部学习强国领衔出品 | 🎬 国家电影局首个 VR 龙标认证\n').bold = True
    end_info.add_run('本方案版权归 [公司名称] 所有，未经许可不得外传').italic = True
    
    # 保存文档
    doc.save('/home/admin/.openclaw/workspace-chief-agent/百校复制方案 - 完整版 V3.0.docx')
    print('✅ Word 文档 V3.0 生成成功：百校复制方案 - 完整版 V3.0.docx')

if __name__ == '__main__':
    create_word_document()
