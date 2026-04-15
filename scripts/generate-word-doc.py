#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成百校思政教育合作方案 Word 文档
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

def create_word_document():
    doc = Document()
    
    # 设置中文字体
    doc.styles['Normal'].font.name = '微软雅黑'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    
    # 标题
    title = doc.add_heading('高校思政 VR 百校复制方案', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 版本信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('版本号：V1.0\n').bold = True
    info.add_run('编制日期：2026 年 4 月\n')
    info.add_run('编制单位：[公司名称]')
    
    doc.add_paragraph()  # 空行
    
    # 目录
    doc.add_heading('目录', level=1)
    sections = [
        '1. 执行摘要',
        '2. 项目背景与政策风口',
        '3. 桂林学院标杆案例',
        '4. 三档合作方案',
        '5. 财务测算与投资回报',
        '6. 百校复制路线图',
        '7. 政策支持与申报指南',
        '8. 交付标准与服务承诺',
        '9. 常见问题 FAQ',
        '10. 附录'
    ]
    for section in sections:
        doc.add_paragraph(section, style='List Number')
    
    doc.add_page_break()
    
    # 第 1 章：执行摘要
    doc.add_heading('1. 执行摘要', level=1)
    
    doc.add_heading('1.1 项目定位', level=2)
    doc.add_paragraph('本项目是全国首个高校思政 VR 轻资产运营标杆项目，以桂林学院为起点，计划 3 年内复制推广至 100 所高校，打造沉浸式思政教育新生态。')
    
    doc.add_heading('1.2 核心优势', level=2)
    advantages = [
        '政策强力支持：教育部大力推进"大思政课"建设，VR+ 教育是重点支持方向',
        '商业模式验证：桂林学院项目已验证可行性，毛利率达 42.5%',
        '轻资产运营：设备租赁 + 内容订阅模式，降低院校一次性投入压力',
        '标准化交付：单校部署周期≤30 天，可快速复制'
    ]
    for adv in advantages:
        doc.add_paragraph(adv, style='List Bullet')
    
    doc.add_heading('1.3 投资规模与回报', level=2)
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '方案档次'
    hdr_cells[1].text = '投资金额'
    hdr_cells[2].text = 'VR 设备数'
    hdr_cells[3].text = '投资回收期'
    
    data = [
        ('标配', '100 万元', '20 台', '18-24 个月'),
        ('高配', '200 万元', '40 台', '20-26 个月'),
        ('顶配', '500 万元', '100 台', '24-30 个月')
    ]
    for i, (level, invest, devices, payback) in enumerate(data):
        row_cells = table.rows[i+1].cells
        row_cells[0].text = level
        row_cells[1].text = invest
        row_cells[2].text = devices
        row_cells[3].text = payback
    
    doc.add_heading('1.4 三年目标', level=2)
    targets = [
        '2026 年：完成 10 所高校部署，实现收入 1000 万元',
        '2027 年：完成 50 所高校部署，实现收入 5000 万元',
        '2028 年：完成 100 所高校部署，实现收入 1 亿元'
    ]
    for target in targets:
        doc.add_paragraph(target, style='List Bullet')
    
    doc.add_page_break()
    
    # 第 2 章：项目背景与政策风口
    doc.add_heading('2. 项目背景与政策风口', level=1)
    
    doc.add_heading('2.1 思政教育现状与痛点', level=2)
    doc.add_paragraph('传统思政教育存在教学方式单一、内容抽象难懂、实践环节薄弱、效果评估困难等问题。Z 世代大学生是数字原住民，习惯沉浸式、互动式学习。')
    
    doc.add_heading('2.2 VR 技术优势', level=2)
    doc.add_paragraph('VR 技术具有沉浸式体验、互动性强、场景还原、安全可控、数据追踪等优势，可提升学习兴趣，加深理解记忆，突破时空限制。')
    
    doc.add_heading('2.3 政策支持', level=2)
    policies = [
        '《全面推进"大思政课"建设的工作方案》（教育部等十部门，2022 年）',
        '《虚拟现实与行业应用融合发展行动计划（2022-2026 年）》（工信部等五部门）',
        '《新时代学校思想政治理论课改革创新实施方案》（教育部，2020 年）'
    ]
    for policy in policies:
        doc.add_paragraph(policy, style='List Bullet')
    
    doc.add_page_break()
    
    # 第 3 章：桂林学院标杆案例
    doc.add_heading('3. 桂林学院标杆案例', level=1)
    
    doc.add_heading('3.1 项目概况', level=2)
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    overview = [
        ('合作院校', '桂林学院'),
        ('签约时间', '2025 年 9 月'),
        ('正式运营', '2025 年 12 月'),
        ('投资金额', '100 万元'),
        ('设备规模', '20 台 VR 一体机'),
        ('场地面积', '80 平方米 VR 实训室'),
        ('服务师生', '年均 3000+ 人次')
    ]
    for i, (item, value) in enumerate(overview):
        table.rows[i].cells[0].text = item
        table.rows[i].cells[1].text = value
    
    doc.add_heading('3.2 课程内容', level=2)
    doc.add_paragraph('包含 12 门 VR 思政精品课程，涵盖党史教育、国情教育、红色文化、价值观教育、校史教育、实践实训六大模块。')
    
    doc.add_heading('3.3 教学效果', level=2)
    effects = [
        '学习兴趣：65 分 → 92 分（+41.5%）',
        '知识掌握：72 分 → 88 分（+22.2%）',
        '课堂参与：58 分 → 94 分（+62.1%）',
        '记忆留存：45% → 78%（+73.3%）',
        '课程满意度：75 分 → 96 分（+28.0%）'
    ]
    for effect in effects:
        doc.add_paragraph(effect, style='List Bullet')
    
    doc.add_heading('3.4 财务表现', level=2)
    doc.add_paragraph('年收入：60 万元\n年成本：34.5 万元\n年毛利润：25.5 万元\n毛利率：42.5%\n投资回收期：约 20 个月')
    
    doc.add_page_break()
    
    # 第 4 章：三档合作方案
    doc.add_heading('4. 三档合作方案', level=1)
    
    doc.add_heading('4.1 标配方案（100 万元）', level=2)
    doc.add_paragraph('适用对象：首次尝试 VR 思政教育的院校、普通本科院校、单校试点项目')
    doc.add_paragraph('配置：20 台 VR 一体机、12 门课程、80㎡VR 实训室、2 天师资培训、3 个月运营指导')
    doc.add_paragraph('交付周期：合同签订后 30 天')
    
    doc.add_heading('4.2 高配方案（200 万元）', level=2)
    doc.add_paragraph('适用对象：重点马院、省级思政示范院校、有一定 VR 应用基础的院校')
    doc.add_paragraph('配置：40 台 VR 一体机、20+2 定制课程、150㎡VR 实训中心、5 天师资培训、6 个月驻校支持')
    doc.add_paragraph('交付周期：合同签订后 45 天')
    
    doc.add_heading('4.3 顶配方案（500 万元）', level=2)
    doc.add_paragraph('适用对象：双一流高校、省级思政教育中心、区域示范标杆项目')
    doc.add_paragraph('配置：100 台 VR 一体机、30+5 定制课程、300㎡VR 思政中心、10 天系统培训、12 个月驻校支持')
    doc.add_paragraph('交付周期：合同签订后 60 天')
    
    doc.add_page_break()
    
    # 第 5 章：财务测算与投资回报
    doc.add_heading('5. 财务测算与投资回报', level=1)
    
    doc.add_heading('5.1 收入模型', level=2)
    doc.add_paragraph('收入来源：设备租赁费（30%）、内容订阅费（15%）、运营服务费（10%）、定制开发费（按需）')
    
    table = doc.add_table(rows=4, cols=6)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '方案'
    hdr_cells[1].text = '设备租赁'
    hdr_cells[2].text = '内容订阅'
    hdr_cells[3].text = '运营服务'
    hdr_cells[4].text = '定制开发'
    hdr_cells[5].text = '合计'
    
    income_data = [
        ('标配', '30 万', '15 万', '10 万', '5 万', '60 万'),
        ('高配', '60 万', '30 万', '20 万', '10 万', '120 万'),
        ('顶配', '150 万', '75 万', '50 万', '25 万', '300 万')
    ]
    for i, row in enumerate(income_data):
        for j, cell in enumerate(row):
            table.rows[i+1].cells[j].text = cell
    
    doc.add_heading('5.2 盈利分析', level=2)
    doc.add_paragraph('单校毛利率：42.5%\n投资回收期：约 20 个月\n三年累计收入：1.6 亿元\n三年累计毛利：6800 万元')
    
    doc.add_page_break()
    
    # 第 6 章：百校复制路线图
    doc.add_heading('6. 百校复制路线图', level=1)
    
    doc.add_heading('6.1 总体目标', level=2)
    doc.add_paragraph('3 年 100 校，打造全国高校思政 VR 教育第一品牌')
    
    doc.add_heading('6.2 阶段规划', level=2)
    stages = [
        '第一阶段（2026 Q2-Q4）：完成 10 所高校部署，收入 1000 万元',
        '第二阶段（2027 全年）：完成 50 所高校部署，收入 5000 万元',
        '第三阶段（2028 全年）：完成 100 所高校部署，收入 1 亿元'
    ]
    for stage in stages:
        doc.add_paragraph(stage, style='List Bullet')
    
    doc.add_heading('6.3 区域布局', level=2)
    regions = [
        '华东（江浙沪皖鲁）：25 所',
        '华北（京津冀晋）：20 所',
        '华南（粤桂琼）：15 所',
        '华中（鄂湘豫）：15 所',
        '西南（川渝云贵）：15 所',
        '西北（陕甘新）：10 所'
    ]
    for region in regions:
        doc.add_paragraph(region, style='List Bullet')
    
    doc.add_page_break()
    
    # 第 7 章：政策支持与申报指南
    doc.add_heading('7. 政策支持与申报指南', level=1)
    
    doc.add_heading('7.1 可申报项目', level=2)
    projects = [
        '国家级：大思政课示范项目（50-100 万）、虚拟仿真实验教学项目（30-80 万）',
        '省级：思政工作精品项目（20-50 万）、智慧马院建设（30-100 万）、虚拟仿真基地（50-200 万）'
    ]
    for project in projects:
        doc.add_paragraph(project, style='List Bullet')
    
    doc.add_heading('7.2 申报策略', level=2)
    doc.add_paragraph('最佳时机：项目落地后 3-6 个月（有运营数据支撑）\n申报成功率提升技巧：突出创新性、数据支撑、示范效应、校政企协同')
    
    doc.add_page_break()
    
    # 第 8 章：交付标准与服务承诺
    doc.add_heading('8. 交付标准与服务承诺', level=1)
    
    doc.add_heading('8.1 交付标准', level=2)
    doc.add_paragraph('硬件交付：100% 开机正常，网络延迟≤20ms\n软件交付：所有功能正常运行，课程内容完整可用\n场地交付：符合设计方案，安全设施完备')
    
    doc.add_heading('8.2 服务承诺', level=2)
    doc.add_paragraph('培训服务：基础培训 1 天、教学培训 1-4 天、认证培训 5-10 天\n运维服务：电话咨询即时响应、远程支持≤30 分钟、现场服务≤4 小时\n内容更新：季度常规更新、重大节日专题更新')
    
    doc.add_heading('8.3 质量保证', level=2)
    doc.add_paragraph('硬件设备质保 2 年、软件系统质保 1 年、装修工程质保 1 年\n系统可用性≥99%、故障响应≤30 分钟、故障解决≤24 小时')
    
    doc.add_page_break()
    
    # 第 9 章：常见问题 FAQ
    doc.add_heading('9. 常见问题 FAQ', level=1)
    
    faqs = [
        ('Q1：VR 设备会不会让学生头晕？', 'A：选用 PICO 4 Enterprise 高分辨率高刷新率，95% 以上学生无眩晕感。'),
        ('Q2：网络条件要求高吗？', 'A：本地内容无需网络，建议带宽≥100Mbps，提供离线内容包。'),
        ('Q3：教师没有 VR 经验怎么办？', 'A：提供系统化培训，2 天培训即可独立开展 VR 教学。'),
        ('Q4：投资是一次性还是分期？', 'A：支持一次性付款和分期付款，分期通常为 30% 首付 +40% 交付 +30% 验收。'),
        ('Q5：能否申请政府专项资金？', 'A：可以，提供全套申报材料支持，可申报 20-200 万元不等专项资金。')
    ]
    for question, answer in faqs:
        p = doc.add_paragraph()
        p.add_run(question).bold = True
        doc.add_paragraph(answer)
    
    doc.add_page_break()
    
    # 第 10 章：附录
    doc.add_heading('10. 附录', level=1)
    
    doc.add_heading('10.1 合同模板（摘要）', level=2)
    doc.add_paragraph('合作协议包含：合作内容、投资金额、交付周期、服务期限、双方权利义务、违约责任、争议解决等条款。')
    
    doc.add_heading('10.2 设备清单（标配方案）', level=2)
    doc.add_paragraph('VR 一体机 20 台、管理主机 1 台、显示设备 2 台、充电存储柜 1 台、网络设备 1 套等，硬件设备成本合计 20.9 万元。')
    
    doc.add_heading('10.3 课程内容目录', level=2)
    courses = [
        '党史教育：《长征路上的抉择》《开国大典》《改革开放春潮》等',
        '国情教育：《脱贫攻坚伟大成就》《大国重器》等',
        '红色文化：《井冈山精神》《延安岁月》《西柏坡赶考》等',
        '价值观教育：《榜样的力量》《青春告白祖国》等',
        '校史教育：《桂林学院发展史》（可定制）',
        '实践实训：《思政微课演练》《演讲比赛模拟》等'
    ]
    for course in courses:
        doc.add_paragraph(course, style='List Bullet')
    
    # 结尾信息
    doc.add_paragraph()
    end_info = doc.add_paragraph()
    end_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end_info.add_run('编制单位：[公司名称]\n').bold = True
    end_info.add_run('编制日期：2026 年 4 月\n')
    end_info.add_run('本方案版权归 [公司名称] 所有，未经许可不得外传').italic = True
    
    # 保存文档
    doc.save('/home/admin/.openclaw/workspace-chief-agent/百校复制方案 - 完整版.docx')
    print('✅ Word 文档生成成功：百校复制方案 - 完整版.docx')

if __name__ == '__main__':
    create_word_document()
