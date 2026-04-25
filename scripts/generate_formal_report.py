#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 XR 项目立项报告 Word 文档（正式报告格式）"""

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_ORIENT

def create_formal_report():
    doc = Document()
    
    # ========== 页面设置 ==========
    section = doc.sections[0]
    section.top_margin = Cm(3)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)
    
    # ========== 封面 ==========
    # 标题（居中，大字号）
    for _ in range(3):
        doc.add_paragraph()
    
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('关于联合学习强国《长征·英雄》VR 大空间项目\n建设"AI+XR 科技思政沉浸空间与数智创研中心"的\n实施方案')
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.name = '黑体'
    
    for _ in range(2):
        doc.add_paragraph()
    
    # 副标题
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('（台长决策版）')
    run.font.size = Pt(14)
    run.font.name = '楷体'
    
    for _ in range(8):
        doc.add_paragraph()
    
    # 呈报信息
    info_table = doc.add_table(rows=4, cols=2)
    info_table.autofit = False
    info_table.allow_autofit = False
    
    # 设置列宽
    info_table.columns[0].width = Cm(4)
    info_table.columns[1].width = Cm(10)
    
    # 填充内容
    data = [
        ['呈报单位：', '融媒体中心'],
        ['呈报日期：', '2026 年 4 月 17 日'],
        ['呈报人：', '强国小马（数字团队总监）'],
        ['联系方式：', '钉钉私信'],
    ]
    for i, row in enumerate(data):
        info_table.rows[i].cells[0].text = row[0]
        info_table.rows[i].cells[1].text = row[1]
        for cell in info_table.rows[i].cells:
            cell.paragraphs[0].runs[0].font.size = Pt(12)
    
    doc.add_page_break()
    
    # ========== 目录 ==========
    doc.add_heading('目  录', level=1)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    toc = [
        ('一、项目背景与战略定位', 1),
        ('  （一）政策依据', 2),
        ('  （二）对标案例：河南台成功经验', 2),
        ('  （三）合作模式', 2),
        ('二、投资估算与资金筹措', 1),
        ('  （一）投资明细', 2),
        ('  （二）资金筹措', 2),
        ('三、财务测算与投资回报', 1),
        ('  （一）收入预测', 2),
        ('  （二）成本结构', 2),
        ('  （三）投资回报分析', 2),
        ('四、项目实施计划', 1),
        ('  （一）时间节点', 2),
        ('  （二）关键里程碑', 2),
        ('五、风险评估与防控措施', 1),
        ('六、结论与建议', 1),
        ('附件清单', 1),
    ]
    
    for item, level in toc:
        p = doc.add_paragraph(item)
        if level == 1:
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.runs[0]
            run.font.size = Pt(14)
            run.font.bold = True
        else:
            p.paragraph_format.left_indent = Cm(2)
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(3)
            run = p.runs[0]
            run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # ========== 正文 ==========
    
    # 一、项目背景与战略定位
    doc.add_heading('一、项目背景与战略定位', level=1)
    
    doc.add_heading('（一）政策依据', level=2)
    p = doc.add_paragraph('本项目是落实台内 2026 年度实施方案的具体行动，第 18、19、33 条直接支撑：')
    
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    # 表头
    headers = ['条款', '原文摘要', '项目对应']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 数据
    data = [
        ['第 18 条', '围绕长征胜利 90 周年，以"历史影像 + 现实寻访 + 互动传播"模式传承红色基因', '《长征·英雄》VR 大空间体验区'],
        ['第 19 条', '谋划推出内蒙古自治区成立 80 周年宣传项目，高水平站位、高品质策划', 'AI 自主研发实验室（内蒙 80 周年主题）'],
        ['第 33 条', '组建 XR 技术团队，搭建 XR 演播室，通过外训与实战建立虚拟场景资源库', '数智创研中心 + XR 技术团队建设'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = cell_text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    
    doc.add_paragraph()
    p = doc.add_paragraph('结论：本项目不是新增项目，是落实台内 2026 年度实施方案的具体行动。')
    p.paragraph_format.left_indent = Cm(2)
    p.runs[0].italic = True
    
    doc.add_heading('（二）对标案例：河南台成功经验', level=2)
    p = doc.add_paragraph('河南电视台《唐宫夜宴》VR 项目已验证模式可行：')
    
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'
    
    data = [
        ['投入成本', '300 万元'],
        ['年营收', '2000 万+'],
        ['年客流', '50 万 + 人次'],
        ['ROI（3 年）', '1:6.7'],
        ['回本周期', '6-8 个月'],
        ['社会影响', '央视报道、文旅部典型案例、全国媒体学习标杆'],
    ]
    for i, row in enumerate(data):
        table.rows[i].cells[0].text = row[0]
        table.rows[i].cells[1].text = row[1]
        for cell in table.rows[i].cells:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    
    doc.add_paragraph()
    p = doc.add_paragraph('我台对比优势：')
    p.runs[0].bold = True
    
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    
    headers = ['维度', '河南台', '我台', '优势对比']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    data = [
        ['投入成本', '300 万', '170 万', '✅ 低 43%'],
        ['回本周期', '6-8 个月', '4-7 个月', '✅ 提前 2-3 个月'],
        ['ROI（3 年）', '1:6.7', '1:8.8', '✅ 高 60%'],
        ['IP 资源', '自有 IP', '学习强国联合', '✅ 政治背书更强'],
        ['政策窗口', '常态化', '双节点驱动', '✅ 更难得'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = cell_text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('（三）合作模式', level=2)
    p = doc.add_paragraph('学习强国城市合伙人协议（三方合作）：')
    
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    
    headers = ['角色', '合作方', '主要职责']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    data = [
        ['甲方', '学习强国学习平台', '开设专题专区、协调政府资源、宣传推广、政策支持'],
        ['乙方', '未来新视界科技（北京）', '提供内容 + 软硬件解决方案（30 台 VR 仅 18 万）+ 空间设计 + 运营管理系统'],
        ['丙方', '内蒙古广播电视台', '场地（旧台体育馆）+ 运营团队 + 票务销售 + 设备运维 + 市场拓展'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = cell_text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    
    doc.add_paragraph()
    p = doc.add_paragraph('合作期限：3 年（可续约）\n分账方式：30 元/台（散客）+ 15 元/人（学生研学），按月结算，无保底压力')
    p.paragraph_format.left_indent = Cm(2)
    
    doc.add_page_break()
    
    # 二、投资估算与资金筹措
    doc.add_heading('二、投资估算与资金筹措', level=1)
    
    doc.add_heading('（一）投资明细', level=2)
    p = doc.add_paragraph('项目总投资 170 万元（一次性投入）：')
    
    table = doc.add_table(rows=5, cols=4)
    table.style = 'Table Grid'
    
    headers = ['序号', '项目', '金额（万元）', '说明']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 调整列宽
    table.columns[0].width = Cm(2)
    table.columns[1].width = Cm(6)
    table.columns[2].width = Cm(3)
    table.columns[3].width = Cm(7)
    
    data = [
        ['1', 'IP 授权费', '50', '学习强国官方授权，一次性支付'],
        ['2', '硬件采购', '18', '30 台 VR 眼镜 + 配套（乙方报价确认）'],
        ['3', '场地改造', '100', '旧台体育馆水电暖 + 声学改造（控制目标）'],
        ['4', '合计', '170', '-'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = cell_text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
            if row[1] == '合计':
                run.bold = True
    
    doc.add_paragraph()
    p = doc.add_paragraph('对比河南台：300 万 → 我台 170 万，节省 130 万元（-43%）')
    p.paragraph_format.left_indent = Cm(2)
    p.runs[0].italic = True
    
    doc.add_heading('（二）资金筹措', level=2)
    p = doc.add_paragraph('资金来源：台内自筹\n支付方式：分阶段支付\n  • IP 授权费：签约后 5 日内支付 50 万元\n  • 硬件采购：设备交付后支付 18 万元\n  • 场地改造：按工程进度分期支付')
    p.paragraph_format.left_indent = Cm(2)
    
    doc.add_page_break()
    
    # 三、财务测算与投资回报
    doc.add_heading('三、财务测算与投资回报', level=1)
    
    doc.add_heading('（一）收入预测', level=2)
    p = doc.add_paragraph('年营收预测（保守预估）：')
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    
    headers = ['收入来源', '金额（万元）', '计算方式']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    data = [
        ['门票收入（散客）', '948', '158 元 × 6 万人次'],
        ['门票收入（团体）', '200', '100 元 × 2 万人次'],
        ['学生研学', '75', '150 元 × 5000 人次'],
        ['文创产品销售', '50', '利润全归我方'],
        ['合计', '1273', '保守预估'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = cell_text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    
    doc.add_paragraph()
    p = doc.add_paragraph('市场依据：呼和浩特市常住人口 395 万，目标客群 80-100 万，年客流 6-9 万人次仅占 0.75%-1.1%')
    p.paragraph_format.left_indent = Cm(2)
    p.runs[0].font.size = Pt(10)
    
    doc.add_heading('（二）成本结构', level=2)
    p = doc.add_paragraph('年运营成本：')
    
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    
    headers = ['成本项', '金额（万元）', '说明']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    data = [
        ['IP 使用费（分账）', '187.5', '30 元×6 万 + 15 元×0.5 万'],
        ['人力成本', '180', '15-20 人团队'],
        ['运维成本', '50', '设备维护、能耗'],
        ['营销推广', '80', '渠道推广、活动'],
        ['其他', '25', '-'],
        ['合计', '522.5', '-'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = cell_text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    
    doc.add_heading('（三）投资回报分析', level=2)
    
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    
    headers = ['指标', '数值', '对比河南台']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    data = [
        ['年营收', '900-1200 万元', '河南台 2000 万+'],
        ['年成本', '522.5 万元', '-'],
        ['年利润', '330-630 万元', '-'],
        ['总投资', '170 万元', '河南台 300 万'],
        ['回本周期', '4-7 个月', '河南台 6-8 个月 ✅'],
        ['ROI（3 年）', '1:8.8', '河南台 1:6.7 ✅'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = cell_text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    
    doc.add_paragraph()
    p = doc.add_paragraph('结论：投得少，赚得多，回本快。')
    p.paragraph_format.left_indent = Cm(2)
    p.runs[0].italic = True
    p.runs[0].bold = True
    
    doc.add_page_break()
    
    # 四、项目实施计划
    doc.add_heading('四、项目实施计划', level=1)
    
    doc.add_heading('（一）时间节点', level=2)
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    
    headers = ['阶段', '时间周期', '核心任务']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    data = [
        ['筹备阶段', '2026.05-2026.07', '方案细化、基建评估、预算审批、协议签署'],
        ['建设阶段', '2026.08-2026.11', '基础设施改造、设备采购部署、内容引进'],
        ['试运营阶段', '2026.12-2027.02', '系统联调、人员培训、压力测试'],
        ['正式运营', '2027.03 起', '全渠道上线、持续迭代优化'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = cell_text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    
    doc.add_heading('（二）关键里程碑', level=2)
    
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    
    headers = ['序号', '时间', '里程碑']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    data = [
        ['1', '2026.07', '完成基建评估与预算审批，签署学习强国城市合伙人协议'],
        ['2', '2026.11', '场馆改造完成，设备部署到位'],
        ['3', '2026.12', '长征胜利 90 周年《长征·英雄》首发 ⭐'],
        ['4', '2027.03', '正式运营启动'],
        ['5', '2027.10', '内蒙古自治区成立 80 周年自主研发内容上线 ⭐'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = cell_text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    
    doc.add_page_break()
    
    # 五、风险评估与防控措施
    doc.add_heading('五、风险评估与防控措施', level=1)
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    
    headers = ['风险类型', '风险等级', '主要防控措施']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    data = [
        ['政治风险', '低', '学习强国联合出品 + 官方媒体资质，内容提前送审'],
        ['财务风险', '低', '无保底条款 + 按人次分账 + 170 万低投资'],
        ['技术风险', '中', '成熟供应商 + 备用设备 + 运维团队'],
        ['运营风险', '中', '多渠道营销推广 + 学校/企事业单位合作 + 学习强国引流'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            cell = table.rows[i+1].cells[j]
            cell.text = cell_text
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(10)
    
    doc.add_paragraph()
    p = doc.add_paragraph('结论：项目风险可控，不做风险更大（错失数字化转型窗口期、被兄弟台站超越、浪费旧台闲置资产、错过双节点）。')
    p.paragraph_format.left_indent = Cm(2)
    p.runs[0].italic = True
    
    # 六、结论与建议
    doc.add_heading('六、结论与建议', level=1)
    
    doc.add_heading('（一）核心优势总结', level=2)
    p = doc.add_paragraph('相较于传统模式，本项目具备四大核心优势：')
    
    items = [
        '财务风险超低：IP 授权费 50 万 + 硬件 18 万 + 按人次分账（30 元/台），无保底压力，4-7 个月回本',
        '政治背书强大：学习强国联合出品 + 长征 90 周年 + 台内年度实施方案支撑',
        '成功案例对标：河南台《唐宫夜宴》VR 项目已验证模式可行（ROI 1:6.7），我台可达 1:8.8',
        '扩展性良好：自主研发能力建设，长期可持续发展，可复制推广至全区乃至全国',
    ]
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(f'{i}. {item}')
        p.paragraph_format.left_indent = Cm(2)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_heading('（二）建议事项', level=2)
    
    items = [
        '尽快立项启动：旧台体育馆基建评估与改造工作',
        '成立专项工作组：台领导牵头，统筹项目推进',
        '签署学习强国城市合伙人协议：锁定呼和浩特市授权（50 万 IP 费 + 分账模式）',
        '以长征胜利 90 周年为首发节点：2026.12《长征·英雄》VR 大空间上线',
        '同步启动自主研发团队建设：为 2027 年内蒙 80 周年节点做准备',
    ]
    for i, item in enumerate(items, 1):
        p = doc.add_paragraph(f'{i}. {item}')
        p.paragraph_format.left_indent = Cm(2)
        p.paragraph_format.space_after = Pt(6)
    
    doc.add_paragraph()
    p = doc.add_paragraph('一句话总结：170 万投资，4-7 个月回本，ROI 1:8.8，学习强国背书，双节点驱动，河南台已验证成功，不做是失职，做了是政绩。')
    p.paragraph_format.left_indent = Cm(2)
    p.paragraph_format.space_before = Pt(12)
    p.runs[0].bold = True
    p.runs[0].italic = True
    
    doc.add_page_break()
    
    # 附件清单
    doc.add_heading('附件清单', level=1)
    
    attachments = [
        '1. 《学习强国〈长征·英雄〉VR 大空间三方城市合伙人协议》',
        '2. 《XR 项目财务测算表》（详细版）',
        '3. 《场馆改造技术方案》',
        '4. 《设备采购清单及预算》',
        '5. 《河南电视台〈唐宫夜宴〉VR 项目案例分析报告》',
        '6. 《台内 2026 年度实施方案》（第 18、19、33 条）',
        '7. 《呼和浩特市市场调研报告》',
        '8. 《项目实施甘特图》',
    ]
    for att in attachments:
        p = doc.add_paragraph(att)
        p.paragraph_format.left_indent = Cm(2)
        p.paragraph_format.space_after = Pt(6)
    
    # 页脚
    doc.add_paragraph()
    doc.add_paragraph('_' * 80)
    p = doc.add_paragraph()
    p.add_run('呈报人：强国小马（数字团队总监）\n')
    p.add_run('联合呈报：内蒙古 XR 项目团队\n')
    p.add_run('联系方式：钉钉私信\n')
    p.add_run('呈报日期：2026 年 4 月 17 日')
    
    # 保存文档
    doc.save('reports/XR 项目立项报告.docx')
    print('✅ 正式报告 Word 文档生成成功！')
    print('文件路径：/home/admin/.openclaw/workspace-chief-agent/reports/XR 项目立项报告.docx')

if __name__ == '__main__':
    create_formal_report()
