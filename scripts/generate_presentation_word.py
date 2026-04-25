#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 XR 项目汇报 Word 文档（精简汇报版·1-2 页）"""

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

def create_presentation_doc():
    doc = Document()
    
    # 设置页面 - A4 横向，更适合汇报
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    
    # 标题
    title = doc.add_heading('XR 项目立项汇报', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题
    subtitle = doc.add_paragraph('联合学习强国《长征·英雄》VR 大空间项目')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    subtitle.runs[0].font.size = Pt(14)
    
    # 呈报信息
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run('呈报单位：融媒体中心  |  呈报日期：2026 年 4 月 17 日  |  呈报人：强国小马')
    info.runs[0].font.size = Pt(10)
    
    doc.add_paragraph('_' * 100)
    
    # 核心数据（用醒目方式呈现）
    doc.add_heading('📊 核心数据（台长最关心）', level=1)
    
    # 核心数据表格
    table = doc.add_table(rows=2, cols=5)
    table.style = 'Table Grid'
    
    # 表头
    headers = ['总投资', '回本周期', 'ROI（3 年）', '年利润', '对比河南台']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.width = Cm(5.5)
    
    # 数据
    data = ['170 万元', '4-7 个月', '1:8.8', '330-630 万', '✅ 全面超越']
    for i, d in enumerate(data):
        cell = table.rows[1].cells[i]
        cell.text = d
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if i == 4:  # 对比河南台
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)  # 绿色
    
    doc.add_paragraph()
    
    # 三大亮点
    doc.add_heading('✨ 三大亮点', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('1️⃣ 投资更低：\n')
    run.bold = True
    p.add_run('   170 万 vs 河南台 300 万，省 130 万（-43%）\n\n')
    
    run = p.add_run('2️⃣ 回本更快：\n')
    run.bold = True
    p.add_run('   4-7 个月 vs 河南台 6-8 个月，提前 2-3 个月\n\n')
    
    run = p.add_run('3️⃣ ROI 更高：\n')
    run.bold = True
    p.add_run('   1:8.8 vs 河南台 1:6.7，多赚 60%')
    
    doc.add_paragraph()
    
    # 为什么必须做
    doc.add_heading('🎯 为什么必须做这个项目？', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # 三列
    cell1 = table.rows[0].cells[0]
    cell1.text = '📋 政策要求\n\n台内 2026 年度实施方案第 18、19、33 条直接支撑\n\n不是新增项目，是落实既有部署'
    
    cell2 = table.rows[0].cells[1]
    cell2.text = '✅ 成功案例\n\n河南台《唐宫夜宴》VR 项目\n投入 300 万，年营收 2000 万+\nROI 1:6.7，央视报道'
    
    cell3 = table.rows[0].cells[2]
    cell3.text = '⏰ 时间窗口\n\n长征胜利 90 周年（2026.12）\n内蒙古自治区成立 80 周年（2027.10）\n双节点驱动，10 年一遇'
    
    # 设置列宽
    for cell in table.rows[0].cells:
        cell.width = Cm(9)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 金句
    quote = doc.add_paragraph()
    quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote.italic = True
    quote.add_run('"不做是失职，做了是政绩"')
    quote.runs[0].font.size = Pt(14)
    quote.runs[0].bold = True
    
    doc.add_paragraph()
    
    # 合作模式
    doc.add_heading('🤝 合作模式（铁三角组合）', level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    cell1 = table.rows[0].cells[0]
    cell1.text = '🔴 甲方：学习强国\n\n• 开设专题专区\n• 协调政府资源\n• 宣传推广\n• 政策支持'
    
    cell2 = table.rows[0].cells[1]
    cell2.text = '🔵 乙方：未来新视界\n\n• 提供内容\n• 软硬件解决方案（30 台 VR 仅 18 万）\n• 空间设计\n• 运营管理系统'
    
    cell3 = table.rows[0].cells[2]
    cell3.text = '🟢 丙方：内蒙古台\n\n• 场地（旧台体育馆）\n• 运营团队\n• 票务销售\n• 设备运维\n• 市场拓展'
    
    for cell in table.rows[0].cells:
        cell.width = Cm(9)
    
    doc.add_paragraph()
    
    # 财务测算
    doc.add_heading('💰 财务测算', level=1)
    
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    
    # 左列：投资
    cell1 = table.rows[0].cells[0]
    cell1.text = '📍 投资（一次性）\n\nIP 授权费：50 万\n硬件采购：18 万（30 台 VR+ 配套）\n场地改造：100 万\n─────────────\n总投资：170 万'
    
    # 右列：回报
    cell2 = table.rows[0].cells[1]
    cell2.text = '📈 回报（年）\n\n营收：900-1200 万\n成本：522.5 万\n─────────────\n利润：330-630 万\n\n回本：4-7 个月\nROI（3 年）：1:8.8'
    
    for cell in table.rows[0].cells:
        cell.width = Cm(13)
    
    doc.add_paragraph()
    
    # 实施计划
    doc.add_heading('📅 实施计划', level=1)
    
    table = doc.add_table(rows=1, cols=5)
    table.style = 'Table Grid'
    
    data = [
        '2026.05-07\n筹备阶段\n\n方案细化\n基建评估\n协议签署',
        '2026.08-11\n建设阶段\n\n场馆改造\n设备部署\n内容引进',
        '2026.12\n首发节点⭐\n\n长征 90 周年\n《长征·英雄》上线',
        '2027.03\n正式运营\n\n全渠道上线\n持续迭代',
        '2027.10\n自主研发⭐\n\n内蒙 80 周年\n内容上线'
    ]
    for i, d in enumerate(data):
        cell = table.rows[0].cells[i]
        cell.text = d
        cell.width = Cm(5.5)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # 风险防控
    doc.add_heading('⚠️ 风险与防控', level=1)
    
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    data = [
        '🔴 政治风险\n等级：低\n\n学习强国背书\n官方媒体资质\n内容提前送审',
        '💰 财务风险\n等级：低\n\n无保底条款\n按人次分账\n170 万低投资',
        '⚙️ 技术风险\n等级：中\n\n成熟供应商\n备用设备\n运维团队',
        '📊 运营风险\n等级：中\n\n多渠道营销\n学校/企事业合作\n学习强国引流'
    ]
    for i, d in enumerate(data):
        cell = table.rows[0].cells[i]
        cell.text = d
        cell.width = Cm(7)
    
    doc.add_paragraph()
    
    # 决策请求
    doc.add_heading('✅ 决策请求（请台长拍板）', level=1)
    
    p = doc.add_paragraph()
    run = p.add_run('1. 批准立项\n')
    run.bold = True
    p.add_run('   启动旧台体育馆基建评估\n')
    run = p.add_run('2. 成立专项工作组\n')
    run.bold = True
    p.add_run('   台领导牵头，统筹项目推进\n')
    run = p.add_run('3. 签署学习强国城市合伙人协议\n')
    run.bold = True
    p.add_run('   锁定呼和浩特市授权（50 万 IP 费 + 分账模式）\n')
    run = p.add_run('4. 以长征胜利 90 周年为首发节点\n')
    run.bold = True
    p.add_run('   2026.12《长征·英雄》VR 大空间上线')
    
    doc.add_paragraph()
    
    # 一句话总结（醒目）
    quote = doc.add_paragraph()
    quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote.italic = True
    quote.add_run('"170 万投资，4-7 个月回本，ROI 1:8.8，学习强国背书，双节点驱动，\n河南台已验证成功，不做是失职，做了是政绩。"')
    quote.runs[0].font.size = Pt(14)
    quote.runs[0].bold = True
    
    doc.add_paragraph()
    doc.add_paragraph('_' * 100)
    
    # 页脚
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run('呈报人：强国小马（数字团队总监）  |  联系方式：钉钉私信')
    footer.runs[0].font.size = Pt(9)
    
    # 保存文档
    doc.save('reports/XR 项目立项汇报 - 精简版.docx')
    print('✅ 汇报 Word 文档生成成功！')
    print('文件路径：/home/admin/.openclaw/workspace-chief-agent/reports/XR 项目立项汇报 - 精简版.docx')
    print('页面设置：A4 横向，适合打印和汇报演示')

if __name__ == '__main__':
    create_presentation_doc()
