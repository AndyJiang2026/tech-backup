#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""生成 XR 项目立项报告 Word 文档（台长决策版）"""

from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()
    
    # 设置页面
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(3)
    
    # 标题
    title = doc.add_heading('关于联合学习强国《长征·英雄》VR 大空间项目\n建设"AI+XR 科技思政沉浸空间与数智创研中心"的实施方案', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 副标题
    subtitle = doc.add_paragraph('（台长决策版·最终）')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].italic = True
    
    # 呈报信息
    doc.add_paragraph()
    info = doc.add_paragraph()
    info.add_run('呈报单位：融媒体中心\n')
    info.add_run('呈报日期：2026 年 4 月 17 日\n')
    info.add_run('密    级：内部')
    
    doc.add_page_break()
    
    # 目录
    doc.add_heading('目  录', level=1)
    toc = [
        '一、台长最关心的 5 个问题（3 分钟读完）',
        '二、为什么必须做这个项目？',
        '三、投资与回报（台长最关心）',
        '四、风险与防控（台长最担心）',
        '五、实施计划（台长最关注）',
        '六、决策请求（台长拍板）',
        '附件清单'
    ]
    for item in toc:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_page_break()
    
    # 第一部分
    doc.add_heading('一、台长最关心的 5 个问题（3 分钟读完）', level=1)
    
    # 问题 1
    doc.add_heading('问题 1：为什么要做这个项目？', level=2)
    p = doc.add_paragraph()
    p.add_run('答：三个理由，不得不做：\n\n')
    p.add_run('1. 台内年度实施方案明确要求\n')
    p.add_run('   第 18、19、33 条直接支撑，不是新增项目，是落实既有部署\n\n')
    p.add_run('2. 河南台已验证成功\n')
    p.add_run('   《唐宫夜宴》VR 项目投入 300 万，年营收 2000 万+，ROI 1:6.7，央视报道、文旅部典型案例\n\n')
    p.add_run('3. 时间窗口难得\n')
    p.add_run('   长征胜利 90 周年（2026.12）+ 内蒙古自治区成立 80 周年（2027.10），双节点驱动，错过再等 10 年')
    
    quote = doc.add_paragraph()
    quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote.italic = True
    quote.add_run('"一句话：政策要求 + 成功案例 + 时间窗口，不做是失职，做了是政绩。"')
    
    # 问题 2
    doc.add_heading('问题 2：要投多少钱？会不会亏？', level=2)
    p = doc.add_paragraph()
    p.add_run('答：投资 170 万，4-7 个月回本，3 年 ROI 1:8.8\n\n')
    
    # 投资表格
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['项目', '金额', '说明']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    
    data = [
        ['IP 授权费', '50 万', '一次性，学习强国官方授权'],
        ['硬件采购', '18 万', '30 台 VR 设备 + 配套（乙方报价确认）'],
        ['场地改造', '100 万', '旧台体育馆水电暖 + 声学改造'],
        ['总投资', '170 万', ''],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            table.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('会不会亏？三个保障：\n')
    p.add_run('1. 无保底条款 - 按实际人次分账（30 元/台），客流少分账少，无固定成本压力\n')
    p.add_run('2. 学习强国背书 - 官方平台引流，客流有保障\n')
    p.add_run('3. 河南台验证 - 同样模式，人家年营收 2000 万+，我们保守预估 900-1200 万')
    
    quote = doc.add_paragraph()
    quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote.italic = True
    quote.add_run('"一句话：170 万投资，4-7 个月回本，亏不了，大概率赚。"')
    
    # 问题 3
    doc.add_heading('问题 3：有什么风险？', level=2)
    p = doc.add_paragraph()
    p.add_run('答：四类风险，都有防控措施：\n\n')
    
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['风险', '等级', '防控措施']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    
    data = [
        ['政治风险', '低', '学习强国联合出品 + 官方媒体资质，内容提前送审'],
        ['财务风险', '低', '无保底条款 + 按人次分账 + 170 万低投资'],
        ['技术风险', '中', '成熟供应商 + 备用设备 + 运维团队'],
        ['运营风险', '中', '多渠道营销 + 学校/企事业合作 + 学习强国引流'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            table.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    quote = doc.add_paragraph()
    quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote.italic = True
    quote.add_run('"一句话：项目风险可控，不做风险更大。"')
    
    # 问题 4
    doc.add_heading('问题 4：谁能做成？', level=2)
    p = doc.add_paragraph()
    p.add_run('答：三方合作，各司其职：\n\n')
    
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['角色', '合作方', '职责']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    
    data = [
        ['甲方', '学习强国学习平台', '开设专题专区、协调政府资源、宣传推广、政策支持'],
        ['乙方', '未来新视界科技（北京）', '提供内容 + 软硬件解决方案（30 台 VR 仅 18 万）+ 空间设计 + 运营管理系统'],
        ['丙方', '内蒙古广播电视台', '场地 + 运营团队 + 票务销售 + 设备运维 + 市场拓展'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            table.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    quote = doc.add_paragraph()
    quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote.italic = True
    quote.add_run('"一句话：学习强国背书 + 成熟技术方 + 我台落地执行，铁三角组合。"')
    
    # 问题 5
    doc.add_heading('问题 5：什么时候见效？', level=2)
    p = doc.add_paragraph()
    p.add_run('答：2026.12 首发，2027.03 正式运营，4-7 个月回本\n\n')
    
    p = doc.add_paragraph()
    p.add_run('关键里程碑：\n')
    p.add_run('• 2026.07：签署学习强国城市合伙人协议\n')
    p.add_run('• 2026.12：长征 90 周年首发（首波政绩）\n')
    p.add_run('• 2027.03：正式运营（持续收益）\n')
    p.add_run('• 2027.10：内蒙 80 周年自主研发（第二波政绩）')
    
    doc.add_paragraph()
    quote = doc.add_paragraph()
    quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote.italic = True
    quote.add_run('"一句话：7 个月建成，12 月首发，当年见效，次年丰收。"')
    
    doc.add_page_break()
    
    # 第二部分
    doc.add_heading('二、为什么必须做这个项目？', level=1)
    
    doc.add_heading('（一）政策要求：台内年度实施方案明确要求', level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['条款', '原文', '项目对应']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    
    data = [
        ['第 18 条', '围绕长征胜利 90 周年，以"历史影像 + 现实寻访 + 互动传播"模式传承红色基因', '《长征·英雄》VR 大空间体验区'],
        ['第 19 条', '谋划推出内蒙古自治区成立 80 周年宣传项目，高水平站位、高品质策划', 'AI 自主研发实验室（内蒙 80 周年主题）'],
        ['第 33 条', '组建 XR 技术团队，搭建 XR 演播室，通过外训与实战建立虚拟场景资源库', '数智创研中心 + XR 技术团队建设'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            table.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    quote = doc.add_paragraph()
    quote.italic = True
    quote.add_run('结论：本项目不是新增项目，是落实台内 2026 年度实施方案的具体行动。')
    
    doc.add_heading('（二）成功案例：河南台已验证模式可行', level=2)
    doc.add_paragraph('河南台《唐宫夜宴》VR 项目数据：\n')
    
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    data = [
        ['投入成本', '300 万元'],
        ['年营收', '2000 万+'],
        ['年客流', '50 万 + 人次'],
        ['ROI（3 年）', '1:6.7'],
        ['回本周期', '6-8 个月'],
        ['社会影响', '央视报道、文旅部典型案例、全国媒体学习标杆'],
    ]
    for i, row in enumerate(data):
        table.rows[i].cells[0].text = row[0]
        table.rows[i].cells[1].text = row[1]
    
    doc.add_paragraph()
    doc.add_paragraph('我台对比优势：\n')
    
    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'
    headers = ['维度', '河南台', '我台', '优势']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    
    data = [
        ['投入成本', '300 万', '170 万', '✅ 低 43%'],
        ['回本周期', '6-8 个月', '4-7 个月', '✅ 提前 2-3 个月'],
        ['ROI（3 年）', '1:6.7', '1:8.8', '✅ 高 60%'],
        ['IP 资源', '自有 IP', '学习强国联合', '✅ 政治背书更强'],
        ['政策窗口', '常态化', '双节点驱动', '✅ 更难得'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            table.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    quote = doc.add_paragraph()
    quote.italic = True
    quote.add_run('结论：河南台模式已验证可行，我台具备更强政治背书和成本优势，成功概率更高。')
    
    doc.add_heading('（三）时间窗口：双节点驱动，错过再等 10 年', level=2)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    headers = ['节点', '时间', '意义']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    
    data = [
        ['长征胜利 90 周年', '2026.10', '国家级纪念活动，政治热度最高'],
        ['内蒙古自治区成立 80 周年', '2027.10', '自治区最高规格庆典'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            table.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    quote = doc.add_paragraph()
    quote.italic = True
    quote.add_run('窗口期价值：政策支持力度最大、媒体关注度最高、客流自然增长最快、政绩显示度最强')
    doc.add_paragraph('结论：双节点叠加，10 年一遇，必须抓住。')
    
    doc.add_page_break()
    
    # 第三部分
    doc.add_heading('三、投资与回报（台长最关心）', level=1)
    
    doc.add_heading('（一）投资明细（一次性）', level=2)
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'
    headers = ['项目', '金额（万元）', '说明']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    
    data = [
        ['IP 授权费', '50', '学习强国官方授权，一次性支付'],
        ['硬件采购', '18', '30 台 VR 眼镜 + 配套（乙方报价确认）'],
        ['场地改造', '100', '旧台体育馆水电暖 + 声学改造（控制目标）'],
        ['总投资', '170', ''],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            table.rows[i+1].cells[j].text = cell_text
            if row[0] == '总投资':
                for run in table.rows[i+1].cells[j].paragraphs[0].runs:
                    run.bold = True
    
    doc.add_paragraph()
    quote = doc.add_paragraph()
    quote.italic = True
    quote.add_run('对比河南台：300 万 → 我台 170 万，省 130 万。')
    
    doc.add_heading('（二）收入预测（年）', level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['收入来源', '金额（万元）', '计算方式']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    
    data = [
        ['门票收入（散客）', '948', '158 元 × 6 万人次'],
        ['门票收入（团体）', '200', '100 元 × 2 万人次'],
        ['学生研学', '75', '150 元 × 5000 人次'],
        ['文创产品销售', '50', '利润全归我方'],
        ['合计', '1273', '保守预估'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            table.rows[i+1].cells[j].text = cell_text
    
    doc.add_heading('（三）成本结构（年）', level=2)
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    headers = ['成本项', '金额（万元）', '说明']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    
    data = [
        ['IP 使用费（分账）', '187.5', '30 元×6 万 + 15 元×0.5 万'],
        ['人力成本', '180', '15-20 人团队'],
        ['运维成本', '50', '设备维护、能耗'],
        ['营销推广', '80', '-'],
        ['其他', '25', '-'],
        ['合计', '522.5', '-'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            table.rows[i+1].cells[j].text = cell_text
    
    doc.add_heading('（四）投资回报', level=2)
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'
    headers = ['指标', '数值', '对比河南台']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    
    data = [
        ['年营收', '900-1200 万元', '河南台 2000 万+'],
        ['年成本', '522.5 万元', '-'],
        ['年利润', '330-630 万元', '-'],
        ['总投资', '170 万元', '河南台 300 万'],
        ['回本周期', '4-7 个月', '河南台 6-8 个月 ✅'],
        ['ROI（3 年）', '1:8.8', '河南台 1:6.7 ✅'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            table.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    quote = doc.add_paragraph()
    quote.italic = True
    quote.add_run('一句话：投得少，赚得多，回本快。')
    
    doc.add_page_break()
    
    # 第四部分
    doc.add_heading('四、风险与防控（台长最担心）', level=1)
    
    doc.add_heading('（一）政治风险（台长最关心）', level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['风险点', '风险等级', '防控措施']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    
    data = [
        ['内容审核', '低', '学习强国联合出品 + 官方媒体资质，内容提前送审'],
        ['意识形态', '低', '红色题材 + 长征主题，政治正确'],
        ['政策变化', '低', '紧跟政策导向，保持项目灵活性'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            table.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    quote = doc.add_paragraph()
    quote.italic = True
    quote.add_run('结论：政治风险极低，学习强国背书是最大保障。')
    
    doc.add_heading('（二）财务风险', level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'
    headers = ['风险点', '风险等级', '防控措施']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    
    data = [
        ['投资超支', '中', '分阶段实施，一期 170 万，二期用一期利润覆盖'],
        ['客流不足', '中', '多渠道营销 + 学校/企事业合作 + 学习强国引流'],
        ['分账压力', '低', '按实际人次结算，无保底条款'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            table.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    quote = doc.add_paragraph()
    quote.italic = True
    quote.add_run('结论：财务风险可控，无保底条款是核心保障。')
    
    doc.add_heading('（三）技术风险', level=2)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    headers = ['风险点', '风险等级', '防控措施']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    
    data = [
        ['设备稳定性', '中', '成熟供应商 + 备用设备 + 运维团队'],
        ['内容更新', '低', '学习强国持续供给 + 自主研发补充'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            table.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    quote = doc.add_paragraph()
    quote.italic = True
    quote.add_run('结论：技术风险中等，成熟方案可防控。')
    
    doc.add_heading('（四）运营风险', level=2)
    table = doc.add_table(rows=3, cols=3)
    table.style = 'Table Grid'
    headers = ['风险点', '风险等级', '防控措施']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    
    data = [
        ['人员培训', '低', '建立标准化培训体系，学习强国联合培训'],
        ['市场竞争', '低', '呼和浩特市独家授权，先发优势'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            table.rows[i+1].cells[j].text = cell_text
    
    doc.add_paragraph()
    quote = doc.add_paragraph()
    quote.italic = True
    quote.add_run('结论：运营风险可控，区域独家是核心优势。')
    
    doc.add_page_break()
    
    # 第五部分
    doc.add_heading('五、实施计划（台长最关注）', level=1)
    
    doc.add_heading('（一）时间节点', level=2)
    p = doc.add_paragraph()
    p.add_run('2026.05-07  筹备阶段 → 方案细化、基建评估、协议签署\n')
    p.add_run('2026.08-11  建设阶段 → 场馆改造、设备部署、内容引进\n')
    p.add_run('2026.12     首发节点 → 长征胜利 90 周年《长征·英雄》上线 ⭐⭐\n')
    p.add_run('2027.03     正式运营 → 全渠道上线、持续迭代\n')
    p.add_run('2027.10     自主研发 → 内蒙古自治区成立 80 周年内容上线 ⭐⭐')
    
    doc.add_heading('（二）关键里程碑', level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['时间', '里程碑', '意义']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    
    data = [
        ['2026.07', '签署学习强国城市合伙人协议', '锁定授权'],
        ['2026.11', '场馆改造完成', '硬件就绪'],
        ['2026.12', '长征 90 周年首发', '首波政绩 ⭐⭐'],
        ['2027.03', '正式运营', '持续收益'],
        ['2027.10', '内蒙 80 周年自主研发', '第二波政绩 ⭐⭐'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            table.rows[i+1].cells[j].text = cell_text
    
    doc.add_heading('（三）责任分工', level=2)
    table = doc.add_table(rows=6, cols=3)
    table.style = 'Table Grid'
    headers = ['工作模块', '责任部门', '配合部门']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
    
    data = [
        ['基建改造', '后勤保障部', '技术部'],
        ['设备采购', '技术部', '财务部'],
        ['学习强国对接', '内容部', '台领导'],
        ['自主研发', '数智创研中心', '技术部'],
        ['运营管理', '运营部', '市场部'],
    ]
    for i, row in enumerate(data):
        for j, cell_text in enumerate(row):
            table.rows[i+1].cells[j].text = cell_text
    
    doc.add_page_break()
    
    # 第六部分
    doc.add_heading('六、决策请求（台长拍板）', level=1)
    
    doc.add_heading('建议事项', level=2)
    p = doc.add_paragraph()
    p.add_run('1. 批准立项\n')
    p.add_run('   启动旧台体育馆基建评估\n\n')
    p.add_run('2. 成立专项工作组\n')
    p.add_run('   台领导牵头，统筹项目推进\n\n')
    p.add_run('3. 签署学习强国城市合伙人协议\n')
    p.add_run('   锁定呼和浩特市授权（50 万 IP 费 + 分账模式）\n\n')
    p.add_run('4. 以长征胜利 90 周年为首发节点\n')
    p.add_run('   2026.12《长征·英雄》VR 大空间上线\n\n')
    p.add_run('5. 同步启动自主研发团队建设\n')
    p.add_run('   为 2027 年内蒙 80 周年节点做准备')
    
    doc.add_paragraph()
    doc.add_heading('一句话总结', level=2)
    quote = doc.add_paragraph()
    quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote.italic = True
    quote.add_run('"170 万投资，4-7 个月回本，ROI 1:8.8，学习强国背书，双节点驱动，河南台已验证成功，不做是失职，做了是政绩。"')
    
    doc.add_page_break()
    
    # 附件清单
    doc.add_heading('附件清单', level=1)
    attachments = [
        '1. 《学习强国〈长征·英雄〉VR 大空间三方城市合伙人协议》',
        '2. 《XR 项目财务测算表》（详细版）',
        '3. 《场馆改造技术方案》',
        '4. 《设备采购清单及预算》',
        '5. 《河南电视台〈唐宫夜宴〉VR 项目案例分析报告》',
        '6. 《台内 2026 年度实施方案》（第 18、19、33 条）',
        '7. 《呼和浩特市市场调研报告》',
        '8. 《项目实施甘特图》',
    ]
    for att in attachments:
        doc.add_paragraph(att, style='List Bullet')
    
    # 页脚
    doc.add_paragraph()
    doc.add_paragraph('_' * 50)
    p = doc.add_paragraph()
    p.add_run('呈报人：强国小马（数字团队总监）\n')
    p.add_run('联合呈报：内蒙古 XR 项目团队\n')
    p.add_run('联系方式：钉钉私信\n')
    p.add_run('呈报日期：2026 年 4 月 17 日')
    
    # 保存文档
    doc.save('reports/XR 项目立项报告 - 台长决策版.docx')
    print('✅ Word 文档生成成功！')
    print('文件路径：/home/admin/.openclaw/workspace-chief-agent/reports/XR 项目立项报告 - 台长决策版.docx')

if __name__ == '__main__':
    create_report()
